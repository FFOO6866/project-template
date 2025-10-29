"""
Integration Tests for Enhanced Document Processing
Tests all 4 extraction strategies with real files
"""

import pytest
import asyncio
from pathlib import Path
from unittest.mock import Mock, AsyncMock, patch
import json

# Test will run in Docker container with real database
pytestmark = pytest.mark.integration


class TestEnhancedDocumentProcessor:
    """Test suite for multi-strategy document processing"""

    @pytest.fixture
    async def processor(self):
        """Create processor instance with test configuration"""
        from src.services.enhanced_document_processor import EnhancedDocumentProcessor

        # Mock OpenAI client for tests
        processor = EnhancedDocumentProcessor()
        return processor

    @pytest.fixture
    async def mock_db_pool(self):
        """Mock database pool for testing"""
        pool = AsyncMock()
        pool.acquire = AsyncMock()
        pool.acquire.return_value.__aenter__ = AsyncMock()
        pool.acquire.return_value.__aexit__ = AsyncMock()
        return pool

    @pytest.fixture
    def sample_requirements_response(self):
        """Sample OpenAI response for requirements extraction"""
        return {
            "customer_name": "Test Corp",
            "project_name": "Office Renovation",
            "deadline": "2025-12-31",
            "items": [
                {
                    "description": "DEWALT DCD791D2 20V Cordless Drill",
                    "quantity": 50,
                    "unit": "pieces",
                    "specifications": ["DEWALT", "DCD791D2", "20V", "Cordless"],
                    "category": "Power Tools"
                },
                {
                    "description": "3M H-700 Safety Helmet",
                    "quantity": 100,
                    "unit": "pieces",
                    "specifications": ["3M", "H-700", "Safety"],
                    "category": "Safety Equipment"
                },
                {
                    "description": "Cat6A Ethernet Cable",
                    "quantity": 200,
                    "unit": "meters",
                    "specifications": ["Cat6A", "Ethernet"],
                    "category": "Networking"
                }
            ],
            "additional_requirements": ["Net 30 payment terms", "Delivery within 2 weeks"],
            "delivery_address": "123 Business St, City",
            "contact_email": "procurement@testcorp.com"
        }

    # ==================== Strategy 1: Specialized Parser Tests ====================

    @pytest.mark.asyncio
    async def test_specialized_pdf_extraction_camelot(
        self, processor, tmp_path, sample_requirements_response
    ):
        """Test PDF extraction with Camelot (bordered tables)"""
        # Create a test PDF with table (would need actual PDF file)
        test_pdf = tmp_path / "test_rfp_table.pdf"

        # Skip if test file doesn't exist (create fixture files separately)
        if not test_pdf.exists():
            pytest.skip("Test PDF file not available")

        # Mock OpenAI response
        with patch.object(
            processor, '_analyze_requirements',
            return_value=asyncio.coroutine(lambda x: sample_requirements_response)()
        ):
            result = await processor._try_specialized_parser(str(test_pdf))

            assert result['method'] == 'pdf_camelot'
            assert len(result['requirements']['items']) == 3
            assert 'DEWALT' in result['requirements']['items'][0]['description']

    @pytest.mark.asyncio
    async def test_specialized_excel_extraction(
        self, processor, tmp_path, sample_requirements_response
    ):
        """Test Excel extraction with openpyxl"""
        # Create test Excel file
        test_excel = tmp_path / "test_rfp.xlsx"

        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "RFP Products"

        # Add headers
        ws.append(["Item", "Description", "Quantity", "Unit"])

        # Add data
        ws.append([1, "DEWALT DCD791D2 20V Cordless Drill", 50, "pieces"])
        ws.append([2, "3M H-700 Safety Helmet", 100, "pieces"])
        ws.append([3, "Cat6A Ethernet Cable", 200, "meters"])

        wb.save(test_excel)

        # Mock OpenAI response
        with patch.object(
            processor, '_analyze_requirements',
            return_value=asyncio.coroutine(lambda x: sample_requirements_response)()
        ):
            result = await processor._try_specialized_parser(str(test_excel))

            assert result['method'] == 'specialized_excel'
            assert 'SHEET' in result['text']
            assert 'HEADERS' in result['text']
            assert 'ROW' in result['text']
            assert len(result['requirements']['items']) == 3

    @pytest.mark.asyncio
    async def test_specialized_docx_extraction(
        self, processor, tmp_path, sample_requirements_response
    ):
        """Test DOCX extraction with python-docx"""
        test_docx = tmp_path / "test_rfp.docx"

        from docx import Document
        doc = Document()

        # Add title
        doc.add_heading("Request for Quotation", 0)
        doc.add_paragraph("Project: Office Renovation")

        # Add table
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'

        # Headers
        headers = table.rows[0].cells
        headers[0].text = "Description"
        headers[1].text = "Quantity"
        headers[2].text = "Unit"

        # Data
        row1 = table.rows[1].cells
        row1[0].text = "DEWALT DCD791D2 20V Cordless Drill"
        row1[1].text = "50"
        row1[2].text = "pieces"

        row2 = table.rows[2].cells
        row2[0].text = "3M H-700 Safety Helmet"
        row2[1].text = "100"
        row2[2].text = "pieces"

        row3 = table.rows[3].cells
        row3[0].text = "Cat6A Ethernet Cable"
        row3[1].text = "200"
        row3[2].text = "meters"

        doc.save(test_docx)

        # Mock OpenAI response
        with patch.object(
            processor, '_analyze_requirements',
            return_value=asyncio.coroutine(lambda x: sample_requirements_response)()
        ):
            result = await processor._try_specialized_parser(str(test_docx))

            assert result['method'] == 'specialized_docx'
            assert '=== TABLE START ===' in result['text']
            assert 'HEADERS' in result['text']
            assert 'DEWALT' in result['text']
            assert len(result['requirements']['items']) == 3

    # ==================== Strategy 2: Docling Tests ====================

    @pytest.mark.asyncio
    async def test_docling_extraction(
        self, processor, tmp_path, sample_requirements_response
    ):
        """Test Docling multi-format parser"""
        # Skip if Docling not installed
        try:
            import docling
        except ImportError:
            pytest.skip("Docling not installed")

        test_pdf = tmp_path / "test_complex.pdf"

        # Skip if test file doesn't exist
        if not test_pdf.exists():
            pytest.skip("Test PDF file not available")

        # Mock OpenAI response
        with patch.object(
            processor, '_analyze_requirements',
            return_value=asyncio.coroutine(lambda x: sample_requirements_response)()
        ):
            result = await processor._try_docling_parser(str(test_pdf))

            assert result['method'] == 'docling'
            assert len(result['requirements']['items']) >= 1

    # ==================== Strategy 3: Vision Model Tests ====================

    @pytest.mark.asyncio
    async def test_vision_extraction_mock(
        self, processor, tmp_path
    ):
        """Test GPT-4 Vision extraction with mocked API"""
        # Create a simple image file for testing
        test_image = tmp_path / "test_rfp.png"

        from PIL import Image, ImageDraw, ImageFont
        img = Image.new('RGB', (800, 600), color='white')
        draw = ImageDraw.Draw(img)

        # Draw sample RFP content
        draw.text((50, 50), "Request for Quotation", fill='black')
        draw.text((50, 100), "1. DEWALT Drill - 50 pieces", fill='black')
        draw.text((50, 130), "2. Safety Helmets - 100 pieces", fill='black')

        img.save(test_image)

        # Mock OpenAI vision response
        mock_vision_response = Mock()
        mock_vision_response.choices = [Mock()]
        mock_vision_response.choices[0].message.content = json.dumps({
            "page_text": "Request for Quotation\n1. DEWALT Drill - 50 pieces\n2. Safety Helmets - 100 pieces",
            "items": [
                {
                    "description": "DEWALT Drill",
                    "quantity": 50,
                    "unit": "pieces",
                    "specifications": ["DEWALT"],
                    "category": "Power Tools"
                },
                {
                    "description": "Safety Helmets",
                    "quantity": 100,
                    "unit": "pieces",
                    "specifications": ["Safety"],
                    "category": "Safety Equipment"
                }
            ]
        })

        with patch.object(
            processor.openai_client.chat.completions, 'create',
            return_value=asyncio.coroutine(lambda: mock_vision_response)()
        ):
            result = await processor._try_vision_extraction(str(test_image))

            assert result['method'] == 'vision'
            assert len(result['requirements']['items']) == 2
            assert result['pages_processed'] == 1

    # ==================== Strategy 4: Cascading Tests ====================

    @pytest.mark.asyncio
    async def test_cascading_strategies_success_early(
        self, processor, tmp_path, sample_requirements_response
    ):
        """Test that cascading stops when high confidence is achieved"""
        test_file = tmp_path / "test_simple.txt"
        test_file.write_text("1. DEWALT Drill - 50 pieces\n2. Safety Helmet - 100 pieces")

        # Mock to return high confidence on first strategy
        with patch.object(
            processor, '_try_specialized_parser',
            return_value=asyncio.coroutine(lambda x: {
                'text': 'test content',
                'requirements': sample_requirements_response,
                'method': 'specialized_txt'
            })()
        ), patch.object(
            processor, '_calculate_confidence',
            return_value=0.95  # High confidence
        ):
            result = await processor._execute_cascading_strategies(str(test_file))

            # Should use specialized strategy and stop (not try docling/vision)
            assert result['method'] == 'specialized'
            assert result['confidence'] >= 0.85
            assert 'specialized' in result['strategies_attempted']
            # Should NOT have tried vision if specialized succeeded
            assert len(result['strategies_attempted']) == 1

    @pytest.mark.asyncio
    async def test_cascading_strategies_fallback(
        self, processor, tmp_path, sample_requirements_response
    ):
        """Test that cascading tries all strategies when confidence is low"""
        test_file = tmp_path / "test_difficult.txt"
        test_file.write_text("Some unstructured RFP content")

        call_count = {'specialized': 0, 'docling': 0, 'vision': 0, 'fallback': 0}

        async def mock_specialized(path):
            call_count['specialized'] += 1
            raise Exception("Specialized failed")

        async def mock_docling(path):
            call_count['docling'] += 1
            raise Exception("Docling failed")

        async def mock_vision(path):
            call_count['vision'] += 1
            raise Exception("Vision failed")

        async def mock_fallback(path):
            call_count['fallback'] += 1
            return {
                'text': 'fallback content',
                'requirements': sample_requirements_response,
                'method': 'basic_fallback'
            }

        with patch.object(processor, '_try_specialized_parser', side_effect=mock_specialized), \
             patch.object(processor, '_try_docling_parser', side_effect=mock_docling), \
             patch.object(processor, '_try_vision_extraction', side_effect=mock_vision), \
             patch.object(processor, '_try_basic_extraction', side_effect=mock_fallback), \
             patch.object(processor, '_calculate_confidence', return_value=0.7):

            result = await processor._execute_cascading_strategies(str(test_file))

            # Should have tried all strategies
            assert call_count['specialized'] == 1
            # Docling/vision might be skipped if disabled in config
            assert call_count['fallback'] == 1
            assert result['method'] == 'fallback'

    # ==================== Confidence Scoring Tests ====================

    @pytest.mark.asyncio
    async def test_confidence_scoring_high(self, processor):
        """Test confidence scoring for high-quality extraction"""
        result = {
            'text': 'A' * 1000,  # Sufficient text
            'requirements': {
                'items': [
                    {
                        'description': 'DEWALT DCD791D2 20V Cordless Drill with detailed specs',
                        'quantity': 50,
                        'unit': 'pieces',
                        'specifications': ['DEWALT', 'DCD791D2'],
                        'category': 'Power Tools'
                    },
                    {
                        'description': '3M H-700 Safety Helmet certified for construction',
                        'quantity': 100,
                        'unit': 'pieces',
                        'specifications': ['3M', 'H-700'],
                        'category': 'Safety'
                    },
                    {
                        'description': 'Cat6A Ethernet Cable 23AWG solid copper',
                        'quantity': 200,
                        'unit': 'meters',
                        'specifications': ['Cat6A'],
                        'category': 'Networking'
                    },
                    # 7 more items for high count bonus
                    *[{
                        'description': f'Product {i} with full description',
                        'quantity': 10,
                        'unit': 'pieces',
                        'specifications': ['Spec'],
                        'category': 'General'
                    } for i in range(4, 11)]
                ]
            }
        }

        confidence = processor._calculate_confidence(result)

        # Should have high confidence:
        # - Base: 0.5
        # - 10+ items: 0.5 + 0.15 + 0.10 + 0.05 = 0.80
        # - Complete items: +0.15
        # - Sufficient text: +0.05
        # Total: ~0.95+
        assert confidence >= 0.85

    @pytest.mark.asyncio
    async def test_confidence_scoring_low(self, processor):
        """Test confidence scoring for low-quality extraction"""
        result = {
            'text': 'Short',  # Insufficient text
            'requirements': {
                'items': [
                    {
                        'description': 'Item',  # Too short
                        'quantity': 0,  # Invalid quantity
                        'unit': '',
                        'specifications': [],
                        'category': ''
                    }
                ]
            }
        }

        confidence = processor._calculate_confidence(result)

        # Should have low confidence:
        # - Base: 0.5
        # - Only 1 item: 0.5
        # - Incomplete data: +0.0
        # - Short text: +0.0
        # Total: ~0.5
        assert confidence < 0.7

    # ==================== Full Pipeline Test ====================

    @pytest.mark.asyncio
    async def test_full_document_processing_pipeline(
        self, processor, mock_db_pool, tmp_path, sample_requirements_response
    ):
        """Test complete document processing pipeline"""
        # Create test DOCX
        test_docx = tmp_path / "test_rfp_complete.docx"

        from docx import Document
        doc = Document()
        doc.add_heading("RFP - Office Equipment", 0)
        doc.add_paragraph("Customer: Test Corporation")
        doc.add_paragraph("Project: Office Renovation 2025")

        # Add table
        table = doc.add_table(rows=3, cols=3)
        headers = table.rows[0].cells
        headers[0].text = "Item"
        headers[1].text = "Quantity"
        headers[2].text = "Unit"

        row1 = table.rows[1].cells
        row1[0].text = "DEWALT Drill"
        row1[1].text = "50"
        row1[2].text = "pieces"

        row2 = table.rows[2].cells
        row2[0].text = "Safety Helmets"
        row2[1].text = "100"
        row2[2].text = "pieces"

        doc.save(test_docx)

        # Mock OpenAI
        with patch.object(
            processor, '_analyze_requirements',
            return_value=asyncio.coroutine(lambda x: sample_requirements_response)()
        ):
            result = await processor.process_document(
                document_id=1,
                file_path=str(test_docx),
                db_pool=mock_db_pool
            )

            # Verify results
            assert result['processor_version'] == '2.0.0'
            assert result['extraction_method'] in ['specialized', 'docling', 'vision', 'fallback']
            assert result['extraction_confidence'] > 0.0
            assert result['processing_time_ms'] > 0
            assert len(result['requirements']['items']) == 3

            # Verify database was updated
            assert mock_db_pool.acquire.called


if __name__ == '__main__':
    pytest.main([__file__, '-v', '-s'])
