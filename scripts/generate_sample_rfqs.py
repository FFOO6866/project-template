#!/usr/bin/env python3
"""
Generate Sample RFQ Documents for Horme POV
============================================
Creates realistic Request for Quotation documents using real products from database

Generates:
- 50 PDF files
- 50 Excel files
- 50 Word (.docx) files

With realistic:
- Singapore company names
- Project descriptions
- Product lists from actual Horme catalog
"""

import os
import sys
import random
from datetime import datetime, timedelta
from pathlib import Path
import asyncio
import asyncpg
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from docx import Document as WordDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

# Fake but realistic-sounding Singapore company names
COMPANY_NAMES = [
    "Pinnacle Development Pte Ltd",
    "Meridian Facilities Management",
    "Skyline Construction Group",
    "Harbour View Commercial Trust",
    "Pacific Marine Engineering",
    "Atlas Building Consultants",
    "Gateway Port Services",
    "Apex Technologies Engineering",
    "Sterling Construction",
    "Golden Circle Group Ltd",
    "Phoenix Development Pte Ltd",
    "Eastpoint Property Management",
    "Hillside Estates",
    "Summit Holdings Construction",
    "Premier Land Management Services",
    "Horizon Property Development",
    "Silverline Facilities Group",
    "Cornerstone Construction",
    "Vertex Building Maintenance",
    "Diamond Construction Pte Ltd",
    "Riverside Contractors",
    "Metro Construction & Trading",
    "Century Development Corporation",
    "Paramount Holdings",
    "Dragon Gate Development",
    "Jade Realty Singapore",
    "Starlight Property",
    "Prosperity Group Limited",
    "Fortune Holdings Facilities",
    "Harmony Property Services",
    "Evergreen Land Pte Ltd",
    "Prestige Holdings Limited",
    "Zenith Asset Management",
    "Northstar Property Management",
    "Orchard Retail Procurement",
    "Metro Transit Maintenance",
    "Citylink Transport Engineering",
    "Premier Leisure Facilities",
    "Industrial Park Projects",
    "Housing Development Services",
    "Transit Solutions Authority",
    "Infrastructure Development Corp",
    "Innovation University",
    "Eastside University",
    "Riverside Business School",
    "Technical Institute Facilities",
    "Applied Sciences College",
    "Vocational Training Centre",
    "Education Services Division",
    "Central Medical Centre",
    "Westside Hospital Operations",
    "Elite Medical Services",
    "Premier Healthcare Management",
    "Family Care Hospital",
    "Wellness Institute",
    "National Sports Complex",
    "Waterfront Resort Procurement",
    "Island Paradise Resort",
    "Botanical Gardens Operations",
    "Wildlife Conservation Park",
    "Nature Discovery Centre",
    "Recreation Island Corporation",
    "Green Spaces Management",
    "Water Services Authority",
    "Environmental Protection Agency",
    "Emergency Response Division",
    "Defense Logistics Command",
    "Guardian Security Services",
    "Shield Protection Management",
    "Airport Services Group",
    "Aviation Hub Corporation",
    "Airways Engineering Company",
    "Airline Services Division",
    "Aerospace Maintenance Corp",
    "Regional Carrier Operations",
    "Budget Airlines Services",
    "First National Bank Property",
    "Metropolitan Bank Facilities",
    "United Overseas Banking Centre",
    "Eastern Bank Singapore",
    "International Banking Services",
    "Telecom Services Group",
    "Mobile Networks Division",
    "Communication Infrastructure",
    "Broadcasting Corporation",
    "Media Publishing Services",
    "Ride Sharing Singapore",
    "Food Delivery Logistics",
    "Online Shopping Warehousing",
    "eCommerce Regional Hub",
    "Marketplace Operations Centre",
    "Digital Services Facilities",
    "Gaming Technology HQ",
    "Tech Innovations Singapore",
    "Cloud Services APAC",
    "Social Media Operations",
    "Data Services Asia Pacific",
    "Software Solutions Campus",
    "Technology Southeast Asia",
    "Innovation Labs Facilities",
    "Consumer Goods Asia Pacific",
    "Regional Manufacturing Hub",
]

# Project types and descriptions
PROJECT_TYPES = [
    {
        "type": "Office Renovation",
        "description": "Complete office renovation including new partitions, flooring, and electrical work",
        "product_categories": ["power_tools", "cleaning", "safety", "hardware"]
    },
    {
        "type": "New Building Construction",
        "description": "Construction of new 12-storey commercial building with basement carpark",
        "product_categories": ["power_tools", "hardware", "safety", "bins"]
    },
    {
        "type": "Facilities Maintenance",
        "description": "Annual facilities maintenance contract for commercial property",
        "product_categories": ["cleaning", "bins", "safety", "power_tools"]
    },
    {
        "type": "Factory Setup",
        "description": "Setting up new manufacturing facility with clean room requirements",
        "product_categories": ["power_tools", "cleaning", "safety", "bins"]
    },
    {
        "type": "Retail Outlet Fitout",
        "description": "Complete fitout for new retail store including fixtures and displays",
        "product_categories": ["power_tools", "hardware", "cleaning"]
    },
    {
        "type": "Hospital Expansion",
        "description": "Addition of new wing to existing hospital facility",
        "product_categories": ["power_tools", "cleaning", "safety", "bins"]
    },
    {
        "type": "School Upgrading",
        "description": "Upgrading works for primary school facilities and classrooms",
        "product_categories": ["power_tools", "hardware", "cleaning", "bins"]
    },
    {
        "type": "Warehouse Refurbishment",
        "description": "Complete refurbishment of logistics warehouse including racking systems",
        "product_categories": ["power_tools", "hardware", "safety"]
    },
    {
        "type": "Hotel Renovation",
        "description": "Renovation of hotel rooms and public areas",
        "product_categories": ["power_tools", "hardware", "cleaning"]
    },
    {
        "type": "Data Centre Construction",
        "description": "Building new Tier III data centre facility",
        "product_categories": ["power_tools", "hardware", "safety", "bins"]
    },
]

# Contact person names (Singapore names)
CONTACT_NAMES = [
    ("Tan", "Wei Ming", "Mr"),
    ("Lim", "Hui Ling", "Ms"),
    ("Wong", "Kai Xiang", "Mr"),
    ("Lee", "Mei Ling", "Ms"),
    ("Ng", "Jun Hao", "Mr"),
    ("Ong", "Siew Lan", "Ms"),
    ("Chua", "Wei Jie", "Mr"),
    ("Teo", "Xin Yi", "Ms"),
    ("Koh", "Ming Wei", "Mr"),
    ("Goh", "Li Ying", "Ms"),
    ("Chong", "Zhi Hao", "Mr"),
    ("Low", "Ai Lin", "Ms"),
    ("Sim", "Jun Wei", "Mr"),
    ("Yeo", "Hui Min", "Ms"),
    ("Seah", "Wei Lun", "Mr"),
    ("Poh", "Mei Hui", "Ms"),
    ("Quek", "Jia Jun", "Mr"),
    ("Tay", "Shi Hui", "Ms"),
    ("Ho", "Wei Ren", "Mr"),
    ("Leong", "Xiao Ting", "Ms"),
]


class RFQGenerator:
    """Generate realistic RFQ documents"""

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        self.products = []

    async def fetch_products_from_db(self):
        """Fetch real products from PostgreSQL database"""
        print("Connecting to database...")

        # Get password from environment or use default
        import os
        from dotenv import load_dotenv

        load_dotenv('.env.production')

        conn = await asyncpg.connect(
            host='localhost',
            port=5432,
            user=os.getenv('POSTGRES_USER', 'horme_user'),
            password=os.getenv('POSTGRES_PASSWORD', 'horme_secure_password_2024'),
            database=os.getenv('POSTGRES_DB', 'horme_db')
        )

        # Get products with prices
        products = await conn.fetch("""
            SELECT name, brand, price, unit, sku, product_code
            FROM products
            WHERE price IS NOT NULL AND price > 0
            AND is_active = true
            ORDER BY RANDOM()
            LIMIT 500
        """)

        self.products = [dict(p) for p in products]
        print(f"Loaded {len(self.products)} products from database")

        await conn.close()

    def generate_rfq_data(self):
        """Generate realistic RFQ data"""
        company = random.choice(COMPANY_NAMES)
        project = random.choice(PROJECT_TYPES)
        contact = random.choice(CONTACT_NAMES)

        # Generate RFQ number
        date_code = datetime.now().strftime("%Y%m")
        rfq_number = f"RFQ-{date_code}-{random.randint(1000, 9999)}"

        # Random number of products (5-20)
        num_products = random.randint(5, 20)
        selected_products = random.sample(self.products, num_products)

        # Create line items
        line_items = []
        for idx, product in enumerate(selected_products, 1):
            quantity = random.choice([5, 10, 15, 20, 25, 30, 50, 100])
            line_items.append({
                "line_no": idx,
                "product_name": product['name'],
                "brand": product['brand'] or "N/A",
                "quantity": quantity,
                "unit": product['unit'] or "pieces",
                "remarks": ""
            })

        # Generate dates
        issue_date = datetime.now()
        due_date = issue_date + timedelta(days=random.randint(7, 21))
        delivery_date = due_date + timedelta(days=random.randint(14, 45))

        return {
            "rfq_number": rfq_number,
            "company": company,
            "contact_person": f"{contact[2]} {contact[0]} {contact[1]}",
            "email": f"{contact[1].lower().replace(' ', '.')}@{company.lower().replace(' ', '').replace('pteltd', '').replace('pte', '').replace('ltd', '')[:20]}.com.sg",
            "phone": f"+65 {random.randint(6000, 9999)} {random.randint(1000, 9999)}",
            "project_name": project["type"],
            "project_description": project["description"],
            "issue_date": issue_date.strftime("%d %B %Y"),
            "due_date": due_date.strftime("%d %B %Y"),
            "delivery_date": delivery_date.strftime("%d %B %Y"),
            "line_items": line_items,
            "terms": [
                "Payment terms: Net 30 days from invoice date",
                "Prices to be quoted in Singapore Dollars (SGD)",
                "Delivery to project site required",
                "GST to be stated separately",
                "Warranty period: Minimum 12 months from delivery"
            ]
        }

    def generate_pdf(self, data: dict, filename: str):
        """Generate PDF RFQ"""
        filepath = self.output_dir / filename

        doc = SimpleDocTemplate(str(filepath), pagesize=A4)
        story = []
        styles = getSampleStyleSheet()

        # Header
        header_style = ParagraphStyle(
            'Header',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#C41E3A'),
            spaceAfter=20
        )
        story.append(Paragraph("REQUEST FOR QUOTATION", header_style))

        # Company info
        story.append(Paragraph(f"<b>{data['company']}</b>", styles['Normal']))
        story.append(Paragraph(f"RFQ Number: {data['rfq_number']}", styles['Normal']))
        story.append(Paragraph(f"Date: {data['issue_date']}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))

        # Contact details
        contact_data = [
            ['Contact Person:', data['contact_person']],
            ['Email:', data['email']],
            ['Phone:', data['phone']],
        ]
        contact_table = Table(contact_data, colWidths=[2*inch, 4*inch])
        story.append(contact_table)
        story.append(Spacer(1, 0.3*inch))

        # Project details
        story.append(Paragraph(f"<b>Project: {data['project_name']}</b>", styles['Heading2']))
        story.append(Paragraph(data['project_description'], styles['Normal']))
        story.append(Spacer(1, 0.2*inch))

        # Line items table
        table_data = [['#', 'Product Description', 'Brand', 'Qty', 'Unit', 'Unit Price', 'Total']]
        for item in data['line_items']:
            table_data.append([
                str(item['line_no']),
                item['product_name'],
                item['brand'],
                str(item['quantity']),
                item['unit'],
                '',  # To be filled by supplier
                ''   # To be filled by supplier
            ])

        items_table = Table(table_data, colWidths=[0.4*inch, 2.5*inch, 1*inch, 0.6*inch, 0.6*inch, 0.9*inch, 0.9*inch])
        items_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))

        story.append(items_table)
        story.append(Spacer(1, 0.3*inch))

        # Terms
        story.append(Paragraph("<b>Terms and Conditions:</b>", styles['Heading3']))
        for term in data['terms']:
            story.append(Paragraph(f"• {term}", styles['Normal']))

        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"<b>Quotation Due Date:</b> {data['due_date']}", styles['Normal']))
        story.append(Paragraph(f"<b>Required Delivery Date:</b> {data['delivery_date']}", styles['Normal']))

        doc.build(story)
        print(f"Created PDF: {filename}")

    def generate_excel(self, data: dict, filename: str):
        """Generate Excel RFQ"""
        filepath = self.output_dir / filename

        wb = Workbook()
        ws = wb.active
        ws.title = "RFQ"

        # Header
        ws['A1'] = "REQUEST FOR QUOTATION"
        ws['A1'].font = Font(size=18, bold=True, color="C41E3A")
        ws.merge_cells('A1:G1')

        # RFQ info
        row = 3
        ws[f'A{row}'] = data['company']
        ws[f'A{row}'].font = Font(bold=True)
        row += 1
        ws[f'A{row}'] = f"RFQ Number: {data['rfq_number']}"
        row += 1
        ws[f'A{row}'] = f"Date: {data['issue_date']}"
        row += 2

        # Contact
        ws[f'A{row}'] = "Contact Person:"
        ws[f'B{row}'] = data['contact_person']
        row += 1
        ws[f'A{row}'] = "Email:"
        ws[f'B{row}'] = data['email']
        row += 1
        ws[f'A{row}'] = "Phone:"
        ws[f'B{row}'] = data['phone']
        row += 2

        # Project
        ws[f'A{row}'] = "Project:"
        ws[f'B{row}'] = data['project_name']
        ws[f'B{row}'].font = Font(bold=True)
        row += 1
        ws[f'A{row}'] = "Description:"
        ws[f'B{row}'] = data['project_description']
        row += 2

        # Line items
        headers = ['#', 'Product Description', 'Brand', 'Qty', 'Unit', 'Unit Price', 'Total']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row, col, header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")

        row += 1
        for item in data['line_items']:
            ws.cell(row, 1, item['line_no'])
            ws.cell(row, 2, item['product_name'])
            ws.cell(row, 3, item['brand'])
            ws.cell(row, 4, item['quantity'])
            ws.cell(row, 5, item['unit'])
            row += 1

        # Terms
        row += 1
        ws[f'A{row}'] = "Terms and Conditions:"
        ws[f'A{row}'].font = Font(bold=True)
        row += 1
        for term in data['terms']:
            ws[f'A{row}'] = f"• {term}"
            row += 1

        row += 1
        ws[f'A{row}'] = f"Quotation Due Date: {data['due_date']}"
        ws[f'A{row}'].font = Font(bold=True)
        row += 1
        ws[f'A{row}'] = f"Required Delivery Date: {data['delivery_date']}"
        ws[f'A{row}'].font = Font(bold=True)

        # Column widths
        ws.column_dimensions['A'].width = 5
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 8
        ws.column_dimensions['E'].width = 10
        ws.column_dimensions['F'].width = 12
        ws.column_dimensions['G'].width = 12

        wb.save(filepath)
        print(f"Created Excel: {filename}")

    def generate_word(self, data: dict, filename: str):
        """Generate Word DOCX RFQ"""
        filepath = self.output_dir / filename

        doc = WordDocument()

        # Header
        heading = doc.add_heading('REQUEST FOR QUOTATION', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.color.rgb = RGBColor(196, 30, 58)

        # Company info
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(data['company']).bold = True
        doc.add_paragraph(f"RFQ Number: {data['rfq_number']}")
        doc.add_paragraph(f"Date: {data['issue_date']}")

        doc.add_paragraph()

        # Contact
        doc.add_paragraph(f"Contact Person: {data['contact_person']}")
        doc.add_paragraph(f"Email: {data['email']}")
        doc.add_paragraph(f"Phone: {data['phone']}")

        doc.add_paragraph()

        # Project
        p = doc.add_paragraph()
        p.add_run('Project: ').bold = True
        p.add_run(data['project_name'])

        doc.add_paragraph(f"Description: {data['project_description']}")

        doc.add_paragraph()

        # Line items table
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        headers = ['#', 'Product Description', 'Brand', 'Qty', 'Unit', 'Unit Price', 'Total']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data rows
        for item in data['line_items']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['line_no'])
            row_cells[1].text = item['product_name']
            row_cells[2].text = item['brand']
            row_cells[3].text = str(item['quantity'])
            row_cells[4].text = item['unit']
            row_cells[5].text = ""
            row_cells[6].text = ""

        doc.add_paragraph()

        # Terms
        p = doc.add_paragraph()
        p.add_run('Terms and Conditions:').bold = True

        for term in data['terms']:
            doc.add_paragraph(f"• {term}", style='List Bullet')

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(f"Quotation Due Date: {data['due_date']}").bold = True

        p = doc.add_paragraph()
        p.add_run(f"Required Delivery Date: {data['delivery_date']}").bold = True

        doc.save(filepath)
        print(f"Created Word: {filename}")

    async def generate_all_documents(self, num_pdf=50, num_excel=50, num_word=50):
        """Generate all sample RFQ documents"""

        # Fetch products from database
        await self.fetch_products_from_db()

        print(f"\nGenerating {num_pdf} PDF documents...")
        for i in range(num_pdf):
            data = self.generate_rfq_data()
            filename = f"RFQ_{data['rfq_number'].replace('-', '_')}_{data['company'][:30].replace(' ', '_').replace('(', '').replace(')', '')}.pdf"
            self.generate_pdf(data, filename)

        print(f"\nGenerating {num_excel} Excel documents...")
        for i in range(num_excel):
            data = self.generate_rfq_data()
            filename = f"RFQ_{data['rfq_number'].replace('-', '_')}_{data['company'][:30].replace(' ', '_').replace('(', '').replace(')', '')}.xlsx"
            self.generate_excel(data, filename)

        print(f"\nGenerating {num_word} Word documents...")
        for i in range(num_word):
            data = self.generate_rfq_data()
            filename = f"RFQ_{data['rfq_number'].replace('-', '_')}_{data['company'][:30].replace(' ', '_').replace('(', '').replace(')', '')}.docx"
            self.generate_word(data, filename)

        print(f"\nCOMPLETE!")
        print(f"Generated {num_pdf + num_excel + num_word} RFQ documents")
        print(f"Location: {self.output_dir}")
        print(f"- {num_pdf} PDF files")
        print(f"- {num_excel} Excel files")
        print(f"- {num_word} Word files")


async def main():
    """Main entry point"""
    output_dir = r"C:\Users\fujif\OneDrive\Desktop\Business Development\Horme\Sample Data"

    print("=" * 80)
    print("HORME POV - Sample RFQ Generator")
    print("=" * 80)
    print(f"Output directory: {output_dir}")
    print()

    generator = RFQGenerator(output_dir)

    try:
        await generator.generate_all_documents(
            num_pdf=50,
            num_excel=50,
            num_word=50
        )
    except Exception as e:
        print(f"\nError: {e}")
        import traceback
        traceback.print_exc()
        return 1

    return 0


if __name__ == "__main__":
    sys.exit(asyncio.run(main()))
