"""
Generate 30 Request for Quotation (RFQ) documents
- 15 Excel files
- 15 Word files
- Random products from Horme's catalog
- Realistic fake companies
- 10-100 items per RFQ
"""

import random
import os
from datetime import datetime, timedelta
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Horme product categories with realistic items
HORME_PRODUCTS = [
    # Safety Equipment
    {"name": "Safety Helmet - White", "category": "Personal Protective Equipment", "unit": "piece", "price_range": (15, 35)},
    {"name": "Safety Helmet - Yellow", "category": "Personal Protective Equipment", "unit": "piece", "price_range": (15, 35)},
    {"name": "Safety Goggles Clear Lens", "category": "Personal Protective Equipment", "unit": "piece", "price_range": (8, 18)},
    {"name": "Face Shield with Frame", "category": "Personal Protective Equipment", "unit": "piece", "price_range": (12, 25)},
    {"name": "Nitrile Gloves - M", "category": "Personal Protective Equipment", "unit": "box", "price_range": (18, 35)},
    {"name": "Nitrile Gloves - L", "category": "Personal Protective Equipment", "unit": "box", "price_range": (18, 35)},
    {"name": "Cotton Work Gloves", "category": "Personal Protective Equipment", "unit": "pair", "price_range": (2, 5)},
    {"name": "Leather Welding Gloves", "category": "Personal Protective Equipment", "unit": "pair", "price_range": (15, 28)},
    {"name": "Safety Boots Size 9", "category": "Personal Protective Equipment", "unit": "pair", "price_range": (45, 85)},
    {"name": "Safety Boots Size 10", "category": "Personal Protective Equipment", "unit": "pair", "price_range": (45, 85)},
    {"name": "High Visibility Vest", "category": "Personal Protective Equipment", "unit": "piece", "price_range": (8, 15)},
    {"name": "Ear Plugs Disposable", "category": "Personal Protective Equipment", "unit": "pair", "price_range": (0.5, 2)},
    {"name": "Ear Muffs Industrial", "category": "Personal Protective Equipment", "unit": "piece", "price_range": (25, 45)},
    {"name": "Respirator Half Face", "category": "Personal Protective Equipment", "unit": "piece", "price_range": (35, 65)},
    {"name": "N95 Respirator Mask", "category": "Personal Protective Equipment", "unit": "piece", "price_range": (2, 5)},

    # Hand Tools
    {"name": "Claw Hammer 16oz", "category": "Hand Tools", "unit": "piece", "price_range": (12, 25)},
    {"name": "Ball Peen Hammer 24oz", "category": "Hand Tools", "unit": "piece", "price_range": (15, 30)},
    {"name": "Screwdriver Set 6pc", "category": "Hand Tools", "unit": "set", "price_range": (18, 35)},
    {"name": "Phillips Screwdriver #2", "category": "Hand Tools", "unit": "piece", "price_range": (5, 12)},
    {"name": "Flathead Screwdriver 6mm", "category": "Hand Tools", "unit": "piece", "price_range": (5, 12)},
    {"name": "Adjustable Wrench 8\"", "category": "Hand Tools", "unit": "piece", "price_range": (15, 28)},
    {"name": "Adjustable Wrench 12\"", "category": "Hand Tools", "unit": "piece", "price_range": (22, 40)},
    {"name": "Combination Wrench Set", "category": "Hand Tools", "unit": "set", "price_range": (45, 95)},
    {"name": "Socket Set 1/2\" Drive", "category": "Hand Tools", "unit": "set", "price_range": (65, 150)},
    {"name": "Pliers Long Nose 8\"", "category": "Hand Tools", "unit": "piece", "price_range": (12, 25)},
    {"name": "Pliers Combination 8\"", "category": "Hand Tools", "unit": "piece", "price_range": (15, 28)},
    {"name": "Wire Cutter 6\"", "category": "Hand Tools", "unit": "piece", "price_range": (18, 35)},
    {"name": "Utility Knife Heavy Duty", "category": "Hand Tools", "unit": "piece", "price_range": (8, 15)},
    {"name": "Tape Measure 5m", "category": "Hand Tools", "unit": "piece", "price_range": (12, 22)},
    {"name": "Spirit Level 24\"", "category": "Hand Tools", "unit": "piece", "price_range": (25, 45)},

    # Power Tools
    {"name": "Drill Driver Cordless 12V", "category": "Power Tools", "unit": "piece", "price_range": (85, 150)},
    {"name": "Impact Driver 18V", "category": "Power Tools", "unit": "piece", "price_range": (120, 220)},
    {"name": "Angle Grinder 4.5\"", "category": "Power Tools", "unit": "piece", "price_range": (65, 125)},
    {"name": "Circular Saw 7.25\"", "category": "Power Tools", "unit": "piece", "price_range": (95, 180)},
    {"name": "Jigsaw 550W", "category": "Power Tools", "unit": "piece", "price_range": (75, 140)},
    {"name": "Rotary Hammer Drill", "category": "Power Tools", "unit": "piece", "price_range": (150, 350)},
    {"name": "Belt Sander 3\"x21\"", "category": "Power Tools", "unit": "piece", "price_range": (85, 165)},
    {"name": "Orbital Sander 5\"", "category": "Power Tools", "unit": "piece", "price_range": (65, 120)},

    # Construction Materials
    {"name": "Portland Cement 50kg", "category": "Construction Materials", "unit": "bag", "price_range": (12, 18)},
    {"name": "Sand Washed 1 Tonne", "category": "Construction Materials", "unit": "tonne", "price_range": (35, 55)},
    {"name": "Aggregate 20mm 1 Tonne", "category": "Construction Materials", "unit": "tonne", "price_range": (40, 60)},
    {"name": "Ready Mix Concrete", "category": "Construction Materials", "unit": "m³", "price_range": (120, 180)},
    {"name": "Bricks Red Clay", "category": "Construction Materials", "unit": "piece", "price_range": (0.5, 1.2)},
    {"name": "Concrete Blocks 6\"", "category": "Construction Materials", "unit": "piece", "price_range": (2.5, 4.5)},
    {"name": "Steel Rebar 10mm", "category": "Construction Materials", "unit": "meter", "price_range": (3, 6)},
    {"name": "Steel Rebar 12mm", "category": "Construction Materials", "unit": "meter", "price_range": (4, 7)},

    # Electrical
    {"name": "Cable 2.5mm² Single Core", "category": "Electrical", "unit": "meter", "price_range": (1.2, 2.5)},
    {"name": "Cable 4mm² Twin Core", "category": "Electrical", "unit": "meter", "price_range": (2.5, 4.5)},
    {"name": "LED Bulb 9W E27", "category": "Electrical", "unit": "piece", "price_range": (3, 8)},
    {"name": "LED Tube Light 18W", "category": "Electrical", "unit": "piece", "price_range": (8, 15)},
    {"name": "Switch Socket 13A", "category": "Electrical", "unit": "piece", "price_range": (5, 12)},
    {"name": "Circuit Breaker 32A", "category": "Electrical", "unit": "piece", "price_range": (15, 30)},
    {"name": "Distribution Board 8-Way", "category": "Electrical", "unit": "piece", "price_range": (45, 85)},
    {"name": "Cable Trunking 50x50mm", "category": "Electrical", "unit": "meter", "price_range": (8, 15)},

    # Plumbing
    {"name": "PVC Pipe 1/2\" 3m", "category": "Plumbing", "unit": "piece", "price_range": (5, 10)},
    {"name": "PVC Pipe 1\" 3m", "category": "Plumbing", "unit": "piece", "price_range": (8, 15)},
    {"name": "PVC Elbow 90° 1/2\"", "category": "Plumbing", "unit": "piece", "price_range": (0.8, 2)},
    {"name": "PVC T-Joint 1\"", "category": "Plumbing", "unit": "piece", "price_range": (1.5, 3)},
    {"name": "Ball Valve Brass 1/2\"", "category": "Plumbing", "unit": "piece", "price_range": (8, 15)},
    {"name": "Gate Valve 1\"", "category": "Plumbing", "unit": "piece", "price_range": (15, 28)},
    {"name": "Flexible Hose 1/2\"", "category": "Plumbing", "unit": "piece", "price_range": (6, 12)},
    {"name": "Sink Mixer Tap", "category": "Plumbing", "unit": "piece", "price_range": (45, 95)},

    # Fasteners
    {"name": "Hex Bolt M10x50mm", "category": "Fasteners", "unit": "piece", "price_range": (0.3, 0.8)},
    {"name": "Hex Nut M10", "category": "Fasteners", "unit": "piece", "price_range": (0.1, 0.3)},
    {"name": "Washer M10", "category": "Fasteners", "unit": "piece", "price_range": (0.05, 0.15)},
    {"name": "Wood Screw 3\"x10", "category": "Fasteners", "unit": "box", "price_range": (5, 12)},
    {"name": "Self-Drilling Screw 6x25mm", "category": "Fasteners", "unit": "box", "price_range": (8, 15)},
    {"name": "Anchor Bolt M12x100mm", "category": "Fasteners", "unit": "piece", "price_range": (1.5, 3.5)},

    # Adhesives & Sealants
    {"name": "Silicone Sealant Clear", "category": "Adhesives & Sealants", "unit": "tube", "price_range": (5, 10)},
    {"name": "PU Foam Expanding", "category": "Adhesives & Sealants", "unit": "can", "price_range": (8, 15)},
    {"name": "Epoxy Adhesive 2-Part", "category": "Adhesives & Sealants", "unit": "set", "price_range": (12, 25)},
    {"name": "Construction Adhesive", "category": "Adhesives & Sealants", "unit": "tube", "price_range": (6, 12)},
    {"name": "Waterproofing Membrane", "category": "Adhesives & Sealants", "unit": "roll", "price_range": (45, 85)},
]

# Realistic Singapore companies
COMPANIES = [
    {
        "name": "Singapore Construction & Engineering Pte Ltd",
        "address": "15 Tuas Avenue 5, Singapore 639340",
        "contact_person": "Mr. Tan Wei Ming",
        "phone": "+65 6861 2345",
        "email": "procurement@scepl.com.sg",
        "registration": "201523456A"
    },
    {
        "name": "Metro Building Solutions Pte Ltd",
        "address": "32 Joo Koon Circle, Singapore 629055",
        "contact_person": "Ms. Linda Chen",
        "phone": "+65 6265 7890",
        "email": "linda.chen@metrobuilding.sg",
        "registration": "201634567B"
    },
    {
        "name": "Prime Infrastructure Development Pte Ltd",
        "address": "8 Pioneer Walk, Singapore 627753",
        "contact_person": "Mr. Kumar Rajesh",
        "phone": "+65 6795 4321",
        "email": "rajesh@primeinfra.com.sg",
        "registration": "201745678C"
    },
    {
        "name": "Apex Engineering & Construction Pte Ltd",
        "address": "50 Gul Road, Singapore 629351",
        "contact_person": "Mr. David Lim",
        "phone": "+65 6861 5555",
        "email": "david.lim@apexeng.sg",
        "registration": "201856789D"
    },
    {
        "name": "Golden Dragon Building Materials Pte Ltd",
        "address": "120 Old Toh Tuck Road, Singapore 597654",
        "contact_person": "Ms. Wang Mei Ling",
        "phone": "+65 6469 8888",
        "email": "meiling.wang@goldendragon.com.sg",
        "registration": "201967890E"
    },
    {
        "name": "Sunrise Contractors & Builders Pte Ltd",
        "address": "23 Woodlands Industrial Park E1, Singapore 757720",
        "contact_person": "Mr. Ahmad bin Hassan",
        "phone": "+65 6752 3456",
        "email": "ahmad@sunrisecontractors.sg",
        "registration": "202078901F"
    },
    {
        "name": "Pacific Alliance Engineering Pte Ltd",
        "address": "45 Jurong Port Road, Singapore 619110",
        "contact_person": "Mr. Jason Ng",
        "phone": "+65 6266 7777",
        "email": "jason.ng@pacificalliance.sg",
        "registration": "202189012G"
    },
    {
        "name": "Excellence Construction Group Pte Ltd",
        "address": "88 Changi North Crescent, Singapore 499613",
        "contact_person": "Ms. Siti Nurhaliza",
        "phone": "+65 6543 9999",
        "email": "siti@excellencecg.com.sg",
        "registration": "202290123H"
    },
    {
        "name": "Skyline Development Corporation Pte Ltd",
        "address": "12 Benoi Crescent, Singapore 629977",
        "contact_person": "Mr. Richard Tan",
        "phone": "+65 6861 4444",
        "email": "richard@skylinedev.sg",
        "registration": "201401234I"
    },
    {
        "name": "United Building Systems Pte Ltd",
        "address": "67 Sungei Kadut Loop, Singapore 729511",
        "contact_person": "Ms. Jessica Lee",
        "phone": "+65 6368 6666",
        "email": "jessica.lee@unitedbuild.com.sg",
        "registration": "201512345J"
    },
    {
        "name": "Titan Engineering Works Pte Ltd",
        "address": "25 Tuas South Avenue 3, Singapore 637342",
        "contact_person": "Mr. Lim Hock Seng",
        "phone": "+65 6863 2222",
        "email": "hockseng@titaneng.sg",
        "registration": "201623456K"
    },
    {
        "name": "Premier Facilities Management Pte Ltd",
        "address": "156 Gul Circle, Singapore 629607",
        "contact_person": "Ms. Catherine Wong",
        "phone": "+65 6896 3333",
        "email": "catherine@premierfm.com.sg",
        "registration": "201734567L"
    },
    {
        "name": "Harmony Infrastructure Pte Ltd",
        "address": "78 Loyang Way, Singapore 508724",
        "contact_person": "Mr. Ravi Shankar",
        "phone": "+65 6542 5555",
        "email": "ravi@harmonyinfra.sg",
        "registration": "201845678M"
    },
    {
        "name": "Global Construction Partners Pte Ltd",
        "address": "90 Pioneer Road, Singapore 639951",
        "contact_person": "Mr. William Chen",
        "phone": "+65 6261 7878",
        "email": "william@globalcp.com.sg",
        "registration": "201956789N"
    },
    {
        "name": "Evergreen Building Solutions Pte Ltd",
        "address": "34 Tuas West Road, Singapore 638383",
        "contact_person": "Ms. Mary Tan",
        "phone": "+65 6898 9090",
        "email": "mary.tan@evergreenbuild.sg",
        "registration": "202067890O"
    }
]

# Additional companies for variety
ADDITIONAL_COMPANIES = [
    {
        "name": "Phoenix Engineering Consultants Pte Ltd",
        "address": "18 Boon Lay Way, Singapore 609966",
        "contact_person": "Mr. Benjamin Goh",
        "phone": "+65 6795 1234",
        "email": "benjamin@phoenixeng.sg",
        "registration": "202178901P"
    },
    {
        "name": "Summit Property Developments Pte Ltd",
        "address": "42 Penjuru Lane, Singapore 609189",
        "contact_person": "Ms. Rachel Koh",
        "phone": "+65 6271 5678",
        "email": "rachel.koh@summitprop.com.sg",
        "registration": "202289012Q"
    },
    {
        "name": "Elite Maintenance Services Pte Ltd",
        "address": "56 Mandai Estate, Singapore 729942",
        "contact_person": "Mr. Faizal Rahman",
        "phone": "+65 6269 8765",
        "email": "faizal@elitemaint.sg",
        "registration": "201390123R"
    },
    {
        "name": "Starlight Construction & Trading Pte Ltd",
        "address": "73 Tagore Lane, Singapore 787517",
        "contact_person": "Ms. Emily Chong",
        "phone": "+65 6453 4321",
        "email": "emily@starlightct.com.sg",
        "registration": "201401234S"
    },
    {
        "name": "Diamond Industrial Supplies Pte Ltd",
        "address": "29 Kaki Bukit Road 3, Singapore 417837",
        "contact_person": "Mr. Henry Yeo",
        "phone": "+65 6745 6789",
        "email": "henry.yeo@diamondindustrial.sg",
        "registration": "201512345T"
    },
    {
        "name": "Neptune Marine & Offshore Pte Ltd",
        "address": "61 Shipyard Road, Singapore 628249",
        "contact_person": "Ms. Sharon Tay",
        "phone": "+65 6265 9876",
        "email": "sharon@neptunemarine.sg",
        "registration": "201623456U"
    },
    {
        "name": "Atlas Engineering Solutions Pte Ltd",
        "address": "14 Senoko South Road, Singapore 758091",
        "contact_person": "Mr. Muthu Kumar",
        "phone": "+65 6755 4567",
        "email": "muthu@atlaseng.com.sg",
        "registration": "201734567V"
    },
    {
        "name": "Bright Future Contractors Pte Ltd",
        "address": "95 Ubi Road 4, Singapore 408612",
        "contact_person": "Ms. Grace Lim",
        "phone": "+65 6844 3210",
        "email": "grace@brightfuture.sg",
        "registration": "201845678W"
    },
    {
        "name": "Reliable Building Services Pte Ltd",
        "address": "38 Jalan Buroh, Singapore 619479",
        "contact_person": "Mr. Peter Ong",
        "phone": "+65 6268 7654",
        "email": "peter.ong@reliablebuild.com.sg",
        "registration": "201956789X"
    },
    {
        "name": "Zenith Property Management Pte Ltd",
        "address": "52 Toh Guan Road East, Singapore 608586",
        "contact_person": "Ms. Angela Foo",
        "phone": "+65 6896 5432",
        "email": "angela@zenithpm.sg",
        "registration": "202067890Y"
    },
    {
        "name": "Victory Construction Materials Pte Ltd",
        "address": "27 Defu Lane 10, Singapore 539222",
        "contact_person": "Mr. Steven Tan",
        "phone": "+65 6285 3456",
        "email": "steven@victorycm.com.sg",
        "registration": "202178901Z"
    },
    {
        "name": "Crown Engineering & Fabrication Pte Ltd",
        "address": "81 Ayer Rajah Crescent, Singapore 139967",
        "contact_person": "Ms. Jasmine Koh",
        "phone": "+65 6775 6543",
        "email": "jasmine@crowneng.sg",
        "registration": "202289012AA"
    },
    {
        "name": "Oceanic Trading & Supplies Pte Ltd",
        "address": "19 Gambas Crescent, Singapore 757087",
        "contact_person": "Mr. Vincent Lee",
        "phone": "+65 6752 7890",
        "email": "vincent.lee@oceanicts.com.sg",
        "registration": "201390123BB"
    },
    {
        "name": "Fortress Industrial Equipment Pte Ltd",
        "address": "46 Kallang Pudding Road, Singapore 349318",
        "contact_person": "Ms. Jennifer Ng",
        "phone": "+65 6294 4567",
        "email": "jennifer@fortressind.sg",
        "registration": "201401234CC"
    },
    {
        "name": "Paramount Building Technologies Pte Ltd",
        "address": "33 Ubi Avenue 3, Singapore 408868",
        "contact_person": "Mr. Daniel Chua",
        "phone": "+65 6846 7891",
        "email": "daniel@paramountbt.com.sg",
        "registration": "201512345DD"
    }
]

ALL_COMPANIES = COMPANIES + ADDITIONAL_COMPANIES


def generate_rfq_number():
    """Generate realistic RFQ number"""
    prefix = random.choice(["RFQ", "PO", "REQ", "QT"])
    year = random.choice([2024, 2025])
    number = random.randint(1000, 9999)
    return f"{prefix}-{year}-{number}"


def generate_date():
    """Generate realistic date within last 3 months"""
    days_ago = random.randint(0, 90)
    return datetime.now() - timedelta(days=days_ago)


def select_random_products(min_items=10, max_items=100):
    """Select random products for RFQ"""
    num_items = random.randint(min_items, max_items)
    selected = random.sample(HORME_PRODUCTS, min(num_items, len(HORME_PRODUCTS)))

    # Create line items with quantities
    line_items = []
    for product in selected:
        quantity = random.choice([
            random.randint(10, 50),      # Small quantities
            random.randint(50, 200),     # Medium quantities
            random.randint(200, 500),    # Large quantities
            random.randint(500, 1000),   # Bulk quantities
        ])

        line_items.append({
            "product": product["name"],
            "category": product["category"],
            "quantity": quantity,
            "unit": product["unit"]
        })

    return line_items


def create_excel_rfq(filename, company, rfq_number, rfq_date, line_items):
    """Create RFQ in Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Request for Quotation"

    # Styling
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    title_font = Font(bold=True, size=14)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Title
    ws.merge_cells('A1:F1')
    ws['A1'] = "REQUEST FOR QUOTATION"
    ws['A1'].font = title_font
    ws['A1'].alignment = Alignment(horizontal='center')

    # Company details
    ws['A3'] = "From:"
    ws['A3'].font = Font(bold=True)
    ws['A4'] = company["name"]
    ws['A5'] = company["address"]
    ws['A6'] = f"Contact: {company['contact_person']}"
    ws['A7'] = f"Phone: {company['phone']}"
    ws['A8'] = f"Email: {company['email']}"
    ws['A9'] = f"Company Reg: {company['registration']}"

    # RFQ details
    ws['E3'] = "RFQ Number:"
    ws['F3'] = rfq_number
    ws['E4'] = "Date:"
    ws['F4'] = rfq_date.strftime("%d %b %Y")
    ws['E5'] = "Valid Until:"
    ws['F5'] = (rfq_date + timedelta(days=30)).strftime("%d %b %Y")

    ws['E3'].font = Font(bold=True)
    ws['E4'].font = Font(bold=True)
    ws['E5'].font = Font(bold=True)

    # To (Horme)
    ws['A11'] = "To:"
    ws['A11'].font = Font(bold=True)
    ws['A12'] = "Horme Hardware Pte Ltd"
    ws['A13'] = "21 Senoko Loop, Singapore 758187"
    ws['A14'] = "Phone: +65 6257 0888"
    ws['A15'] = "Email: sales@horme.com.sg"

    # Items table header
    start_row = 17
    headers = ["No.", "Item Description", "Category", "Quantity", "Unit", "Unit Price (SGD)", "Total (SGD)"]
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=start_row, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = border
        cell.alignment = Alignment(horizontal='center')

    # Items
    for idx, item in enumerate(line_items, start=1):
        row = start_row + idx
        ws.cell(row=row, column=1, value=idx).border = border
        ws.cell(row=row, column=2, value=item["product"]).border = border
        ws.cell(row=row, column=3, value=item["category"]).border = border
        ws.cell(row=row, column=4, value=item["quantity"]).border = border
        ws.cell(row=row, column=4).alignment = Alignment(horizontal='right')
        ws.cell(row=row, column=5, value=item["unit"]).border = border
        ws.cell(row=row, column=6, value="").border = border  # For supplier to fill
        ws.cell(row=row, column=7, value="").border = border  # For supplier to fill

    # Footer
    footer_row = start_row + len(line_items) + 2
    ws.merge_cells(f'A{footer_row}:G{footer_row}')
    ws[f'A{footer_row}'] = "Please provide your best quotation including delivery terms and payment terms."
    ws[f'A{footer_row}'].font = Font(italic=True)

    # Notes
    notes_row = footer_row + 2
    ws[f'A{notes_row}'] = "Notes:"
    ws[f'A{notes_row}'].font = Font(bold=True)
    ws[f'A{notes_row+1}'] = "1. Prices should be in Singapore Dollars (SGD)"
    ws[f'A{notes_row+2}'] = "2. Delivery location: Project site in Singapore"
    ws[f'A{notes_row+3}'] = "3. Payment terms: 30 days from delivery"
    ws[f'A{notes_row+4}'] = "4. Quotation validity: 30 days from RFQ date"

    # Column widths
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 45
    ws.column_dimensions['C'].width = 30
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 18
    ws.column_dimensions['G'].width = 18

    wb.save(filename)


def create_word_rfq(filename, company, rfq_number, rfq_date, line_items):
    """Create RFQ in Word format"""
    doc = Document()

    # Title
    title = doc.add_heading('REQUEST FOR QUOTATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Blank line
    doc.add_paragraph()

    # Create header table
    header_table = doc.add_table(rows=2, cols=2)
    header_table.style = 'Light Grid Accent 1'

    # From section
    from_cell = header_table.cell(0, 0)
    from_para = from_cell.paragraphs[0]
    from_para.add_run('From:\n').bold = True
    from_para.add_run(f"{company['name']}\n")
    from_para.add_run(f"{company['address']}\n")
    from_para.add_run(f"Contact: {company['contact_person']}\n")
    from_para.add_run(f"Phone: {company['phone']}\n")
    from_para.add_run(f"Email: {company['email']}\n")
    from_para.add_run(f"Company Reg: {company['registration']}")

    # RFQ details section
    rfq_cell = header_table.cell(0, 1)
    rfq_para = rfq_cell.paragraphs[0]
    rfq_para.add_run(f"RFQ Number: {rfq_number}\n").bold = True
    rfq_para.add_run(f"Date: {rfq_date.strftime('%d %b %Y')}\n")
    rfq_para.add_run(f"Valid Until: {(rfq_date + timedelta(days=30)).strftime('%d %b %Y')}\n")

    # To section
    to_cell = header_table.cell(1, 0)
    to_para = to_cell.paragraphs[0]
    to_para.add_run('To:\n').bold = True
    to_para.add_run("Horme Hardware Pte Ltd\n")
    to_para.add_run("21 Senoko Loop, Singapore 758187\n")
    to_para.add_run("Phone: +65 6257 0888\n")
    to_para.add_run("Email: sales@horme.com.sg")

    # Merge the second row second cell with first row second cell for better layout
    doc.add_paragraph()

    # Items table
    items_table = doc.add_table(rows=len(line_items) + 1, cols=7)
    items_table.style = 'Light Grid Accent 1'

    # Headers
    headers = ["No.", "Item Description", "Category", "Quantity", "Unit", "Unit Price (SGD)", "Total (SGD)"]
    header_cells = items_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Items
    for idx, item in enumerate(line_items, start=1):
        row_cells = items_table.rows[idx].cells
        row_cells[0].text = str(idx)
        row_cells[1].text = item["product"]
        row_cells[2].text = item["category"]
        row_cells[3].text = str(item["quantity"])
        row_cells[4].text = item["unit"]
        row_cells[5].text = ""  # For supplier to fill
        row_cells[6].text = ""  # For supplier to fill

    # Footer
    doc.add_paragraph()
    doc.add_paragraph("Please provide your best quotation including delivery terms and payment terms.").italic = True

    # Notes
    doc.add_paragraph()
    notes_heading = doc.add_paragraph()
    notes_heading.add_run('Notes:').bold = True
    doc.add_paragraph('1. Prices should be in Singapore Dollars (SGD)', style='List Bullet')
    doc.add_paragraph('2. Delivery location: Project site in Singapore', style='List Bullet')
    doc.add_paragraph('3. Payment terms: 30 days from delivery', style='List Bullet')
    doc.add_paragraph('4. Quotation validity: 30 days from RFQ date', style='List Bullet')

    doc.save(filename)


def main():
    """Generate 30 RFQ documents"""
    output_dir = r"C:\Users\fujif\OneDrive\Desktop\Business Development\Horme\Sample Data"
    os.makedirs(output_dir, exist_ok=True)

    # Shuffle companies for randomness
    companies_pool = random.sample(ALL_COMPANIES, len(ALL_COMPANIES))

    print("Generating RFQ documents...")

    # Generate 15 Excel files
    for i in range(15):
        company = companies_pool[i]
        rfq_number = generate_rfq_number()
        rfq_date = generate_date()
        line_items = select_random_products(min_items=10, max_items=100)

        filename = os.path.join(output_dir, f"RFQ_{rfq_number}_{company['name'][:30]}.xlsx")
        create_excel_rfq(filename, company, rfq_number, rfq_date, line_items)
        print(f"[OK] Created Excel: {os.path.basename(filename)} ({len(line_items)} items)")

    # Generate 15 Word files
    for i in range(15, 30):
        company = companies_pool[i] if i < len(companies_pool) else random.choice(ALL_COMPANIES)
        rfq_number = generate_rfq_number()
        rfq_date = generate_date()
        line_items = select_random_products(min_items=10, max_items=100)

        filename = os.path.join(output_dir, f"RFQ_{rfq_number}_{company['name'][:30]}.docx")
        create_word_rfq(filename, company, rfq_number, rfq_date, line_items)
        print(f"[OK] Created Word: {os.path.basename(filename)} ({len(line_items)} items)")

    print(f"\n[SUCCESS] Generated 30 RFQ documents in: {output_dir}")
    print(f"   - 15 Excel files (.xlsx)")
    print(f"   - 15 Word files (.docx)")


if __name__ == "__main__":
    main()
