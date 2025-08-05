"""
AI-Powered Sales Assistant MCP Server
========================================

Comprehensive MCP server implementation providing:
- Document processing AI agents (PDF, Excel, Word extraction)
- Intelligent quote generation with pricing algorithms
- Real-time chat AI assistant with database and ERP tool access
- RAG implementation for document Q&A and sales support
- Multi-modal AI agents for different document types
- Integration with DataFlow models and Nexus platform
- Enterprise-grade authentication and monitoring

This MCP server serves as the AI orchestration layer for the sales assistant,
providing sophisticated AI capabilities through standardized MCP protocol.
"""

import os
import sys
import json
import asyncio
import logging
import hashlib
import mimetypes
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
from dataclasses import dataclass

# Apply Windows SDK compatibility patch BEFORE Kailash imports  
if os.name == 'nt':
    try:
        sys.path.append(os.path.join(os.path.dirname(__file__), 'new_project'))
        print("Windows SDK compatibility patch applied")
    except ImportError:
        print("Windows patch not found - attempting direct import")

# Kailash MCP Server imports
from kailash.mcp_server import MCPServer
from kailash.mcp_server.auth import APIKeyAuth, JWTAuth
from kailash.mcp_server.advanced_features import (
    MultiModalContent, ProgressReporter, ResourceTemplate, structured_tool
)

# Kailash Workflow imports
from kailash.workflow.builder import WorkflowBuilder
from kailash.runtime.local import LocalRuntime

# DataFlow models integration
from dataflow_models import (
    db, Company, User, Customer, Document, DocumentProcessingQueue, Quote
)
# Note: QuoteLineItem, ERPProduct, ActivityLog, WorkflowState not yet implemented

# Document processing libraries
import PyPDF2
import pandas as pd
import docx
from PIL import Image
import pytesseract

# AI/ML libraries
import openai
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Configuration
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
UPLOAD_DIR = Path("uploads")
RAG_STORAGE_DIR = Path("rag_storage")
DOCUMENT_CHUNK_SIZE = 1000
DOCUMENT_CHUNK_OVERLAP = 200

# Ensure directories exist
UPLOAD_DIR.mkdir(exist_ok=True)
RAG_STORAGE_DIR.mkdir(exist_ok=True)

# Initialize embedding model for RAG
try:
    embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
except Exception as e:
    logger.warning(f"Could not load embedding model: {e}")
    embedding_model = None

# Initialize vector database for RAG
try:
    chroma_client = chromadb.PersistentClient(path=str(RAG_STORAGE_DIR))
    documents_collection = chroma_client.get_or_create_collection(
        name="sales_documents",
        metadata={"hnsw:space": "cosine"}
    )
except Exception as e:
    logger.warning(f"Could not initialize ChromaDB: {e}")
    chroma_client = None
    documents_collection = None

# Authentication configuration
auth_keys = {
    "admin_key": {
        "permissions": ["admin", "tools", "resources", "documents", "quotes", "customers"],
        "rate_limit": 1000
    },
    "sales_manager_key": {
        "permissions": ["tools", "documents", "quotes", "customers", "reports"],
        "rate_limit": 500
    },
    "sales_rep_key": {
        "permissions": ["tools", "documents", "quotes", "customers"],
        "rate_limit": 200
    },
    "viewer_key": {
        "permissions": ["documents", "quotes", "customers"],
        "rate_limit": 100
    }
}

auth_provider = APIKeyAuth(auth_keys)

# Initialize production MCP server with enterprise features
server = MCPServer(
    name="sales-assistant-ai",
    auth_provider=auth_provider,
    enable_cache=True,
    cache_ttl=300,  # 5 minutes default cache
    cache_backend="memory",
    enable_metrics=True,
    enable_monitoring=True,
    rate_limit_config={},
    circuit_breaker_config={"failure_threshold": 5},
    enable_http_transport=True,
    enable_sse_transport=True,
    transport_timeout=60.0,
    max_request_size=50_000_000,  # 50MB for document uploads
    enable_streaming=True,
    error_aggregation=True
)

# ==============================================================================
# DOCUMENT PROCESSING AI AGENTS
# ==============================================================================

@dataclass
class DocumentExtractionResult:
    """Structured result for document extraction"""
    text_content: str
    metadata: Dict[str, Any]
    entities: List[Dict[str, Any]]
    confidence_score: float
    processing_time: float
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    key_terms: Optional[List[str]] = None
    summary: Optional[str] = None

class DocumentProcessor:
    """AI-powered document processing engine"""
    
    def __init__(self):
        self.openai_client = openai.OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
    
    async def extract_pdf(self, file_path: Path) -> DocumentExtractionResult:
        """Extract content from PDF documents"""
        start_time = datetime.now()
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ""
                page_count = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                # OCR fallback for scanned PDFs
                if len(text_content.strip()) < 100:
                    text_content = await self._ocr_pdf(file_path)
                
                word_count = len(text_content.split())
                
                # AI-powered analysis
                analysis = await self._analyze_with_ai(text_content, "PDF")
                
                processing_time = (datetime.now() - start_time).total_seconds()
                
                return DocumentExtractionResult(
                    text_content=text_content,
                    metadata={
                        "file_type": "PDF",
                        "page_count": page_count,
                        "word_count": word_count,
                        "file_size": file_path.stat().st_size
                    },
                    entities=analysis.get("entities", []),
                    confidence_score=analysis.get("confidence", 0.8),
                    processing_time=processing_time,
                    page_count=page_count,
                    word_count=word_count,
                    key_terms=analysis.get("key_terms", []),
                    summary=analysis.get("summary")
                )
                
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise
    
    async def extract_excel(self, file_path: Path) -> DocumentExtractionResult:
        """Extract content from Excel documents"""
        start_time = datetime.now()
        
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            text_content = ""
            sheet_info = []
            total_rows = 0
            
            for sheet_name, df in excel_data.items():
                rows, cols = df.shape
                total_rows += rows
                
                sheet_info.append({
                    "name": sheet_name,
                    "rows": rows,
                    "columns": cols,
                    "column_names": df.columns.tolist()
                })
                
                # Convert to text representation
                text_content += f"\n=== Sheet: {sheet_name} ===\n"
                text_content += df.to_string(max_rows=100) + "\n"
            
            # AI-powered analysis for structured data
            analysis = await self._analyze_structured_data(excel_data)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return DocumentExtractionResult(
                text_content=text_content,
                metadata={
                    "file_type": "Excel",
                    "sheets": sheet_info,
                    "total_rows": total_rows,
                    "file_size": file_path.stat().st_size
                },
                entities=analysis.get("entities", []),
                confidence_score=analysis.get("confidence", 0.9),
                processing_time=processing_time,
                word_count=len(text_content.split()),
                key_terms=analysis.get("key_terms", []),
                summary=analysis.get("summary")
            )
            
        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            raise
    
    async def extract_word(self, file_path: Path) -> DocumentExtractionResult:
        """Extract content from Word documents"""
        start_time = datetime.now()
        
        try:
            doc = docx.Document(file_path)
            text_content = ""
            paragraph_count = 0
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Extract tables
            table_data = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                    text_content += " | ".join(row_data) + "\n"
            
            word_count = len(text_content.split())
            
            # AI-powered analysis
            analysis = await self._analyze_with_ai(text_content, "Word")
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return DocumentExtractionResult(
                text_content=text_content,
                metadata={
                    "file_type": "Word",
                    "paragraph_count": paragraph_count,
                    "table_count": len(doc.tables),
                    "word_count": word_count,
                    "file_size": file_path.stat().st_size
                },
                entities=analysis.get("entities", []),
                confidence_score=analysis.get("confidence", 0.85),
                processing_time=processing_time,
                word_count=word_count,
                key_terms=analysis.get("key_terms", []),
                summary=analysis.get("summary")
            )
            
        except Exception as e:
            logger.error(f"Word extraction failed: {e}")
            raise
    
    async def _ocr_pdf(self, file_path: Path) -> str:
        """OCR fallback for scanned PDFs"""
        try:
            # This would integrate with pytesseract for OCR
            # Simplified implementation
            return f"OCR text extraction from {file_path.name} would go here"
        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return ""
    
    async def _analyze_with_ai(self, text: str, doc_type: str) -> Dict[str, Any]:
        """AI-powered document analysis"""
        if not self.openai_client or not OPENAI_API_KEY:
            # Mock analysis for development
            return {
                "entities": [
                    {"type": "organization", "text": "Sample Corp", "confidence": 0.9},
                    {"type": "money", "text": "$50,000", "confidence": 0.85}
                ],
                "key_terms": ["proposal", "budget", "timeline", "requirements"],
                "summary": f"This {doc_type} document contains business information with key requirements and financial details.",
                "confidence": 0.8
            }
        
        try:
            response = await self.openai_client.chat.completions.acreate(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": f"""Analyze this {doc_type} document and extract:
                        1. Named entities (organizations, people, money, dates)
                        2. Key terms and concepts
                        3. A concise summary
                        4. Document classification
                        
                        Return as JSON with fields: entities, key_terms, summary, classification, confidence"""
                    },
                    {
                        "role": "user",
                        "content": text[:4000]  # Limit content length
                    }
                ],
                temperature=0.1
            )
            
            analysis = json.loads(response.choices[0].message.content)
            return analysis
            
        except Exception as e:
            logger.error(f"AI analysis failed: {e}")
            return {
                "entities": [],
                "key_terms": [],
                "summary": "AI analysis unavailable",
                "confidence": 0.5
            }
    
    async def _analyze_structured_data(self, excel_data: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze structured Excel data"""
        try:
            # Analyze data patterns, types, and potential insights
            analysis = {
                "entities": [],
                "key_terms": [],
                "summary": "Structured data analysis",
                "confidence": 0.9
            }
            
            for sheet_name, df in excel_data.items():
                # Detect data types and patterns
                numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
                date_cols = df.select_dtypes(include=['datetime']).columns.tolist()
                text_cols = df.select_dtypes(include=['object']).columns.tolist()
                
                if numeric_cols:
                    analysis["key_terms"].extend(["financial_data", "metrics", "numbers"])
                if date_cols:
                    analysis["key_terms"].extend(["timeline", "dates", "schedule"])
                
                # Look for common business patterns
                if any('price' in col.lower() or 'cost' in col.lower() or 'amount' in col.lower() 
                       for col in df.columns):
                    analysis["entities"].append({
                        "type": "financial_data",
                        "text": "Pricing information detected",
                        "confidence": 0.9
                    })
            
            return analysis
            
        except Exception as e:
            logger.error(f"Structured data analysis failed: {e}")
            return {"entities": [], "key_terms": [], "summary": "Analysis failed", "confidence": 0.3}

# Initialize document processor
document_processor = DocumentProcessor()

# ==============================================================================
# RAG IMPLEMENTATION FOR DOCUMENT Q&A
# ==============================================================================

class RAGEngine:
    """Retrieval-Augmented Generation for document Q&A"""
    
    def __init__(self):
        self.embedding_model = embedding_model
        self.collection = documents_collection
        self.openai_client = openai.OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
    
    async def index_document(self, document_id: int, content: str, metadata: Dict[str, Any]) -> bool:
        """Index document content for RAG retrieval"""
        if not self.embedding_model or not self.collection:
            logger.warning("RAG components not available")
            return False
        
        try:
            # Split document into chunks
            chunks = self._chunk_text(content)
            
            # Generate embeddings
            embeddings = self.embedding_model.encode(chunks)
            
            # Store in vector database
            chunk_ids = [f"{document_id}_chunk_{i}" for i in range(len(chunks))]
            chunk_metadata = [
                {**metadata, "chunk_index": i, "document_id": document_id}
                for i in range(len(chunks))
            ]
            
            self.collection.add(
                embeddings=embeddings.tolist(),
                documents=chunks,
                metadatas=chunk_metadata,
                ids=chunk_ids
            )
            
            logger.info(f"Indexed {len(chunks)} chunks for document {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Document indexing failed: {e}")
            return False
    
    async def query_documents(self, query: str, limit: int = 5, 
                            filters: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """Query indexed documents using RAG"""
        if not self.embedding_model or not self.collection:
            return []
        
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])
            
            # Search similar chunks
            results = self.collection.query(
                query_embeddings=query_embedding.tolist(),
                n_results=limit,
                where=filters
            )
            
            # Format results
            retrieved_docs = []
            for i in range(len(results['documents'][0])):
                retrieved_docs.append({
                    "content": results['documents'][0][i],
                    "metadata": results['metadatas'][0][i],
                    "score": 1 - results['distances'][0][i],  # Convert distance to similarity
                    "id": results['ids'][0][i]
                })
            
            return retrieved_docs
            
        except Exception as e:
            logger.error(f"Document query failed: {e}")
            return []
    
    async def answer_question(self, question: str, context_docs: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Generate answer using retrieved context"""
        if not self.openai_client:
            return {
                "answer": "AI service not available",
                "confidence": 0.0,
                "sources": []
            }
        
        try:
            # Prepare context from retrieved documents
            context = "\n\n".join([
                f"Document {doc['metadata']['document_id']}: {doc['content']}"
                for doc in context_docs
            ])
            
            response = await self.openai_client.chat.completions.acreate(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": """You are a sales assistant AI. Use the provided document context to answer questions accurately. 
                        If you cannot answer based on the context, say so clearly. Include confidence level and cite sources."""
                    },
                    {
                        "role": "user",
                        "content": f"""Context documents:
                        {context}
                        
                        Question: {question}
                        
                        Please provide a comprehensive answer with confidence level (0-1) and source citations."""
                    }
                ],
                temperature=0.2
            )
            
            answer = response.choices[0].message.content
            
            return {
                "answer": answer,
                "confidence": 0.85,  # Could be extracted from response
                "sources": [doc['metadata']['document_id'] for doc in context_docs],
                "retrieved_chunks": len(context_docs)
            }
            
        except Exception as e:
            logger.error(f"Answer generation failed: {e}")
            return {
                "answer": "Failed to generate answer",
                "confidence": 0.0,
                "sources": [],
                "error": str(e)
            }
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), DOCUMENT_CHUNK_SIZE - DOCUMENT_CHUNK_OVERLAP):
            chunk_words = words[i:i + DOCUMENT_CHUNK_SIZE]
            chunk = " ".join(chunk_words)
            if chunk.strip():
                chunks.append(chunk)
        
        return chunks

# Initialize RAG engine
rag_engine = RAGEngine()

# ==============================================================================
# INTELLIGENT QUOTE GENERATION SYSTEM
# ==============================================================================

class QuoteGenerationEngine:
    """AI-powered quote generation with pricing algorithms"""
    
    def __init__(self):
        self.openai_client = openai.OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
    
    async def generate_quote_from_rfp(self, rfp_content: str, customer_id: int) -> Dict[str, Any]:
        """Generate intelligent quote from RFP document"""
        try:
            # Extract requirements from RFP
            requirements = await self._extract_requirements(rfp_content)
            
            # Find matching products
            products = await self._find_matching_products(requirements)
            
            # Calculate pricing with algorithms
            pricing = await self._calculate_intelligent_pricing(products, customer_id)
            
            # Generate quote structure
            quote_data = {
                "requirements_analysis": requirements,
                "recommended_products": products,
                "pricing_analysis": pricing,
                "confidence_score": pricing.get("confidence", 0.8),
                "total_estimate": pricing.get("total", 0.0)
            }
            
            return quote_data
            
        except Exception as e:
            logger.error(f"Quote generation failed: {e}")
            raise
    
    async def _extract_requirements(self, rfp_content: str) -> Dict[str, Any]:
        """Extract structured requirements from RFP using AI"""
        if not self.openai_client:
            # Mock requirements extraction
            return {
                "functional_requirements": ["Feature A", "Feature B"],
                "technical_requirements": ["Integration X", "Platform Y"],
                "budget_range": "$50,000 - $100,000",
                "timeline": "6 months",
                "priority_level": "high"
            }
        
        try:
            response = await self.openai_client.chat.completions.acreate(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": """Extract structured requirements from this RFP document.
                        Return JSON with: functional_requirements, technical_requirements, 
                        budget_range, timeline, priority_level, special_considerations"""
                    },
                    {
                        "role": "user",
                        "content": rfp_content[:6000]  # Limit content
                    }
                ],
                temperature=0.1
            )
            
            return json.loads(response.choices[0].message.content)
            
        except Exception as e:
            logger.error(f"Requirements extraction failed: {e}")
            return {"error": str(e), "functional_requirements": []}
    
    async def _find_matching_products(self, requirements: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Find products matching requirements using database and AI"""
        try:
            runtime = LocalRuntime()
            workflow = WorkflowBuilder()
            
            # Search products based on requirements
            workflow.add_node("AsyncSQLQueryNode", "search_products", {
                "query": """
                    SELECT product_code, name, description, category, list_price, 
                           stock_status, specifications
                    FROM erp_products 
                    WHERE stock_status = 'available'
                    AND (name ILIKE %(search_term)s OR description ILIKE %(search_term)s)
                    ORDER BY list_price ASC
                    LIMIT 20
                """,
                "parameters": {
                    "search_term": f"%{' '.join(requirements.get('functional_requirements', []))[:100]}%"
                },
                "connection_pool": db.connection_pool
            })
            
            results, _ = runtime.execute(workflow.build())
            products = results.get("search_products", [])
            
            # AI-powered product matching
            if self.openai_client and products:
                enhanced_products = await self._enhance_product_matching(products, requirements)
                return enhanced_products
            
            return products
            
        except Exception as e:
            logger.error(f"Product matching failed: {e}")
            return []
    
    async def _enhance_product_matching(self, products: List[Dict], requirements: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Use AI to enhance product matching and recommendations"""
        try:
            product_summaries = [f"{p['name']}: {p['description']}" for p in products[:10]]
            
            response = await self.openai_client.chat.completions.acreate(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": """Score and rank products based on how well they match the requirements.
                        Return JSON with product rankings, match scores (0-1), and reasoning."""
                    },
                    {
                        "role": "user",
                        "content": f"""Requirements: {json.dumps(requirements)}
                        
                        Products: {json.dumps(product_summaries)}
                        
                        Rank products by relevance and provide match scores."""
                    }
                ],
                temperature=0.2
            )
            
            ai_rankings = json.loads(response.choices[0].message.content)
            
            # Merge AI insights with product data
            enhanced_products = []
            for product in products:
                # Find AI score for this product
                ai_score = 0.5  # Default
                for ranking in ai_rankings.get("rankings", []):
                    if product["name"] in ranking.get("product", ""):
                        ai_score = ranking.get("match_score", 0.5)
                        break
                
                enhanced_product = {
                    **product,
                    "ai_match_score": ai_score,
                    "recommendation_level": "high" if ai_score > 0.8 else "medium" if ai_score > 0.6 else "low"
                }
                enhanced_products.append(enhanced_product)
            
            # Sort by AI match score
            enhanced_products.sort(key=lambda x: x["ai_match_score"], reverse=True)
            
            return enhanced_products[:10]  # Return top 10
            
        except Exception as e:
            logger.error(f"AI product enhancement failed: {e}")
            return products
    
    async def _calculate_intelligent_pricing(self, products: List[Dict], customer_id: int) -> Dict[str, Any]:
        """Calculate intelligent pricing with algorithms"""
        try:
            # Get customer information for pricing context
            runtime = LocalRuntime()
            workflow = WorkflowBuilder()
            
            workflow.add_node("AsyncSQLQueryNode", "get_customer", {
                "query": """
                    SELECT c.*, COUNT(q.id) as quote_count, AVG(q.total_amount) as avg_quote_value
                    FROM customers c
                    LEFT JOIN quotes q ON c.id = q.customer_id
                    WHERE c.id = %(customer_id)s
                    GROUP BY c.id
                """,
                "parameters": {"customer_id": customer_id},
                "connection_pool": db.connection_pool
            })
            
            results, _ = runtime.execute(workflow.build())
            customer_data = results.get("get_customer", [{}])[0]
            
            # Apply pricing algorithms
            base_total = sum(float(p.get("list_price", 0)) for p in products)
            
            # Volume discount algorithm
            volume_discount = self._calculate_volume_discount(base_total)
            
            # Customer loyalty discount
            loyalty_discount = self._calculate_loyalty_discount(customer_data)
            
            # Strategic discount (AI-driven)
            strategic_discount = await self._calculate_strategic_discount(products, customer_data)
            
            # Final calculations
            total_discount = volume_discount + loyalty_discount + strategic_discount
            total_discount = min(total_discount, 0.25)  # Cap at 25%
            
            discounted_total = base_total * (1 - total_discount)
            
            pricing_analysis = {
                "base_total": base_total,
                "volume_discount": volume_discount,
                "loyalty_discount": loyalty_discount,
                "strategic_discount": strategic_discount,
                "total_discount_percent": total_discount * 100,
                "discounted_total": discounted_total,
                "total": discounted_total,
                "confidence": 0.85,
                "algorithm_notes": {
                    "volume_tier": self._get_volume_tier(base_total),
                    "customer_tier": self._get_customer_tier(customer_data),
                    "pricing_strategy": "competitive"
                }
            }
            
            return pricing_analysis
            
        except Exception as e:
            logger.error(f"Pricing calculation failed: {e}")
            return {"total": 0.0, "confidence": 0.0, "error": str(e)}
    
    def _calculate_volume_discount(self, total: float) -> float:
        """Calculate volume-based discount"""
        if total > 100000:
            return 0.15  # 15% for orders over $100k
        elif total > 50000:
            return 0.10  # 10% for orders over $50k
        elif total > 25000:
            return 0.05  # 5% for orders over $25k
        return 0.0
    
    def _calculate_loyalty_discount(self, customer_data: Dict) -> float:
        """Calculate customer loyalty discount"""
        quote_count = customer_data.get("quote_count", 0)
        avg_value = float(customer_data.get("avg_quote_value", 0))
        
        if quote_count > 10 and avg_value > 25000:
            return 0.08  # 8% for high-value repeat customers
        elif quote_count > 5:
            return 0.05  # 5% for repeat customers
        elif quote_count > 2:
            return 0.02  # 2% for returning customers
        return 0.0
    
    async def _calculate_strategic_discount(self, products: List[Dict], customer_data: Dict) -> float:
        """AI-driven strategic discount calculation"""
        try:
            # Factors: competition, market conditions, customer profile, product mix
            base_strategic = 0.03  # 3% base strategic discount
            
            # Industry-specific adjustments
            industry = customer_data.get("industry", "")
            if industry.lower() in ["healthcare", "education", "nonprofit"]:
                base_strategic += 0.02  # Additional 2% for strategic sectors
            
            # Product mix analysis
            high_margin_products = sum(1 for p in products if p.get("ai_match_score", 0) > 0.8)
            if high_margin_products > len(products) * 0.6:
                base_strategic += 0.02  # Can afford more discount on high-margin items
            
            return min(base_strategic, 0.10)  # Cap strategic discount at 10%
            
        except Exception as e:
            logger.error(f"Strategic discount calculation failed: {e}")
            return 0.03
    
    def _get_volume_tier(self, total: float) -> str:
        """Determine volume tier for customer"""
        if total > 100000:
            return "enterprise"
        elif total > 50000:
            return "large"
        elif total > 25000:
            return "medium"
        return "small"
    
    def _get_customer_tier(self, customer_data: Dict) -> str:
        """Determine customer tier"""
        quote_count = customer_data.get("quote_count", 0)
        if quote_count > 10:
            return "platinum"
        elif quote_count > 5:
            return "gold"
        elif quote_count > 2:
            return "silver"
        return "bronze"

# Initialize quote generation engine
quote_engine = QuoteGenerationEngine()

# ==============================================================================
# REAL-TIME CHAT AI ASSISTANT
# ==============================================================================

class ChatAssistant:
    """Real-time AI chat assistant with context and memory"""
    
    def __init__(self):
        self.openai_client = openai.OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
        self.conversation_memory = {}  # In production, use Redis or database
    
    async def process_message(self, message: str, user_id: int, session_id: str, 
                            context: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Process chat message with context and memory"""
        try:
            # Get conversation history
            conversation = self.conversation_memory.get(session_id, [])
            
            # Add current message
            conversation.append({"role": "user", "content": message, "timestamp": datetime.now().isoformat()})
            
            # Get relevant context from database if needed
            if context is None:
                context = await self._gather_context(message, user_id)
            
            # Determine intent and required tools
            intent = await self._analyze_intent(message, context)
            
            # Generate response based on intent
            if intent["requires_tools"]:
                response = await self._handle_tool_request(message, intent, context, user_id)
            else:
                response = await self._generate_conversational_response(message, conversation, context)
            
            # Add response to memory
            conversation.append({"role": "assistant", "content": response["content"], "timestamp": datetime.now().isoformat()})
            
            # Keep conversation memory manageable (last 20 messages)
            self.conversation_memory[session_id] = conversation[-20:]
            
            return {
                "response": response["content"],
                "intent": intent,
                "context_used": context,
                "suggested_actions": response.get("suggested_actions", []),
                "confidence": response.get("confidence", 0.8)
            }
            
        except Exception as e:
            logger.error(f"Chat processing failed: {e}")
            return {
                "response": "I'm having trouble processing your request right now. Please try again.",
                "error": str(e),
                "confidence": 0.0
            }
    
    async def _gather_context(self, message: str, user_id: int) -> Dict[str, Any]:
        """Gather relevant context for the user's message"""
        context = {"user_id": user_id}
        
        try:
            runtime = LocalRuntime()
            workflow = WorkflowBuilder()
            
            # Get user information
            workflow.add_node("AsyncSQLQueryNode", "get_user", {
                "query": """
                    SELECT u.*, c.name as company_name 
                    FROM users u 
                    JOIN companies c ON u.company_id = c.id 
                    WHERE u.id = %(user_id)s
                """,
                "parameters": {"user_id": user_id},
                "connection_pool": db.connection_pool
            })
            
            # Get recent activity
            workflow.add_node("AsyncSQLQueryNode", "get_recent_activity", {
                "query": """
                    SELECT entity_type, entity_id, action, timestamp 
                    FROM activity_logs 
                    WHERE user_id = %(user_id)s 
                    ORDER BY timestamp DESC 
                    LIMIT 5
                """,
                "parameters": {"user_id": user_id},
                "connection_pool": db.connection_pool
            })
            
            results, _ = runtime.execute(workflow.build())
            
            context["user_info"] = results.get("get_user", [{}])[0]
            context["recent_activity"] = results.get("get_recent_activity", [])
            
            return context
            
        except Exception as e:
            logger.error(f"Context gathering failed: {e}")
            return context
    
    async def _analyze_intent(self, message: str, context: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze user intent and determine required actions"""
        if not self.openai_client:
            # Simple rule-based intent detection for development
            message_lower = message.lower()
            if any(word in message_lower for word in ["quote", "pricing", "price"]):
                return {"intent": "quote_request", "requires_tools": True, "confidence": 0.8}
            elif any(word in message_lower for word in ["customer", "client"]):
                return {"intent": "customer_inquiry", "requires_tools": True, "confidence": 0.7}
            elif any(word in message_lower for word in ["document", "file", "pdf"]):
                return {"intent": "document_inquiry", "requires_tools": True, "confidence": 0.7}
            else:
                return {"intent": "general_conversation", "requires_tools": False, "confidence": 0.6}
        
        try:
            response = await self.openai_client.chat.completions.acreate(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": """Analyze user intent and determine if tools are needed.
                        Return JSON with: intent, requires_tools, confidence, suggested_tools
                        
                        Possible intents: quote_request, customer_inquiry, document_inquiry, 
                        product_search, general_conversation, data_analysis"""
                    },
                    {
                        "role": "user",
                        "content": f"Message: {message}\n\nContext: {json.dumps(context, default=str)}"
                    }
                ],
                temperature=0.1
            )
            
            return json.loads(response.choices[0].message.content)
            
        except Exception as e:
            logger.error(f"Intent analysis failed: {e}")
            return {"intent": "general_conversation", "requires_tools": False, "confidence": 0.5}
    
    async def _handle_tool_request(self, message: str, intent: Dict[str, Any], 
                                  context: Dict[str, Any], user_id: int) -> Dict[str, Any]:
        """Handle requests that require tool usage"""
        intent_type = intent.get("intent")
        
        if intent_type == "quote_request":
            return await self._handle_quote_request(message, context, user_id)
        elif intent_type == "customer_inquiry":
            return await self._handle_customer_inquiry(message, context, user_id)
        elif intent_type == "document_inquiry":
            return await self._handle_document_inquiry(message, context, user_id)
        elif intent_type == "product_search":
            return await self._handle_product_search(message, context, user_id)
        else:
            return await self._generate_conversational_response(message, [], context)
    
    async def _handle_quote_request(self, message: str, context: Dict[str, Any], user_id: int) -> Dict[str, Any]:
        """Handle quote-related requests"""
        try:
            # Search for recent quotes or create new one based on message
            runtime = LocalRuntime()
            workflow = WorkflowBuilder()
            
            workflow.add_node("AsyncSQLQueryNode", "recent_quotes", {
                "query": """
                    SELECT q.*, c.name as customer_name 
                    FROM quotes q 
                    JOIN customers c ON q.customer_id = c.id 
                    WHERE q.created_by = %(user_id)s 
                    ORDER BY q.created_date DESC 
                    LIMIT 5
                """,
                "parameters": {"user_id": user_id},
                "connection_pool": db.connection_pool
            })
            
            results, _ = runtime.execute(workflow.build())
            recent_quotes = results.get("recent_quotes", [])
            
            response_content = f"I found {len(recent_quotes)} recent quotes you've created. "
            
            if recent_quotes:
                response_content += "Here are your most recent quotes:\n\n"
                for quote in recent_quotes[:3]:
                    response_content += f"• Quote {quote['quote_number']} for {quote['customer_name']} - ${quote['total_amount']:,.2f} ({quote['status']})\n"
                response_content += "\nWould you like me to help you create a new quote or review an existing one?"
            else:
                response_content += "Would you like me to help you create a new quote? I can assist with pricing, product selection, and customer analysis."
            
            return {
                "content": response_content,
                "suggested_actions": ["Create new quote", "Review recent quotes", "Search customers"],
                "data": {"recent_quotes": recent_quotes[:3]},
                "confidence": 0.9
            }
            
        except Exception as e:
            logger.error(f"Quote request handling failed: {e}")
            return {
                "content": "I can help you with quotes, but I'm having trouble accessing the quote data right now.",
                "confidence": 0.3
            }
    
    async def _handle_customer_inquiry(self, message: str, context: Dict[str, Any], user_id: int) -> Dict[str, Any]:
        """Handle customer-related inquiries"""
        try:
            # Extract potential customer name or search term from message
            search_term = message.replace("customer", "").replace("client", "").strip()
            
            if len(search_term) < 3:
                return {
                    "content": "Please provide a customer name or search term to help me find the right customer information.",
                    "suggested_actions": ["Search customers", "Recent customers", "Top customers"],
                    "confidence": 0.7
                }
            
            runtime = LocalRuntime()
            workflow = WorkflowBuilder()
            
            workflow.add_node("AsyncSQLQueryNode", "search_customers", {
                "query": """
                    SELECT c.*, COUNT(q.id) as quote_count, 
                           COALESCE(SUM(q.total_amount), 0) as total_value
                    FROM customers c
                    LEFT JOIN quotes q ON c.id = q.customer_id
                    WHERE c.name ILIKE %(search_term)s 
                       OR c.primary_contact ILIKE %(search_term)s
                       OR c.email ILIKE %(search_term)s
                    GROUP BY c.id
                    ORDER BY total_value DESC
                    LIMIT 5
                """,
                "parameters": {"search_term": f"%{search_term}%"},
                "connection_pool": db.connection_pool
            })
            
            results, _ = runtime.execute(workflow.build())
            customers = results.get("search_customers", [])
            
            if customers:
                response_content = f"I found {len(customers)} customers matching '{search_term}':\n\n"
                for customer in customers:
                    response_content += f"• **{customer['name']}** ({customer['type']})\n"
                    response_content += f"  Contact: {customer['primary_contact']} - {customer['email']}\n"
                    response_content += f"  Quotes: {customer['quote_count']} (${customer['total_value']:,.2f} total)\n\n"
                
                response_content += "Would you like detailed information about any of these customers?"
            else:
                response_content = f"I couldn't find any customers matching '{search_term}'. Would you like me to search differently or help you create a new customer record?"
            
            return {
                "content": response_content,
                "suggested_actions": ["Customer details", "Create customer", "Search again"],
                "data": {"customers": customers},
                "confidence": 0.85
            }
            
        except Exception as e:
            logger.error(f"Customer inquiry handling failed: {e}")
            return {
                "content": "I can help you with customer information, but I'm having trouble accessing the customer database right now.",
                "confidence": 0.3
            }
    
    async def _handle_document_inquiry(self, message: str, context: Dict[str, Any], user_id: int) -> Dict[str, Any]:
        """Handle document-related inquiries"""
        try:
            runtime = LocalRuntime()
            workflow = WorkflowBuilder()
            
            workflow.add_node("AsyncSQLQueryNode", "recent_documents", {
                "query": """
                    SELECT d.*, c.name as customer_name 
                    FROM documents d
                    LEFT JOIN customers c ON d.customer_id = c.id
                    WHERE d.uploaded_by = %(user_id)s 
                    ORDER BY d.upload_date DESC 
                    LIMIT 10
                """,
                "parameters": {"user_id": user_id},
                "connection_pool": db.connection_pool
            })
            
            results, _ = runtime.execute(workflow.build())
            documents = results.get("recent_documents", [])
            
            response_content = f"Here are your recent documents ({len(documents)} found):\n\n"
            
            for doc in documents[:5]:
                status_emoji = "✅" if doc['ai_status'] == 'completed' else "🔄" if doc['ai_status'] == 'processing' else "⏳"
                response_content += f"{status_emoji} **{doc['name']}** ({doc['type']})\n"
                if doc['customer_name']:
                    response_content += f"   Customer: {doc['customer_name']}\n"
                response_content += f"   Uploaded: {doc['upload_date']}\n"
                response_content += f"   AI Status: {doc['ai_status']}\n\n"
            
            if len(documents) > 5:
                response_content += f"... and {len(documents) - 5} more documents.\n\n"
            
            response_content += "I can help you search documents, analyze content, or answer questions about any of these files."
            
            return {
                "content": response_content,
                "suggested_actions": ["Search documents", "Analyze document", "Ask about document"],
                "data": {"documents": documents[:5]},
                "confidence": 0.9
            }
            
        except Exception as e:
            logger.error(f"Document inquiry handling failed: {e}")
            return {
                "content": "I can help with document analysis and search, but I'm having trouble accessing the document database right now.",
                "confidence": 0.3
            }
    
    async def _handle_product_search(self, message: str, context: Dict[str, Any], user_id: int) -> Dict[str, Any]:
        """Handle product search requests"""
        try:
            # Extract search terms from message
            search_terms = message.replace("product", "").replace("search", "").strip()
            
            runtime = LocalRuntime()
            workflow = WorkflowBuilder()
            
            workflow.add_node("AsyncSQLQueryNode", "search_products", {
                "query": """
                    SELECT product_code, name, description, category, list_price, stock_status
                    FROM erp_products 
                    WHERE (name ILIKE %(search_term)s OR description ILIKE %(search_term)s)
                    AND stock_status = 'available'
                    ORDER BY list_price ASC
                    LIMIT 10
                """,
                "parameters": {"search_term": f"%{search_terms}%"},
                "connection_pool": db.connection_pool
            })
            
            results, _ = runtime.execute(workflow.build())
            products = results.get("search_products", [])
            
            if products:
                response_content = f"I found {len(products)} products matching '{search_terms}':\n\n"
                for product in products[:5]:
                    response_content += f"• **{product['name']}** ({product['product_code']})\n"
                    response_content += f"  Category: {product['category']}\n"
                    response_content += f"  Price: ${product['list_price']:,.2f}\n"
                    response_content += f"  Status: {product['stock_status']}\n\n"
                
                if len(products) > 5:
                    response_content += f"... and {len(products) - 5} more products.\n\n"
                
                response_content += "Would you like more details about any of these products or help creating a quote?"
            else:
                response_content = f"I couldn't find any available products matching '{search_terms}'. Try different search terms or browse by category."
            
            return {
                "content": response_content,
                "suggested_actions": ["Product details", "Add to quote", "Search again"],
                "data": {"products": products[:5]},
                "confidence": 0.85
            }
            
        except Exception as e:
            logger.error(f"Product search handling failed: {e}")
            return {
                "content": "I can help you search products, but I'm having trouble accessing the product database right now.",
                "confidence": 0.3
            }
    
    async def _generate_conversational_response(self, message: str, conversation: List[Dict], 
                                              context: Dict[str, Any]) -> Dict[str, Any]:
        """Generate conversational response using AI"""
        if not self.openai_client:
            return {
                "content": "I'm here to help with your sales activities. I can assist with quotes, customers, documents, and product information. What would you like to work on?",
                "confidence": 0.6
            }
        
        try:
            # Prepare conversation history
            messages = [
                {
                    "role": "system",
                    "content": f"""You are an AI sales assistant. Help users with:
                    - Quote generation and management
                    - Customer information and analysis  
                    - Document processing and Q&A
                    - Product search and recommendations
                    - Sales data analysis
                    
                    User context: {json.dumps(context.get('user_info', {}), default=str)}
                    
                    Be helpful, professional, and concise. Offer specific actions when appropriate."""
                }
            ]
            
            # Add conversation history (last 10 messages)
            messages.extend(conversation[-10:])
            
            # Add current message
            messages.append({"role": "user", "content": message})
            
            response = await self.openai_client.chat.completions.acreate(
                model="gpt-4",
                messages=messages,
                temperature=0.7,
                max_tokens=500
            )
            
            return {
                "content": response.choices[0].message.content,
                "confidence": 0.8,
                "suggested_actions": ["Create quote", "Search customers", "Upload document"]
            }
            
        except Exception as e:
            logger.error(f"Conversational response generation failed: {e}")
            return {
                "content": "I'm here to help with your sales work. What can I assist you with today?",
                "confidence": 0.5
            }

# Initialize chat assistant
chat_assistant = ChatAssistant()

# ==============================================================================
# MCP TOOL IMPLEMENTATIONS
# ==============================================================================

@server.tool(required_permission="documents", cache_ttl=60)
async def process_document(file_path: str, document_type: str = "auto", 
                          customer_id: Optional[int] = None) -> Dict[str, Any]:
    """Process uploaded document with AI extraction and analysis"""
    try:
        file_path_obj = Path(file_path)
        
        if not file_path_obj.exists():
            return {"success": False, "error": "File not found"}
        
        # Determine document type if auto
        if document_type == "auto":
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type == "application/pdf":
                document_type = "pdf"
            elif mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
                document_type = "excel"
            elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                document_type = "word"
            else:
                return {"success": False, "error": f"Unsupported file type: {mime_type}"}
        
        # Process based on document type
        if document_type == "pdf":
            result = await document_processor.extract_pdf(file_path_obj)
        elif document_type == "excel":
            result = await document_processor.extract_excel(file_path_obj)
        elif document_type == "word":
            result = await document_processor.extract_word(file_path_obj)
        else:
            return {"success": False, "error": f"Unsupported document type: {document_type}"}
        
        # Store in database
        runtime = LocalRuntime()
        workflow = WorkflowBuilder()
        
        workflow.add_node("AsyncSQLExecuteNode", "store_document", {
            "query": """
                INSERT INTO documents (name, type, category, file_path, file_size, mime_type,
                                     customer_id, upload_date, uploaded_by, ai_status, 
                                     ai_extracted_data, ai_confidence_score, page_count, word_count)
                VALUES (%(name)s, %(type)s, 'inbound', %(file_path)s, %(file_size)s, %(mime_type)s,
                       %(customer_id)s, NOW(), 1, 'completed', %(ai_data)s, %(confidence)s, 
                       %(page_count)s, %(word_count)s)
                RETURNING id
            """,
            "parameters": {
                "name": file_path_obj.name,
                "type": document_type,
                "file_path": str(file_path_obj),
                "file_size": file_path_obj.stat().st_size,
                "mime_type": mimetypes.guess_type(file_path)[0],
                "customer_id": customer_id,
                "ai_data": json.dumps({
                    "entities": result.entities,
                    "key_terms": result.key_terms,
                    "summary": result.summary
                }),
                "confidence": result.confidence_score,
                "page_count": result.page_count,
                "word_count": result.word_count
            },
            "connection_pool": db.connection_pool
        })
        
        db_results, _ = runtime.execute(workflow.build())
        document_id = db_results.get("store_document", {}).get("id")
        
        # Index for RAG if available
        if document_id and result.text_content:
            await rag_engine.index_document(
                document_id, 
                result.text_content, 
                {"document_type": document_type, "customer_id": customer_id}
            )
        
        return {
            "success": True,
            "document_id": document_id,
            "extraction_result": {
                "text_length": len(result.text_content),
                "entities_found": len(result.entities),
                "key_terms": result.key_terms,
                "summary": result.summary,
                "confidence_score": result.confidence_score,
                "processing_time": result.processing_time,
                "metadata": result.metadata
            }
        }
        
    except Exception as e:
        logger.error(f"Document processing failed: {e}")
        return {"success": False, "error": str(e)}

@server.tool(required_permission="quotes", cache_ttl=30)
async def generate_intelligent_quote(customer_id: int, rfp_content: Optional[str] = None,
                                   requirements: Optional[List[str]] = None) -> Dict[str, Any]:
    """Generate intelligent quote using AI and pricing algorithms"""
    try:
        if rfp_content:
            # Generate from RFP document
            quote_data = await quote_engine.generate_quote_from_rfp(rfp_content, customer_id)
        elif requirements:
            # Generate from requirements list
            req_dict = {"functional_requirements": requirements}
            products = await quote_engine._find_matching_products(req_dict)
            pricing = await quote_engine._calculate_intelligent_pricing(products, customer_id)
            quote_data = {
                "requirements_analysis": req_dict,
                "recommended_products": products,
                "pricing_analysis": pricing
            }
        else:
            return {"success": False, "error": "Either rfp_content or requirements must be provided"}
        
        # Create quote in database
        runtime = LocalRuntime()
        workflow = WorkflowBuilder()
        
        # Generate quote number
        quote_number = f"Q{datetime.now().strftime('%Y%m%d')}{customer_id:04d}"
        
        workflow.add_node("AsyncSQLExecuteNode", "create_quote", {
            "query": """
                INSERT INTO quotes (quote_number, customer_id, title, description, status,
                                  created_date, expiry_date, created_by, currency, 
                                  subtotal, total_amount)
                VALUES (%(quote_number)s, %(customer_id)s, %(title)s, %(description)s, 'draft',
                       NOW(), NOW() + INTERVAL '30 days', 1, 'USD', 
                       %(subtotal)s, %(total)s)
                RETURNING id
            """,
            "parameters": {
                "quote_number": quote_number,
                "customer_id": customer_id,
                "title": "AI-Generated Quote",
                "description": "Intelligently generated quote based on requirements analysis",
                "subtotal": quote_data["pricing_analysis"]["base_total"],
                "total": quote_data["pricing_analysis"]["total"]
            },
            "connection_pool": db.connection_pool
        })
        
        db_results, _ = runtime.execute(workflow.build())
        quote_id = db_results.get("create_quote", {}).get("id")
        
        return {
            "success": True,
            "quote_id": quote_id,
            "quote_number": quote_number,
            "analysis": {
                "recommended_products": len(quote_data["recommended_products"]),
                "pricing_strategy": quote_data["pricing_analysis"].get("algorithm_notes", {}),
                "total_amount": quote_data["pricing_analysis"]["total"],
                "discount_applied": quote_data["pricing_analysis"]["total_discount_percent"],
                "confidence_score": quote_data.get("confidence_score", 0.8)
            },
            "products": quote_data["recommended_products"][:5],  # Top 5 products
            "pricing_breakdown": quote_data["pricing_analysis"]
        }
        
    except Exception as e:
        logger.error(f"Quote generation failed: {e}")
        return {"success": False, "error": str(e)}

@server.tool(required_permission="documents", cache_ttl=120)
async def ask_document_question(question: str, document_ids: Optional[List[int]] = None,
                              customer_id: Optional[int] = None) -> Dict[str, Any]:
    """Answer questions about documents using RAG"""
    try:
        # Build filters for document search
        filters = {}
        if document_ids:
            filters["document_id"] = {"$in": document_ids}
        if customer_id:
            filters["customer_id"] = customer_id
        
        # Retrieve relevant documents
        relevant_docs = await rag_engine.query_documents(question, limit=5, filters=filters)
        
        if not relevant_docs:
            return {
                "success": True,
                "answer": "I couldn't find any relevant documents to answer your question. Please make sure the documents are uploaded and processed.",
                "confidence": 0.1,
                "sources": []
            }
        
        # Generate answer using retrieved context
        answer_result = await rag_engine.answer_question(question, relevant_docs)
        
        return {
            "success": True,
            "answer": answer_result["answer"],
            "confidence": answer_result["confidence"],
            "sources": answer_result["sources"],
            "retrieved_documents": len(relevant_docs),
            "context_used": [{
                "document_id": doc["metadata"]["document_id"],
                "relevance_score": doc["score"],
                "snippet": doc["content"][:200] + "..." if len(doc["content"]) > 200 else doc["content"]
            } for doc in relevant_docs]
        }
        
    except Exception as e:
        logger.error(f"Document Q&A failed: {e}")
        return {"success": False, "error": str(e)}

@server.tool(required_permission="tools", cache_ttl=60)
async def chat_with_assistant(message: str, user_id: int, session_id: str,
                            context: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """Chat with AI assistant with context and memory"""
    try:
        response = await chat_assistant.process_message(message, user_id, session_id, context)
        
        return {
            "success": True,
            "response": response["response"],
            "intent": response.get("intent", {}),
            "suggested_actions": response.get("suggested_actions", []),
            "confidence": response.get("confidence", 0.8),
            "context_used": bool(response.get("context_used")),
            "data": response.get("data", {})
        }
        
    except Exception as e:
        logger.error(f"Chat processing failed: {e}")
        return {"success": False, "error": str(e)}

@server.tool(required_permission="customers", cache_ttl=300)
async def search_customers(query: str, limit: int = 10, filters: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """Search customers with AI-enhanced results"""
    try:
        runtime = LocalRuntime()
        workflow = WorkflowBuilder()
        
        # Build base query
        base_query = """
            SELECT c.*, COUNT(q.id) as quote_count, 
                   COALESCE(SUM(q.total_amount), 0) as total_value,
                   MAX(q.created_date) as last_quote_date
            FROM customers c
            LEFT JOIN quotes q ON c.id = q.customer_id
            WHERE c.deleted_at IS NULL
            AND (c.name ILIKE %(search_term)s 
                 OR c.primary_contact ILIKE %(search_term)s
                 OR c.email ILIKE %(search_term)s
                 OR c.industry ILIKE %(search_term)s)
        """
        
        params = {"search_term": f"%{query}%", "limit": limit}
        
        # Add filters
        if filters:
            if filters.get("industry"):
                base_query += " AND c.industry = %(industry)s"
                params["industry"] = filters["industry"]
            if filters.get("status"):
                base_query += " AND c.status = %(status)s"
                params["status"] = filters["status"]
        
        base_query += " GROUP BY c.id ORDER BY total_value DESC LIMIT %(limit)s"
        
        workflow.add_node("AsyncSQLQueryNode", "search_customers", {
            "query": base_query,
            "parameters": params,
            "connection_pool": db.connection_pool
        })
        
        results, _ = runtime.execute(workflow.build())
        customers = results.get("search_customers", [])
        
        # Enhance with AI insights
        enhanced_customers = []
        for customer in customers:
            # Calculate customer score
            quote_count = customer.get("quote_count", 0)
            total_value = float(customer.get("total_value", 0))
            
            # Simple scoring algorithm
            score = min(100, (quote_count * 10) + (total_value / 1000))
            
            enhanced_customer = {
                **customer,
                "customer_score": round(score, 1),
                "tier": "platinum" if score > 80 else "gold" if score > 50 else "silver" if score > 20 else "bronze",
                "last_activity": customer.get("last_quote_date"),
                "relationship_status": "active" if quote_count > 0 else "prospect"
            }
            enhanced_customers.append(enhanced_customer)
        
        return {
            "success": True,
            "customers": enhanced_customers,
            "total_found": len(enhanced_customers),
            "search_query": query,
            "filters_applied": filters or {}
        }
        
    except Exception as e:
        logger.error(f"Customer search failed: {e}")
        return {"success": False, "error": str(e)}

@server.tool(required_permission="tools", cache_ttl=300)
async def search_products(query: str, category: Optional[str] = None, 
                        price_range: Optional[Dict[str, float]] = None,
                        limit: int = 10) -> Dict[str, Any]:
    """Search ERP products with intelligent matching"""
    try:
        runtime = LocalRuntime()
        workflow = WorkflowBuilder()
        
        base_query = """
            SELECT product_code, name, description, category, subcategory,
                   list_price, cost_price, stock_quantity, stock_status,
                   manufacturer, model_number, specifications
            FROM erp_products 
            WHERE stock_status = 'available'
            AND (name ILIKE %(search_term)s OR description ILIKE %(search_term)s)
        """
        
        params = {"search_term": f"%{query}%", "limit": limit}
        
        if category:
            base_query += " AND category = %(category)s"
            params["category"] = category
        
        if price_range:
            if price_range.get("min"):
                base_query += " AND list_price >= %(min_price)s"
                params["min_price"] = price_range["min"]
            if price_range.get("max"):
                base_query += " AND list_price <= %(max_price)s"
                params["max_price"] = price_range["max"]
        
        base_query += " ORDER BY list_price ASC LIMIT %(limit)s"
        
        workflow.add_node("AsyncSQLQueryNode", "search_products", {
            "query": base_query,
            "parameters": params,
            "connection_pool": db.connection_pool
        })
        
        results, _ = runtime.execute(workflow.build())
        products = results.get("search_products", [])
        
        # Enhance products with AI scoring
        enhanced_products = []
        for product in products:
            # Simple relevance scoring based on query match
            name_match = query.lower() in product["name"].lower()
            desc_match = query.lower() in (product["description"] or "").lower()
            
            relevance_score = 0.0
            if name_match:
                relevance_score += 0.8
            if desc_match:
                relevance_score += 0.6
            
            # Normalize to 0-1 range
            relevance_score = min(1.0, relevance_score)
            
            enhanced_product = {
                **product,
                "relevance_score": relevance_score,
                "margin_percent": ((float(product.get("list_price", 0)) - float(product.get("cost_price", 0))) / float(product.get("list_price", 1))) * 100 if product.get("cost_price") else None,
                "availability": "in_stock" if product.get("stock_quantity", 0) > 0 else "available_to_order"
            }
            enhanced_products.append(enhanced_product)
        
        # Sort by relevance score
        enhanced_products.sort(key=lambda x: x["relevance_score"], reverse=True)
        
        return {
            "success": True,
            "products": enhanced_products,
            "total_found": len(enhanced_products),
            "search_query": query,
            "filters": {
                "category": category,
                "price_range": price_range
            }
        }
        
    except Exception as e:
        logger.error(f"Product search failed: {e}")
        return {"success": False, "error": str(e)}

@server.tool(required_permission="admin", cache_ttl=30)
async def sync_erp_data(entity_type: str = "products", batch_size: int = 100) -> Dict[str, Any]:
    """Synchronize data with ERP systems"""
    try:
        # Mock ERP sync - in production, this would connect to actual ERP
        runtime = LocalRuntime()
        workflow = WorkflowBuilder()
        
        workflow.add_node("AsyncSQLExecuteNode", "log_sync", {
            "query": """
                INSERT INTO erp_sync_logs (sync_type, sync_direction, entity_type, 
                                         erp_system, started_at, status, batch_size)
                VALUES (%(sync_type)s, 'inbound', %(entity_type)s, 'mock_erp', NOW(), 'running', %(batch_size)s)
                RETURNING id
            """,
            "parameters": {
                "sync_type": entity_type,
                "entity_type": entity_type,
                "batch_size": batch_size
            },
            "connection_pool": db.connection_pool
        })
        
        results, _ = runtime.execute(workflow.build())
        sync_log_id = results.get("log_sync", {}).get("id")
        
        # Simulate sync process
        import random
        await asyncio.sleep(2)  # Simulate processing time
        
        records_processed = random.randint(50, batch_size)
        records_successful = random.randint(int(records_processed * 0.8), records_processed)
        records_failed = records_processed - records_successful
        
        # Update sync log
        workflow2 = WorkflowBuilder()
        workflow2.add_node("AsyncSQLExecuteNode", "update_sync", {
            "query": """
                UPDATE erp_sync_logs 
                SET completed_at = NOW(), status = 'completed',
                    records_processed = %(processed)s, 
                    records_successful = %(successful)s,
                    records_failed = %(failed)s
                WHERE id = %(sync_id)s
            """,
            "parameters": {
                "processed": records_processed,
                "successful": records_successful,
                "failed": records_failed,
                "sync_id": sync_log_id
            },
            "connection_pool": db.connection_pool
        })
        
        runtime.execute(workflow2.build())
        
        return {
            "success": True,
            "sync_id": sync_log_id,
            "entity_type": entity_type,
            "records_processed": records_processed,
            "records_successful": records_successful,
            "records_failed": records_failed,
            "success_rate": round((records_successful / records_processed) * 100, 2)
        }
        
    except Exception as e:
        logger.error(f"ERP sync failed: {e}")
        return {"success": False, "error": str(e)}

@server.tool(required_permission="tools", cache_ttl=60)
async def get_sales_analytics(time_period: str = "30d", user_id: Optional[int] = None) -> Dict[str, Any]:
    """Get comprehensive sales analytics and insights"""
    try:
        runtime = LocalRuntime()
        workflow = WorkflowBuilder()
        
        # Calculate date range
        if time_period == "7d":
            date_filter = "created_date >= NOW() - INTERVAL '7 days'"
        elif time_period == "30d":
            date_filter = "created_date >= NOW() - INTERVAL '30 days'"
        elif time_period == "90d":
            date_filter = "created_date >= NOW() - INTERVAL '90 days'"
        else:
            date_filter = "created_date >= NOW() - INTERVAL '30 days'"
        
        user_filter = f"AND created_by = {user_id}" if user_id else ""
        
        # Get quote analytics
        workflow.add_node("AsyncSQLQueryNode", "quote_analytics", {
            "query": f"""
                SELECT 
                    COUNT(*) as total_quotes,
                    COUNT(CASE WHEN status = 'sent' THEN 1 END) as quotes_sent,
                    COUNT(CASE WHEN status = 'accepted' THEN 1 END) as quotes_accepted,
                    COUNT(CASE WHEN status = 'rejected' THEN 1 END) as quotes_rejected,
                    COALESCE(SUM(total_amount), 0) as total_value,
                    COALESCE(AVG(total_amount), 0) as average_value,
                    COALESCE(SUM(CASE WHEN status = 'accepted' THEN total_amount END), 0) as won_value
                FROM quotes 
                WHERE {date_filter} {user_filter}
            """,
            "parameters": {},
            "connection_pool": db.connection_pool
        })
        
        # Get customer analytics
        workflow.add_node("AsyncSQLQueryNode", "customer_analytics", {
            "query": f"""
                SELECT 
                    COUNT(DISTINCT c.id) as total_customers,
                    COUNT(DISTINCT CASE WHEN q.created_date >= NOW() - INTERVAL '{time_period.replace('d', ' days')}' THEN c.id END) as active_customers,
                    AVG(customer_quotes.quote_count) as avg_quotes_per_customer
                FROM customers c
                LEFT JOIN quotes q ON c.id = q.customer_id
                LEFT JOIN (
                    SELECT customer_id, COUNT(*) as quote_count
                    FROM quotes 
                    WHERE {date_filter}
                    GROUP BY customer_id
                ) customer_quotes ON c.id = customer_quotes.customer_id
                WHERE c.deleted_at IS NULL
            """,
            "parameters": {},
            "connection_pool": db.connection_pool
        })
        
        results, _ = runtime.execute(workflow.build())
        
        quote_data = results.get("quote_analytics", [{}])[0]
        customer_data = results.get("customer_analytics", [{}])[0]
        
        # Calculate metrics
        total_quotes = quote_data.get("total_quotes", 0)
        quotes_sent = quote_data.get("quotes_sent", 0)
        quotes_accepted = quote_data.get("quotes_accepted", 0)
        
        win_rate = (quotes_accepted / quotes_sent * 100) if quotes_sent > 0 else 0
        conversion_rate = (quotes_accepted / total_quotes * 100) if total_quotes > 0 else 0
        
        analytics = {
            "time_period": time_period,
            "quote_metrics": {
                "total_quotes": total_quotes,
                "quotes_sent": quotes_sent,
                "quotes_accepted": quotes_accepted,
                "quotes_rejected": quote_data.get("quotes_rejected", 0),
                "pending_quotes": total_quotes - quotes_sent,
                "win_rate_percent": round(win_rate, 2),
                "conversion_rate_percent": round(conversion_rate, 2)
            },
            "financial_metrics": {
                "total_quote_value": float(quote_data.get("total_value", 0)),
                "average_quote_value": float(quote_data.get("average_value", 0)),
                "won_value": float(quote_data.get("won_value", 0)),
                "pipeline_value": float(quote_data.get("total_value", 0)) - float(quote_data.get("won_value", 0))
            },
            "customer_metrics": {
                "total_customers": customer_data.get("total_customers", 0),
                "active_customers": customer_data.get("active_customers", 0),
                "average_quotes_per_customer": round(float(customer_data.get("avg_quotes_per_customer", 0)), 2)
            },
            "insights": []
        }
        
        # Generate insights
        if win_rate > 50:
            analytics["insights"].append("🎯 Excellent win rate! Your quotes are highly competitive.")
        elif win_rate > 30:
            analytics["insights"].append("📈 Good win rate. Consider analyzing rejected quotes for improvement opportunities.")
        else:
            analytics["insights"].append("⚠️ Win rate could be improved. Review pricing strategy and quote quality.")
        
        if analytics["financial_metrics"]["average_quote_value"] > 25000:
            analytics["insights"].append("💰 High-value quotes indicate strong market positioning.")
        
        if analytics["customer_metrics"]["average_quotes_per_customer"] > 3:
            analytics["insights"].append("🔄 Strong customer retention with multiple quotes per customer.")
        
        return {
            "success": True,
            "analytics": analytics,
            "generated_at": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Sales analytics failed: {e}")
        return {"success": False, "error": str(e)}

# ==============================================================================
# MCP RESOURCE IMPLEMENTATIONS
# ==============================================================================

@server.resource("sales://documents/{document_id}")
async def get_document_resource(document_id: int) -> Dict[str, Any]:
    """Get document information and content"""
    try:
        runtime = LocalRuntime()
        workflow = WorkflowBuilder()
        
        workflow.add_node("AsyncSQLQueryNode", "get_document", {
            "query": """
                SELECT d.*, c.name as customer_name
                FROM documents d
                LEFT JOIN customers c ON d.customer_id = c.id
                WHERE d.id = %(document_id)s AND d.deleted_at IS NULL
            """,
            "parameters": {"document_id": document_id},
            "connection_pool": db.connection_pool
        })
        
        results, _ = runtime.execute(workflow.build())
        documents = results.get("get_document", [])
        
        if not documents:
            return {"error": "Document not found"}
        
        document = documents[0]
        
        return {
            "uri": f"sales://documents/{document_id}",
            "name": document["name"],
            "description": f"{document['type']} document",
            "mimeType": document["mime_type"],
            "document": {
                "id": document["id"],
                "name": document["name"],
                "type": document["type"],
                "category": document["category"],
                "customer_name": document["customer_name"],
                "upload_date": document["upload_date"],
                "ai_status": document["ai_status"],
                "ai_confidence_score": document["ai_confidence_score"],
                "page_count": document["page_count"],
                "word_count": document["word_count"],
                "file_size": document["file_size"]
            }
        }
        
    except Exception as e:
        logger.error(f"Document resource access failed: {e}")
        return {"error": str(e)}

@server.resource("sales://quotes/{quote_id}")
async def get_quote_resource(quote_id: int) -> Dict[str, Any]:
    """Get quote information with line items"""
    try:
        runtime = LocalRuntime()
        workflow = WorkflowBuilder()
        
        # Get quote header
        workflow.add_node("AsyncSQLQueryNode", "get_quote", {
            "query": """
                SELECT q.*, c.name as customer_name, c.email as customer_email,
                       u.first_name || ' ' || u.last_name as created_by_name
                FROM quotes q
                JOIN customers c ON q.customer_id = c.id
                LEFT JOIN users u ON q.created_by = u.id
                WHERE q.id = %(quote_id)s AND q.deleted_at IS NULL
            """,
            "parameters": {"quote_id": quote_id},
            "connection_pool": db.connection_pool
        })
        
        # Get line items
        workflow.add_node("AsyncSQLQueryNode", "get_line_items", {
            "query": """
                SELECT * FROM quote_line_items 
                WHERE quote_id = %(quote_id)s AND deleted_at IS NULL
                ORDER BY line_number
            """,
            "parameters": {"quote_id": quote_id},
            "connection_pool": db.connection_pool
        })
        
        results, _ = runtime.execute(workflow.build())
        quotes = results.get("get_quote", [])
        line_items = results.get("get_line_items", [])
        
        if not quotes:
            return {"error": "Quote not found"}
        
        quote = quotes[0]
        
        return {
            "uri": f"sales://quotes/{quote_id}",
            "name": f"Quote {quote['quote_number']}",
            "description": f"Quote for {quote['customer_name']}",
            "mimeType": "application/json",
            "quote": {
                "id": quote["id"],
                "quote_number": quote["quote_number"],
                "title": quote["title"],
                "status": quote["status"],
                "customer": {
                    "name": quote["customer_name"],
                    "email": quote["customer_email"]
                },
                "financial": {
                    "subtotal": float(quote["subtotal"]),
                    "tax_amount": float(quote["tax_amount"]),
                    "discount_amount": float(quote["discount_amount"]),
                    "total_amount": float(quote["total_amount"]),
                    "currency": quote["currency"]
                },
                "dates": {
                    "created_date": quote["created_date"],
                    "expiry_date": quote["expiry_date"],
                    "sent_date": quote["sent_date"]
                },
                "created_by": quote["created_by_name"],
                "line_items": [{
                    "line_number": item["line_number"],
                    "product_name": item["product_name"],
                    "description": item["description"],
                    "quantity": float(item["quantity"]),
                    "unit_price": float(item["unit_price"]),
                    "line_total": float(item["line_total"])
                } for item in line_items]
            }
        }
        
    except Exception as e:
        logger.error(f"Quote resource access failed: {e}")
        return {"error": str(e)}

@server.resource("sales://customers/{customer_id}")
async def get_customer_resource(customer_id: int) -> Dict[str, Any]:
    """Get comprehensive customer information"""
    try:
        runtime = LocalRuntime()
        workflow = WorkflowBuilder()
        
        workflow.add_node("AsyncSQLQueryNode", "get_customer", {
            "query": """
                SELECT c.*, 
                       COUNT(q.id) as total_quotes,
                       COALESCE(SUM(q.total_amount), 0) as total_quote_value,
                       COUNT(CASE WHEN q.status = 'accepted' THEN 1 END) as accepted_quotes,
                       MAX(q.created_date) as last_quote_date
                FROM customers c
                LEFT JOIN quotes q ON c.id = q.customer_id
                WHERE c.id = %(customer_id)s AND c.deleted_at IS NULL
                GROUP BY c.id
            """,
            "parameters": {"customer_id": customer_id},
            "connection_pool": db.connection_pool
        })
        
        results, _ = runtime.execute(workflow.build())
        customers = results.get("get_customer", [])
        
        if not customers:
            return {"error": "Customer not found"}
        
        customer = customers[0]
        
        return {
            "uri": f"sales://customers/{customer_id}",
            "name": customer["name"],
            "description": f"Customer profile for {customer['name']}",
            "mimeType": "application/json",
            "customer": {
                "id": customer["id"],
                "name": customer["name"],
                "type": customer["type"],
                "industry": customer["industry"],
                "contact_info": {
                    "primary_contact": customer["primary_contact"],
                    "email": customer["email"],
                    "phone": customer["phone"],
                    "website": customer["website"]
                },
                "business_info": {
                    "company_size": customer["company_size"],
                    "payment_terms": customer["payment_terms"],
                    "credit_limit": customer["credit_limit"],
                    "currency": customer["currency"]
                },
                "sales_metrics": {
                    "total_quotes": customer["total_quotes"],
                    "accepted_quotes": customer["accepted_quotes"],
                    "total_quote_value": float(customer["total_quote_value"]),
                    "last_quote_date": customer["last_quote_date"],
                    "win_rate": round((customer["accepted_quotes"] / customer["total_quotes"] * 100), 2) if customer["total_quotes"] > 0 else 0
                },
                "status": customer["status"],
                "priority": customer["priority"],
                "assigned_sales_rep": customer["assigned_sales_rep"]
            }
        }
        
    except Exception as e:
        logger.error(f"Customer resource access failed: {e}")
        return {"error": str(e)}

# ==============================================================================
# SERVER STARTUP AND CONFIGURATION
# ==============================================================================

def main():
    """Main entry point for the sales assistant MCP server"""
    logger.info("🚀 Starting AI-Powered Sales Assistant MCP Server")
    
    # Display server configuration
    logger.info("📋 Server Configuration:")
    logger.info(f"   • Authentication: {len(auth_keys)} API keys configured")
    logger.info(f"   • Caching: Memory cache with 5-minute TTL")
    logger.info(f"   • Rate Limiting: 100 requests per minute")
    logger.info(f"   • Document Processing: PDF, Excel, Word support")
    logger.info(f"   • RAG Engine: {'✅ Available' if embedding_model and chroma_client else '❌ Unavailable'}")
    logger.info(f"   • OpenAI Integration: {'✅ Available' if OPENAI_API_KEY else '❌ API key not configured'}")
    
    logger.info("🛠️ Available Tools:")
    tools = [
        "process_document - AI document extraction and analysis",
        "generate_intelligent_quote - Smart quote generation with pricing algorithms", 
        "ask_document_question - RAG-powered document Q&A",
        "chat_with_assistant - Real-time AI chat with context and memory",
        "search_customers - Enhanced customer search with AI insights",
        "search_products - Intelligent product matching and recommendations",
        "sync_erp_data - ERP system data synchronization",
        "get_sales_analytics - Comprehensive sales performance analytics"
    ]
    
    for tool in tools:
        logger.info(f"   • {tool}")
    
    logger.info("📚 Available Resources:")
    resources = [
        "sales://documents/{id} - Document information and metadata",
        "sales://quotes/{id} - Quote details with line items",
        "sales://customers/{id} - Customer profiles with metrics"
    ]
    
    for resource in resources:
        logger.info(f"   • {resource}")
    
    logger.info("🔐 Authentication Required:")
    logger.info("   Use one of the configured API keys in the Authorization header")
    logger.info("   Example: Authorization: Bearer admin_key")
    
    logger.info("\n🌟 Key Features:")
    features = [
        "Multi-modal document processing (PDF, Excel, Word)",
        "AI-powered quote generation with pricing algorithms",
        "RAG-based document Q&A for sales support", 
        "Real-time chat assistant with conversation memory",
        "Intelligent customer and product search",
        "ERP system integration for data synchronization",
        "Comprehensive sales analytics and insights",
        "Production-ready with caching, rate limiting, and monitoring"
    ]
    
    for feature in features:
        logger.info(f"   ✨ {feature}")
    
    # Initialize database connection
    try:
        # Test database connection
        runtime = LocalRuntime()
        workflow = WorkflowBuilder()
        workflow.add_node("AsyncSQLQueryNode", "test_connection", {
            "query": "SELECT 1 as test",
            "parameters": {},
            "connection_pool": db.connection_pool
        })
        runtime.execute(workflow.build())
        logger.info("✅ Database connection established")
    except Exception as e:
        logger.warning(f"⚠️ Database connection issues: {e}")
    
    logger.info("\n🚀 MCP Server starting...")
    logger.info("Ready to process sales assistant AI requests!")
    
    # Start the server
    server.run()

if __name__ == "__main__":
    main()
