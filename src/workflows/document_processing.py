"""
Document Processing Workflow for RFP Analysis

Handles document reading, text extraction, and requirement parsing using Kailash SDK nodes.
Supports multiple document formats: PDF, DOCX, TXT, and raw text input.
"""

import sys
from pathlib import Path
from typing import Dict, List, Any, Optional
import logging

# Add project root to path for imports
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from kailash.workflow.builder import WorkflowBuilder
from kailash.runtime.local import LocalRuntime

logger = logging.getLogger(__name__)

class DocumentProcessingWorkflow:
    """Workflow for processing RFP documents and extracting requirements."""
    
    def __init__(self):
        self.workflow_builder = WorkflowBuilder()
        self.runtime = LocalRuntime()
    
    def create_document_reader_workflow(self) -> Any:
        """Create workflow for reading different document types."""
        workflow = WorkflowBuilder()
        
        # File format detection and reading
        workflow.add_node("PythonCodeNode", "detect_format", {
            "code": """
import os
from pathlib import Path

# Detect file format based on extension
if 'file_path' in locals() and file_path:
    file_path = Path(file_path)
    if not file_path.exists():
        result = {'error': f'File not found: {file_path}', 'format': 'unknown'}
    else:
        extension = file_path.suffix.lower()
        format_map = {
            '.pdf': 'pdf',
            '.docx': 'docx', 
            '.doc': 'doc',
            '.txt': 'text',
            '.rtf': 'rtf'
        }
        detected_format = format_map.get(extension, 'unknown')
        result = {
            'file_path': str(file_path),
            'format': detected_format,
            'size': file_path.stat().st_size if file_path.exists() else 0
        }
elif 'text_content' in locals() and text_content:
    # Direct text input
    result = {
        'format': 'text',
        'text_content': text_content,
        'size': len(text_content)
    }
else:
    result = {'error': 'No file_path or text_content provided', 'format': 'unknown'}
"""
        })
        
        # PDF text extraction
        workflow.add_node("PythonCodeNode", "extract_pdf", {
            "code": """
result = {'text': '', 'success': False, 'error': None}

if format == 'pdf' and 'file_path' in locals():
    try:
        # Try multiple PDF libraries for better compatibility
        text_content = ""
        
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\\n"
        except ImportError:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text_content += page.extract_text() + "\\n"
            except ImportError:
                result['error'] = 'No PDF library available (PyPDF2, pdfplumber)'
        
        if text_content.strip():
            result = {'text': text_content.strip(), 'success': True, 'error': None}
        else:
            result['error'] = 'No text extracted from PDF'
            
    except Exception as e:
        result['error'] = f'PDF extraction failed: {str(e)}'
else:
    result['error'] = 'Not a PDF file or file_path missing'
"""
        })
        
        # DOCX text extraction  
        workflow.add_node("PythonCodeNode", "extract_docx", {
            "code": """
result = {'text': '', 'success': False, 'error': None}

if format == 'docx' and 'file_path' in locals():
    try:
        import docx
        doc = docx.Document(file_path)
        
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + " "
                text_content += "\\n"
        
        if text_content.strip():
            result = {'text': text_content.strip(), 'success': True, 'error': None}
        else:
            result['error'] = 'No text extracted from DOCX'
            
    except ImportError:
        result['error'] = 'python-docx library not available'
    except Exception as e:
        result['error'] = f'DOCX extraction failed: {str(e)}'
else:
    result['error'] = 'Not a DOCX file or file_path missing'
"""
        })
        
        # Plain text file reading
        workflow.add_node("PythonCodeNode", "extract_text", {
            "code": """
result = {'text': '', 'success': False, 'error': None}

if format == 'text' and 'file_path' in locals():
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text_content = file.read()
        result = {'text': text_content, 'success': True, 'error': None}
    except Exception as e:
        try:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                text_content = file.read()
            result = {'text': text_content, 'success': True, 'error': None}
        except Exception as e2:
            result['error'] = f'Text file reading failed: {str(e2)}'
elif format == 'text' and 'text_content' in locals():
    # Direct text input
    result = {'text': text_content, 'success': True, 'error': None}
else:
    result['error'] = 'Not a text file or content missing'
"""
        })
        
        # Text consolidation
        workflow.add_node("PythonCodeNode", "consolidate_text", {
            "code": """
# Consolidate text from different extraction methods
final_text = ""
success = False
errors = []

# Check PDF extraction result
if 'pdf_result' in locals() and pdf_result.get('success'):
    final_text = pdf_result['text']
    success = True
elif 'pdf_result' in locals() and pdf_result.get('error'):
    errors.append(f"PDF: {pdf_result['error']}")

# Check DOCX extraction result  
if 'docx_result' in locals() and docx_result.get('success'):
    final_text = docx_result['text']
    success = True
elif 'docx_result' in locals() and docx_result.get('error'):
    errors.append(f"DOCX: {docx_result['error']}")

# Check text extraction result
if 'text_result' in locals() and text_result.get('success'):
    final_text = text_result['text']
    success = True
elif 'text_result' in locals() and text_result.get('error'):
    errors.append(f"Text: {text_result['error']}")

result = {
    'extracted_text': final_text,
    'success': success,
    'errors': errors,
    'char_count': len(final_text) if final_text else 0
}
"""
        })
        
        return workflow.build()
    
    def create_requirement_parser_workflow(self) -> Any:
        """Create workflow for parsing requirements from extracted text."""
        workflow = WorkflowBuilder()
        
        workflow.add_node("PythonCodeNode", "parse_requirements", {
            "code": """
import re
from dataclasses import dataclass
from typing import Dict, List, Optional

@dataclass
class RequirementItem:
    category: str
    description: str
    quantity: int
    specifications: Dict[str, str]
    keywords: List[str]
    brand: Optional[str] = None
    model: Optional[str] = None
    sku: Optional[str] = None

def extract_keywords_from_description(description: str) -> List[str]:
    if not description:
        return []
    
    # Remove special characters and split into words
    words = re.findall(r'\\b[a-zA-Z0-9]+\\b', description.lower())
    
    # Filter out common stop words
    stop_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'a', 'an'}
    keywords = [word for word in words if word not in stop_words and len(word) > 2]
    
    return list(set(keywords))

def infer_category(description: str, keywords: List[str]) -> str:
    category_mappings = {
        'power tools': ['drill', 'saw', 'driver', 'impact', 'grinder', 'sander', 'dewalt', 'makita', 'milwaukee', 'cordless', '20v', '18v'],
        'safety equipment': ['helmet', 'hat', 'hard hat', 'respirator', 'mask', 'safety', 'vest', 'visibility', '3m', 'n95', 'protection'],
        'lighting': ['light', 'led', 'lamp', 'fixture', 'illumination', 'bulb'],
        'security': ['camera', 'sensor', 'security', 'surveillance', 'alarm', 'detector'],
        'networking': ['cable', 'ethernet', 'network', 'switch', 'router', 'wifi', 'cat6', 'cat6a', 'utp', 'rj45'],
        'power': ['power', 'supply', 'battery', 'transformer', 'ups', 'voltage'],
        'electronics': ['circuit', 'board', 'component', 'resistor', 'capacitor', 'ic'],
        'hardware': ['screw', 'bolt', 'mounting', 'bracket', 'enclosure', 'cabinet'],
        'accessories': ['ties', 'zip', 'cable ties', 'assortment', 'pack', 'kit']
    }
    
    description_lower = description.lower()
    
    for category, category_keywords in category_mappings.items():
        for keyword in category_keywords:
            if keyword in description_lower or keyword in keywords:
                return category.title()
    
    return 'General'

def extract_specifications(description: str) -> Dict[str, str]:
    specs = {}
    
    # Common specification patterns
    spec_patterns = [
        (r'(\\d+)v(?:olt)?', 'voltage'),
        (r'(\\d+)a(?:mp)?', 'current'),
        (r'(\\d+)w(?:att)?', 'power'),
        (r'(\\d+)mm', 'dimension'),
        (r'(\\d+)m(?:eter)?', 'length'),
        (r'ip(\\d+)', 'ip_rating'),
        (r'(\\d+)k', 'temperature'),
    ]
    
    description_lower = description.lower()
    
    for pattern, spec_name in spec_patterns:
        matches = re.findall(pattern, description_lower)
        if matches:
            specs[spec_name] = matches[0]
    
    return specs

def extract_brand_model_from_description(description: str) -> tuple:
    brand_patterns = {
        r'\\b(DEWALT|DeWalt)\\b': 'DEWALT',
        r'\\b(3M)\\b': '3M',  
        r'\\b(MAKITA|Makita)\\b': 'MAKITA',
        r'\\b(MILWAUKEE|Milwaukee)\\b': 'MILWAUKEE',
        r'\\b(BOSCH|Bosch)\\b': 'BOSCH',
        r'\\b(BLACK[\\s&]+DECKER|Black[\\s&]+Decker)\\b': 'BLACK&DECKER',
        r'\\b(RYOBI|Ryobi)\\b': 'RYOBI',
        r'\\b(CRAFTSMAN|Craftsman)\\b': 'CRAFTSMAN',
        r'\\b(HONEYWELL|Honeywell)\\b': 'HONEYWELL',
        r'\\b(MSA)\\b': 'MSA',
        r'\\b(FLUKE|Fluke)\\b': 'FLUKE',
        r'\\b(KLEIN|Klein)\\b': 'KLEIN',
        r'\\b(COMMSCOPE|CommScope)\\b': 'COMMSCOPE'
    }
    
    brand = None
    model = None
    
    # Find brand first
    for pattern, brand_name in brand_patterns.items():
        match = re.search(pattern, description, re.IGNORECASE)
        if match:
            brand = brand_name
            # Look for model after brand
            brand_pos = match.end()
            remaining_text = description[brand_pos:].strip()
            
            # Model pattern: alphanumeric with dashes, typically follows brand
            model_match = re.match(r'^\\s*([A-Z0-9-]{3,15})', remaining_text, re.IGNORECASE)
            if model_match:
                model = model_match.group(1)
            break
    
    return brand, model

# Enhanced patterns for brand/model extraction
brand_model_patterns = [
    # Pattern: "50 units of DEWALT DCD791D2 20V cordless drill"
    r'(\\d+)\\s*(?:units?|pcs?|pieces?)\\s+of\\s+([A-Z][A-Z0-9&]+)\\s+([A-Z0-9-]+)\\s+(.+?)(?:\\n|$|\\.)',
    # Pattern: "100 units of 3M H-700 Series safety helmets"  
    r'(\\d+)\\s*(?:units?|pcs?|pieces?)\\s+of\\s+([A-Z0-9]+)\\s+([A-Z0-9-]+(?:\\s+Series)?)\\s+(.+?)(?:\\n|$|\\.)',
    # Pattern: "50 units DEWALT DCD791D2 20V cordless drill" (no "of")
    r'(\\d+)\\s*(?:units?|pcs?|pieces?)\\s+([A-Z][A-Z0-9&]+)\\s+([A-Z0-9-]+)\\s+(.+?)(?:\\n|$|\\.)',
    # Pattern: "200 meters of Cat6A ethernet cable"
    r'(\\d+)\\s*(?:meters?|feet?|ft)\\s+of\\s+([A-Za-z0-9]+)\\s+(.+?)(?:\\n|$|\\.)',
    # Pattern: "DEWALT DCD791D2: 50 units"
    r'([A-Z][A-Z0-9&]+)\\s+([A-Z0-9-]+):\\s*(\\d+)\\s*(?:units?|pcs?|pieces?)',
]

# Basic requirement patterns
requirement_patterns = [
    r'(\\d+)\\s*(?:units?|pcs?|pieces?)\\s+(?:of\\s+)?(.+?)(?:\\n|$)',
    r'(.+?):\\s*(\\d+)\\s*(?:units?|pcs?|pieces?)',
    r'require\\s+(\\d+)\\s+(.+?)(?:\\n|\\.)',
    r'need\\s+(\\d+)\\s+(.+?)(?:\\n|\\.)',
]

# Parse the extracted text
requirements = []
processed_lines = set()

if 'extracted_text' not in locals() or not extracted_text:
    result = {'requirements': [], 'error': 'No extracted_text provided'}
else:
    lines = extracted_text.split('\\n')
    
    for line in lines:
        line = line.strip()
        if not line or line in processed_lines:
            continue
        processed_lines.add(line)
        
        # Try brand/model patterns first
        found_match = False
        for pattern in brand_model_patterns:
            matches = re.findall(pattern, line, re.IGNORECASE)
            for match in matches:
                try:
                    brand = None
                    model = None
                    
                    if len(match) == 4:  # Quantity, brand, model, description
                        quantity_str, brand, model, description = match
                        quantity = int(quantity_str)
                        full_description = f"{brand} {model} {description}"
                    elif len(match) == 3:
                        if match[2].isdigit():  # Brand, model, quantity (reverse format)
                            brand, model, quantity_str = match
                            quantity = int(quantity_str)
                            full_description = f"{brand} {model}"
                        else:  # Quantity, brand_or_model, description (like Cat6A)
                            quantity_str, brand_or_model, description = match
                            quantity = int(quantity_str)
                            if brand_or_model.upper() in ['CAT6A', 'CAT6', 'CAT5E']:
                                model = brand_or_model.upper()
                                brand = None
                                full_description = f"{brand_or_model} {description}"
                            else:
                                brand = brand_or_model
                                model = None
                                full_description = f"{brand_or_model} {description}"
                    else:
                        continue
                    
                    # Skip invalid brands
                    if brand and brand.lower() in ['of', 'the', 'a', 'an', 'and', 'or', 'but', 'with']:
                        continue
                    
                    # Extract additional information
                    keywords = extract_keywords_from_description(full_description)
                    category = infer_category(full_description, keywords)
                    specifications = extract_specifications(full_description)
                    
                    if brand and model:
                        keywords.extend([brand.lower(), model.lower()])
                    
                    requirement = {
                        'category': category,
                        'description': full_description.strip(),
                        'quantity': quantity,
                        'specifications': specifications,
                        'keywords': list(set(keywords)),
                        'brand': brand.upper() if brand else None,
                        'model': model.upper() if model else None,
                        'sku': model.upper() if model else None
                    }
                    requirements.append(requirement)
                    found_match = True
                    break
                    
                except (ValueError, IndexError):
                    continue
            
            if found_match:
                break
        
        # If no brand/model pattern matched, try general patterns
        if not found_match:
            for pattern in requirement_patterns:
                matches = re.findall(pattern, line, re.IGNORECASE)
                for match in matches:
                    try:
                        if len(match) == 2:
                            quantity_str, description = match
                            try:
                                quantity = int(quantity_str)
                            except ValueError:
                                description, quantity_str = match
                                quantity = int(quantity_str)
                        else:
                            continue
                        
                        # Extract brand/model from description if present
                        brand, model = extract_brand_model_from_description(description)
                        
                        # Extract keywords and specifications
                        keywords = extract_keywords_from_description(description)
                        category = infer_category(description, keywords)
                        specifications = extract_specifications(description)
                        
                        if brand:
                            keywords.append(brand.lower())
                        if model:
                            keywords.append(model.lower())
                        
                        requirement = {
                            'category': category,
                            'description': description.strip(),
                            'quantity': quantity,
                            'specifications': specifications,
                            'keywords': list(set(keywords)),
                            'brand': brand.upper() if brand else None,
                            'model': model.upper() if model else None,
                            'sku': model.upper() if model else None
                        }
                        requirements.append(requirement)
                        found_match = True
                        break
                        
                    except (ValueError, IndexError):
                        continue
                
                if found_match:
                    break

    result = {
        'requirements': requirements,
        'count': len(requirements),
        'success': len(requirements) > 0
    }
"""
        })
        
        return workflow.build()
    
    def create_complete_document_processing_workflow(self) -> Any:
        """Create complete workflow combining document reading and requirement parsing."""
        workflow = WorkflowBuilder()
        
        # Document format detection and reading
        workflow.add_node("PythonCodeNode", "process_document", {
            "code": """
import os
import re
from pathlib import Path
from typing import Dict, List, Optional

# Document processing function
def process_document_input(file_path=None, text_content=None):
    result = {'extracted_text': '', 'success': False, 'errors': []}
    
    if text_content:
        # Direct text input
        result = {'extracted_text': text_content, 'success': True, 'errors': []}
    elif file_path:
        file_path = Path(file_path)
        if not file_path.exists():
            result['errors'] = [f'File not found: {file_path}']
            return result
        
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                # PDF processing
                text = ""
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            text += page.extract_text() + "\\n"
                except ImportError:
                    try:
                        import pdfplumber
                        with pdfplumber.open(file_path) as pdf:
                            for page in pdf.pages:
                                if page.extract_text():
                                    text += page.extract_text() + "\\n"
                    except ImportError:
                        result['errors'] = ['No PDF library available (PyPDF2, pdfplumber)']
                        return result
                
                if text.strip():
                    result = {'extracted_text': text.strip(), 'success': True, 'errors': []}
                else:
                    result['errors'] = ['No text extracted from PDF']
                    
            elif extension in ['.docx', '.doc']:
                # DOCX processing
                try:
                    import docx
                    doc = docx.Document(file_path)
                    
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\\n"
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                            text += "\\n"
                    
                    if text.strip():
                        result = {'extracted_text': text.strip(), 'success': True, 'errors': []}
                    else:
                        result['errors'] = ['No text extracted from DOCX']
                        
                except ImportError:
                    result['errors'] = ['python-docx library not available']
                    
            else:
                # Plain text processing
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        text = file.read()
                    result = {'extracted_text': text, 'success': True, 'errors': []}
                except UnicodeDecodeError:
                    try:
                        with open(file_path, 'r', encoding='latin-1') as file:
                            text = file.read()
                        result = {'extracted_text': text, 'success': True, 'errors': []}
                    except Exception as e:
                        result['errors'] = [f'Text file reading failed: {str(e)}']
                        
        except Exception as e:
            result['errors'] = [f'Document processing failed: {str(e)}']
    else:
        result['errors'] = ['No file_path or text_content provided']
    
    return result

# Process the input
doc_result = process_document_input(
    file_path=file_path if 'file_path' in locals() else None,
    text_content=text_content if 'text_content' in locals() else None
)

result = doc_result
"""
        })
        
        # Requirement parsing
        workflow.add_node("PythonCodeNode", "parse_requirements", {
            "code": """
import re
from typing import Dict, List, Optional

def extract_keywords_from_description(description: str) -> List[str]:
    if not description:
        return []
    
    words = re.findall(r'\\b[a-zA-Z0-9]+\\b', description.lower())
    stop_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'a', 'an'}
    keywords = [word for word in words if word not in stop_words and len(word) > 2]
    
    return list(set(keywords))

def infer_category(description: str, keywords: List[str]) -> str:
    category_mappings = {
        'power tools': ['drill', 'saw', 'driver', 'impact', 'grinder', 'sander', 'dewalt', 'makita', 'milwaukee', 'cordless', '20v', '18v'],
        'safety equipment': ['helmet', 'hat', 'hard hat', 'respirator', 'mask', 'safety', 'vest', 'visibility', '3m', 'n95', 'protection'],
        'lighting': ['light', 'led', 'lamp', 'fixture', 'illumination', 'bulb'],
        'security': ['camera', 'sensor', 'security', 'surveillance', 'alarm', 'detector'],
        'networking': ['cable', 'ethernet', 'network', 'switch', 'router', 'wifi', 'cat6', 'cat6a', 'utp', 'rj45'],
        'power': ['power', 'supply', 'battery', 'transformer', 'ups', 'voltage'],
        'electronics': ['circuit', 'board', 'component', 'resistor', 'capacitor', 'ic'],
        'hardware': ['screw', 'bolt', 'mounting', 'bracket', 'enclosure', 'cabinet'],
        'accessories': ['ties', 'zip', 'cable ties', 'assortment', 'pack', 'kit']
    }
    
    description_lower = description.lower()
    
    for category, category_keywords in category_mappings.items():
        for keyword in category_keywords:
            if keyword in description_lower or keyword in keywords:
                return category.title()
    
    return 'General'

def extract_specifications(description: str) -> Dict[str, str]:
    specs = {}
    spec_patterns = [
        (r'(\\d+)v(?:olt)?', 'voltage'),
        (r'(\\d+)a(?:mp)?', 'current'), 
        (r'(\\d+)w(?:att)?', 'power'),
        (r'(\\d+)mm', 'dimension'),
        (r'(\\d+)m(?:eter)?', 'length'),
        (r'ip(\\d+)', 'ip_rating'),
        (r'(\\d+)k', 'temperature'),
    ]
    
    description_lower = description.lower()
    
    for pattern, spec_name in spec_patterns:
        matches = re.findall(pattern, description_lower)
        if matches:
            specs[spec_name] = matches[0]
    
    return specs

def extract_brand_model_from_description(description: str) -> tuple:
    brand_patterns = {
        r'\\b(DEWALT|DeWalt)\\b': 'DEWALT',
        r'\\b(3M)\\b': '3M',
        r'\\b(MAKITA|Makita)\\b': 'MAKITA',
        r'\\b(MILWAUKEE|Milwaukee)\\b': 'MILWAUKEE',
        r'\\b(BOSCH|Bosch)\\b': 'BOSCH',
        r'\\b(BLACK[\\s&]+DECKER|Black[\\s&]+Decker)\\b': 'BLACK&DECKER',
        r'\\b(RYOBI|Ryobi)\\b': 'RYOBI',
        r'\\b(CRAFTSMAN|Craftsman)\\b': 'CRAFTSMAN',
        r'\\b(HONEYWELL|Honeywell)\\b': 'HONEYWELL',
        r'\\b(MSA)\\b': 'MSA',
        r'\\b(FLUKE|Fluke)\\b': 'FLUKE',
        r'\\b(KLEIN|Klein)\\b': 'KLEIN',
        r'\\b(COMMSCOPE|CommScope)\\b': 'COMMSCOPE'
    }
    
    brand = None
    model = None
    
    for pattern, brand_name in brand_patterns.items():
        match = re.search(pattern, description, re.IGNORECASE)
        if match:
            brand = brand_name
            brand_pos = match.end()
            remaining_text = description[brand_pos:].strip()
            
            model_match = re.match(r'^\\s*([A-Z0-9-]{3,15})', remaining_text, re.IGNORECASE)
            if model_match:
                model = model_match.group(1)
            break
    
    return brand, model

# Parse requirements from extracted text
requirements = []
processed_lines = set()

if not extracted_text:
    result = {'requirements': [], 'error': 'No extracted text to parse'}
else:
    # Enhanced patterns for brand/model extraction
    brand_model_patterns = [
        r'(\\d+)\\s*(?:units?|pcs?|pieces?)\\s+of\\s+([A-Z][A-Z0-9&]+)\\s+([A-Z0-9-]+)\\s+(.+?)(?:\\n|$|\\.)',
        r'(\\d+)\\s*(?:units?|pcs?|pieces?)\\s+of\\s+([A-Z0-9]+)\\s+([A-Z0-9-]+(?:\\s+Series)?)\\s+(.+?)(?:\\n|$|\\.)',
        r'(\\d+)\\s*(?:units?|pcs?|pieces?)\\s+([A-Z][A-Z0-9&]+)\\s+([A-Z0-9-]+)\\s+(.+?)(?:\\n|$|\\.)',
        r'(\\d+)\\s*(?:meters?|feet?|ft)\\s+of\\s+([A-Za-z0-9]+)\\s+(.+?)(?:\\n|$|\\.)',
        r'([A-Z][A-Z0-9&]+)\\s+([A-Z0-9-]+):\\s*(\\d+)\\s*(?:units?|pcs?|pieces?)',
    ]
    
    # Basic requirement patterns
    requirement_patterns = [
        r'(\\d+)\\s*(?:units?|pcs?|pieces?)\\s+(?:of\\s+)?(.+?)(?:\\n|$)',
        r'(.+?):\\s*(\\d+)\\s*(?:units?|pcs?|pieces?)',
        r'require\\s+(\\d+)\\s+(.+?)(?:\\n|\\.)',
        r'need\\s+(\\d+)\\s+(.+?)(?:\\n|\\.)',
    ]
    
    lines = extracted_text.split('\\n')
    
    for line in lines:
        line = line.strip()
        if not line or line in processed_lines:
            continue
        processed_lines.add(line)
        
        found_match = False
        
        # Try brand/model patterns first
        for pattern in brand_model_patterns:
            matches = re.findall(pattern, line, re.IGNORECASE)
            for match in matches:
                try:
                    brand = None
                    model = None
                    
                    if len(match) == 4:
                        quantity_str, brand, model, description = match
                        quantity = int(quantity_str)
                        full_description = f"{brand} {model} {description}"
                    elif len(match) == 3:
                        if match[2].isdigit():
                            brand, model, quantity_str = match
                            quantity = int(quantity_str)
                            full_description = f"{brand} {model}"
                        else:
                            quantity_str, brand_or_model, description = match
                            quantity = int(quantity_str)
                            if brand_or_model.upper() in ['CAT6A', 'CAT6', 'CAT5E']:
                                model = brand_or_model.upper()
                                brand = None
                                full_description = f"{brand_or_model} {description}"
                            else:
                                brand = brand_or_model
                                model = None
                                full_description = f"{brand_or_model} {description}"
                    else:
                        continue
                    
                    if brand and brand.lower() in ['of', 'the', 'a', 'an', 'and', 'or', 'but', 'with']:
                        continue
                    
                    keywords = extract_keywords_from_description(full_description)
                    category = infer_category(full_description, keywords)
                    specifications = extract_specifications(full_description)
                    
                    if brand and model:
                        keywords.extend([brand.lower(), model.lower()])
                    
                    requirement = {
                        'category': category,
                        'description': full_description.strip(),
                        'quantity': quantity,
                        'specifications': specifications,
                        'keywords': list(set(keywords)),
                        'brand': brand.upper() if brand else None,
                        'model': model.upper() if model else None,
                        'sku': model.upper() if model else None
                    }
                    requirements.append(requirement)
                    found_match = True
                    break
                    
                except (ValueError, IndexError):
                    continue
            
            if found_match:
                break
        
        # Try general patterns if no brand/model match
        if not found_match:
            for pattern in requirement_patterns:
                matches = re.findall(pattern, line, re.IGNORECASE)
                for match in matches:
                    try:
                        if len(match) == 2:
                            quantity_str, description = match
                            try:
                                quantity = int(quantity_str)
                            except ValueError:
                                description, quantity_str = match
                                quantity = int(quantity_str)
                        else:
                            continue
                        
                        brand, model = extract_brand_model_from_description(description)
                        keywords = extract_keywords_from_description(description)
                        category = infer_category(description, keywords)
                        specifications = extract_specifications(description)
                        
                        if brand:
                            keywords.append(brand.lower())
                        if model:
                            keywords.append(model.lower())
                        
                        requirement = {
                            'category': category,
                            'description': description.strip(),
                            'quantity': quantity,
                            'specifications': specifications,
                            'keywords': list(set(keywords)),
                            'brand': brand.upper() if brand else None,
                            'model': model.upper() if model else None,
                            'sku': model.upper() if model else None
                        }
                        requirements.append(requirement)
                        found_match = True
                        break
                        
                    except (ValueError, IndexError):
                        continue
                
                if found_match:
                    break

result = {
    'requirements': requirements,
    'count': len(requirements),
    'success': len(requirements) > 0,
    'extracted_text': extracted_text if 'extracted_text' in locals() else '',
    'char_count': len(extracted_text) if 'extracted_text' in locals() else 0
}
"""
        })
        
        # Connect the nodes
        workflow.add_connection("process_document", "extracted_text", "parse_requirements", "extracted_text")
        
        return workflow.build()
    
    def execute_document_processing(self, file_path: Optional[str] = None, 
                                  text_content: Optional[str] = None) -> Dict[str, Any]:
        """Execute the complete document processing workflow."""
        try:
            workflow = self.create_complete_document_processing_workflow()
            
            # Prepare parameters
            params = {}
            if file_path:
                params['file_path'] = file_path
            if text_content:
                params['text_content'] = text_content
            
            # Execute workflow
            results, run_id = self.runtime.execute(workflow, parameters=params)
            
            return {
                'success': True,
                'results': results,
                'run_id': run_id
            }
            
        except Exception as e:
            logger.error(f"Document processing workflow failed: {e}")
            return {
                'success': False,
                'error': str(e),
                'results': None,
                'run_id': None
            }

# Example usage and testing
if __name__ == "__main__":
    # Test with sample RFP text
    sample_rfp = """
    Request for Proposal - Office Equipment
    
    We require the following items:
    
    1. 50 units of DEWALT DCD791D2 20V cordless drill
    2. 100 units of 3M H-700 Series safety helmets  
    3. 25 pieces motion sensors for lighting control
    4. 200 meters of Cat6A ethernet cable
    5. 10 units security cameras for lobby area
    
    Please provide detailed quotation with pricing and specifications.
    """
    
    processor = DocumentProcessingWorkflow()
    result = processor.execute_document_processing(text_content=sample_rfp)
    
    if result['success']:
        print("Document Processing Results:")
        print(f"- Found {result['results']['count']} requirements")
        for req in result['results']['requirements']:
            print(f"  * {req['quantity']}x {req['description']} [{req['category']}]")
            if req['brand']:
                print(f"    Brand: {req['brand']}, Model: {req['model']}")
    else:
        print(f"Processing failed: {result['error']}")