"""
Document Processing Service
Handles document text extraction and requirement analysis using OpenAI
NO MOCK DATA - All processing uses real files and AI
"""

import os
import json
import logging
import time
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime

from openai import AsyncOpenAI
import asyncpg

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document processing with real AI extraction"""

    def __init__(self):
        openai_api_key = os.getenv('OPENAI_API_KEY')
        if not openai_api_key:
            raise ValueError("OPENAI_API_KEY environment variable not set")

        self.openai_client = AsyncOpenAI(api_key=openai_api_key)
        self.openai_model = os.getenv('OPENAI_MODEL', 'gpt-4-turbo-preview')

    async def process_document(self, document_id: int, file_path: str, db_pool: asyncpg.Pool) -> Dict[str, Any]:
        """
        Process uploaded document: extract text, analyze requirements

        Args:
            document_id: Database ID of document
            file_path: Path to uploaded file
            db_pool: Database connection pool

        Returns:
            Dict with extracted requirements and metadata

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: For unsupported file formats
        """
        logger.info(f"Processing document {document_id}: {file_path}")
        start_time = time.time()
        extraction_method = "unknown"

        try:
            # Update status to processing
            await self._update_status(db_pool, document_id, 'processing', None)

            # Extract text from document (returns tuple: text, method)
            extracted_text, extraction_method = await self._extract_text(file_path)

            if not extracted_text or len(extracted_text.strip()) < 10:
                raise ValueError("Insufficient text extracted from document")

            logger.info(f"Extracted {len(extracted_text)} characters from document {document_id} using {extraction_method}")

            # Analyze requirements using OpenAI
            requirements = await self._analyze_requirements(extracted_text)

            # Calculate confidence score
            confidence = self._calculate_extraction_confidence(requirements)

            # Calculate processing time
            processing_time_ms = int((time.time() - start_time) * 1000)

            # Prepare results
            results = {
                'extracted_text_preview': extracted_text[:1000],
                'full_text_length': len(extracted_text),
                'requirements': requirements,
                'processed_at': datetime.utcnow().isoformat(),
                'processor_version': '1.1.0',  # Updated version with enhanced extraction
                'extraction_method': extraction_method,
                'extraction_confidence': confidence,
                'processing_time_ms': processing_time_ms
            }

            # Update database with results and metadata
            await self._update_status(
                db_pool,
                document_id,
                'completed',
                results
            )

            logger.info(
                f"Successfully processed document {document_id}, "
                f"found {len(requirements.get('items', []))} items, "
                f"confidence: {confidence:.2f}, "
                f"method: {extraction_method}, "
                f"time: {processing_time_ms}ms"
            )

            return results

        except Exception as e:
            logger.error(f"Error processing document {document_id}: {str(e)}", exc_info=True)

            # Calculate processing time even for failures
            processing_time_ms = int((time.time() - start_time) * 1000)

            # Update status to failed with error details
            error_data = {
                'error': str(e),
                'error_type': type(e).__name__,
                'failed_at': datetime.utcnow().isoformat(),
                'extraction_method': extraction_method,
                'processing_time_ms': processing_time_ms
            }

            await self._update_status(
                db_pool,
                document_id,
                'failed',
                error_data
            )

            raise

    async def _extract_text(self, file_path: str) -> tuple[str, str]:
        """
        Extract text from document based on file type

        Returns:
            tuple: (extracted_text, extraction_method)
        """

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = Path(file_path).suffix.lower()
        logger.info(f"Extracting text from {file_ext} file")

        if file_ext == '.txt':
            text = await self._extract_txt(file_path)
            return (text, 'txt')
        elif file_ext == '.pdf':
            return await self._extract_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            text = await self._extract_docx(file_path)
            return (text, 'docx_tables')
        elif file_ext in ['.xlsx', '.xls']:
            text = await self._extract_excel(file_path)
            return (text, 'excel_sheets')
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    async def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    async def _extract_pdf(self, file_path: str) -> tuple[str, str]:
        """
        Extract text and tables from PDF file using cascading strategy

        Strategy order:
        1. Camelot (best for PDFs with bordered tables)
        2. PDFPlumber (good for text and borderless tables)
        3. PyPDF2 (fallback for simple text extraction)

        Returns:
            tuple: (extracted_text, extraction_method)
        """

        # Strategy 1: Try Camelot for table extraction first
        try:
            import camelot

            logger.info("Trying Camelot PDF table extraction")
            tables = camelot.read_pdf(file_path, pages='all', flavor='lattice')

            # If we found tables with good accuracy, use Camelot
            if len(tables) > 0 and tables[0].accuracy > 50:
                logger.info(f"Camelot found {len(tables)} tables with good accuracy")
                content_parts = []

                # Also get text from pdfplumber for context
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        for page_num, page in enumerate(pdf.pages, 1):
                            text = page.extract_text()
                            if text:
                                content_parts.append(f"=== PAGE {page_num} TEXT ===")
                                content_parts.append(text)
                except:
                    pass

                # Add Camelot tables
                for i, table in enumerate(tables):
                    content_parts.append(f"\n=== TABLE {i+1} START (Page {table.page}) ===")
                    content_parts.append(self._format_camelot_table(table))
                    content_parts.append(f"=== TABLE {i+1} END ===\n")

                return ('\n'.join(content_parts), 'pdf_camelot')

        except Exception as e:
            logger.info(f"Camelot extraction failed: {e}, trying PDFPlumber")

        # Strategy 2: Try PDFPlumber (good for text and tables)
        try:
            import pdfplumber

            logger.info("Trying PDFPlumber extraction")
            content_parts = []

            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"=== PAGE {page_num} ===")
                        content_parts.append(text)

                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        for table_num, table in enumerate(tables, 1):
                            content_parts.append(f"\n=== TABLE START (Page {page_num}, Table {table_num}) ===")
                            content_parts.append(self._format_list_table(table))
                            content_parts.append("=== TABLE END ===\n")

            result_text = '\n'.join(content_parts)
            if result_text.strip():
                return (result_text, 'pdf_pdfplumber')

        except Exception as e:
            logger.warning(f"PDFPlumber extraction failed: {e}, trying PyPDF2")

        # Strategy 3: Fallback to PyPDF2 for simple text extraction
        try:
            import PyPDF2

            logger.info("Using PyPDF2 fallback")
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"=== PAGE {page_num} ===")
                        text_parts.append(text)

            return ('\n'.join(text_parts), 'pdf_pypdf2')

        except ImportError:
            raise ImportError("No PDF library available (install camelot-py, pdfplumber, or PyPDF2)")

    async def _extract_docx(self, file_path: str) -> str:
        """Extract text AND tables from DOCX file"""
        try:
            from docx import Document

            doc = Document(file_path)

            # Extract all content in order (paragraphs and tables)
            content_parts = []

            # Get all paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    content_parts.append(p.text)

            # Extract all tables (THIS IS THE KEY FIX!)
            for table in doc.tables:
                content_parts.append("\n=== TABLE START ===")

                # Extract table header and rows
                for i, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)

                    if row_data:
                        # First row is usually headers
                        if i == 0:
                            content_parts.append("HEADERS: " + " | ".join(row_data))
                        else:
                            content_parts.append("ROW: " + " | ".join(row_data))

                content_parts.append("=== TABLE END ===\n")

            return '\n'.join(content_parts)

        except ImportError:
            raise ImportError("python-docx library required for DOCX processing (install python-docx)")

    async def _extract_excel(self, file_path: str) -> str:
        """
        Extract text and tables from Excel file

        Handles multiple sheets, merged cells, and formatting
        NO MOCK DATA - Returns real Excel content or raises error
        """
        try:
            from openpyxl import load_workbook

            logger.info(f"Extracting from Excel file: {file_path}")
            wb = load_workbook(file_path, data_only=True)  # data_only=True gets calculated values

            content_parts = []

            for sheet_idx, sheet in enumerate(wb.worksheets, 1):
                content_parts.append(f"\n=== SHEET {sheet_idx}: {sheet.title} ===")
                content_parts.append("=== TABLE START ===")

                # Get all rows with data
                rows_with_data = []
                for row in sheet.iter_rows(values_only=True):
                    # Skip completely empty rows
                    if any(cell is not None and str(cell).strip() for cell in row):
                        rows_with_data.append(row)

                if rows_with_data:
                    # First row is typically headers
                    headers = [str(cell) if cell is not None else "" for cell in rows_with_data[0]]
                    content_parts.append("HEADERS: " + " | ".join(headers))

                    # Remaining rows are data
                    for row in rows_with_data[1:]:
                        row_data = [str(cell) if cell is not None else "" for cell in row]
                        content_parts.append("ROW: " + " | ".join(row_data))
                else:
                    content_parts.append("(Empty sheet)")

                content_parts.append("=== TABLE END ===\n")

            result = '\n'.join(content_parts)

            if not result.strip():
                raise ValueError("No data extracted from Excel file")

            logger.info(f"Successfully extracted {len(rows_with_data)} rows from {len(wb.worksheets)} sheets")
            return result

        except ImportError:
            raise ImportError("openpyxl library required for Excel processing (install openpyxl)")
        except Exception as e:
            logger.error(f"Error extracting Excel file: {e}")
            raise ValueError(f"Failed to extract Excel file: {str(e)}")

    def _format_camelot_table(self, table) -> str:
        """
        Format Camelot table object to text

        Args:
            table: Camelot table object

        Returns:
            Formatted table as text with HEADERS and ROW markers
        """
        try:
            # Convert to DataFrame
            df = table.df

            # First row is headers
            headers = df.iloc[0].tolist()
            content_parts = ["HEADERS: " + " | ".join(str(h) for h in headers)]

            # Remaining rows are data
            for idx in range(1, len(df)):
                row = df.iloc[idx].tolist()
                content_parts.append("ROW: " + " | ".join(str(cell) for cell in row))

            return '\n'.join(content_parts)
        except Exception as e:
            logger.warning(f"Error formatting Camelot table: {e}")
            return "(Table formatting error)"

    def _format_list_table(self, table: List[List]) -> str:
        """
        Format list-of-lists table to text

        Args:
            table: List of lists representing table rows

        Returns:
            Formatted table as text with HEADERS and ROW markers
        """
        if not table or len(table) == 0:
            return "(Empty table)"

        content_parts = []

        # First row is usually headers
        if table[0]:
            headers = [str(cell) if cell else "" for cell in table[0]]
            content_parts.append("HEADERS: " + " | ".join(headers))

        # Remaining rows are data
        for row in table[1:]:
            if row:  # Skip empty rows
                row_data = [str(cell) if cell else "" for cell in row]
                content_parts.append("ROW: " + " | ".join(row_data))

        return '\n'.join(content_parts)

    async def _analyze_requirements(self, text: str) -> Dict[str, Any]:
        """
        Analyze document text using OpenAI to extract structured requirements
        NO FALLBACK DATA - If AI fails, we return error
        """

        prompt = f"""Analyze this Request for Quotation (RFQ) document and extract ALL product line items.

Document text:
{text[:8000]}

IMPORTANT INSTRUCTIONS - Product Extraction:
1. TABLES: Look for tables marked with "=== TABLE START ===" and "=== TABLE END ==="
   - Each "ROW:" line represents ONE product item - extract ALL of them
   - Table columns typically include: Product Description, Brand, Quantity, Unit

2. BULLETED LISTS: Look for product lists with bullet points (-, *, •)
   - Extract each bulleted item as a separate product
   - Parse quantity and description from each line

3. NUMBERED LISTS: Look for numbered lists (1., 2., etc.)
   - Extract each numbered item as a separate product
   - Parse quantity and description from each line

4. TEXT DESCRIPTIONS: Look for product specifications in paragraphs
   - Extract any mentioned product names, quantities, and specifications
   - Look for patterns like "10 pieces of...", "50 units of...", etc.

5. MULTIPLE SHEETS/PAGES: Document may have multiple sheets or pages
   - Extract products from ALL sheets/pages, not just the first
   - Look for "=== SHEET N:" or "=== PAGE N ===" markers

EXTRACTION RULES:
- Extract EVERY single product mentioned - do NOT summarize or group
- If there are 15 line items, return 15 separate items
- Parse quantities from text (e.g., "10 pcs" → quantity: 10, unit: "pcs")
- If quantity is not specified, use 1 as default
- Include brand names in specifications if mentioned

Extract and return a JSON object with this exact structure:
{{
    "customer_name": "company name from the RFQ header",
    "project_name": "project name from the document",
    "deadline": "quotation due date in YYYY-MM-DD format or null",
    "items": [
        {{
            "description": "exact product name/description",
            "quantity": actual_number_or_1_if_not_specified,
            "unit": "unit (pieces, boxes, units, pcs, etc.)",
            "specifications": ["brand name if available", "model", "other specs"],
            "category": "infer from product name (tools, cleaning, safety, electrical, hardware, etc.)"
        }}
        // ... one object for EACH product item found
    ],
    "additional_requirements": ["payment terms", "delivery terms", "warranty", "certifications", etc],
    "delivery_address": "delivery location if mentioned",
    "contact_email": "contact email from document"
}}

CRITICAL: Extract EVERY product from ALL formats (tables, lists, text). No summaries - individual items only.
Return valid JSON only."""

        try:
            response = await self.openai_client.chat.completions.create(
                model=self.openai_model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert at analyzing RFPs and extracting structured product requirements. Always return valid JSON."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )

            requirements_json = response.choices[0].message.content
            requirements = json.loads(requirements_json)

            logger.info(f"OpenAI extracted {len(requirements.get('items', []))} requirement items")

            # Validate we got items
            if not requirements.get('items'):
                logger.warning("No items extracted from document")

            return requirements

        except Exception as e:
            logger.error(f"OpenAI API error: {str(e)}")
            # NO FALLBACK - Let the error propagate
            raise RuntimeError(f"AI requirement extraction failed: {str(e)}")

    def _calculate_extraction_confidence(self, requirements: Dict[str, Any]) -> float:
        """
        Calculate confidence score for extraction based on completeness of extracted data

        Args:
            requirements: Dictionary containing extracted requirements

        Returns:
            Confidence score between 0.0 and 1.0
        """
        if not requirements or not isinstance(requirements, dict):
            return 0.0

        items = requirements.get('items', [])
        if not items:
            return 0.0

        # Base confidence on number of items and their completeness
        total_fields = 0
        filled_fields = 0

        for item in items:
            if not isinstance(item, dict):
                continue

            # Check key fields
            required_fields = ['description', 'quantity']
            optional_fields = ['unit_price', 'specifications', 'category']

            for field in required_fields:
                total_fields += 1
                if item.get(field):
                    filled_fields += 1

            for field in optional_fields:
                total_fields += 0.5  # Optional fields count less
                if item.get(field):
                    filled_fields += 0.5

        if total_fields == 0:
            return 0.0

        base_confidence = filled_fields / total_fields

        # Bonus for having multiple items (more data = more confidence)
        item_count_bonus = min(0.1, len(items) * 0.01)

        # Cap at 0.95 (never 100% confident in AI extraction)
        confidence = min(0.95, base_confidence + item_count_bonus)

        return round(confidence, 2)

    async def _update_status(
        self,
        db_pool: asyncpg.Pool,
        document_id: int,
        status: str,
        extracted_data: Optional[Dict[str, Any]]
    ) -> None:
        """Update document processing status in database"""

        async with db_pool.acquire() as conn:
            await conn.execute("""
                UPDATE documents
                SET ai_status = $1,
                    ai_extracted_data = $2,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = $3
            """, status, json.dumps(extracted_data) if extracted_data else None, document_id)

            logger.info(f"Updated document {document_id} status to: {status}")
