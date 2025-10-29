"""
Enhanced Document Processing Service - Multi-Strategy Implementation
Implements industry best practices from RFP_DOCUMENT_PARSING_RESEARCH.md

FEATURES:
- Phase 1: Specialized parsers (Camelot, PDFPlumber, openpyxl) ✓
- Phase 2: Multi-format parser (Docling/Unstructured.io) ✓
- Phase 3: Vision model fallback (GPT-4 Vision) ✓
- Phase 4: Intelligent orchestration with confidence scoring ✓

NO MOCK DATA - All processing uses real files and AI
"""

import os
import json
import logging
import time
import base64
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime
from io import BytesIO

from openai import AsyncOpenAI
import asyncpg

logger = logging.getLogger(__name__)


class EnhancedDocumentProcessor:
    """Multi-strategy document processor with intelligent fallbacks"""

    def __init__(self):
        openai_api_key = os.getenv('OPENAI_API_KEY')
        if not openai_api_key:
            raise ValueError("OPENAI_API_KEY environment variable not set")

        self.openai_client = AsyncOpenAI(api_key=openai_api_key)
        self.openai_model = os.getenv('OPENAI_MODEL', 'gpt-4-turbo-preview')
        self.vision_model = os.getenv('OPENAI_VISION_MODEL', 'gpt-4o')

        # Strategy configuration
        self.confidence_threshold = float(os.getenv('EXTRACTION_CONFIDENCE_THRESHOLD', '0.85'))
        self.enable_vision_fallback = os.getenv('ENABLE_VISION_FALLBACK', 'true').lower() == 'true'
        self.enable_docling = os.getenv('ENABLE_DOCLING', 'true').lower() == 'true'

    async def process_document(
        self,
        document_id: int,
        file_path: str,
        db_pool: asyncpg.Pool
    ) -> Dict[str, Any]:
        """
        Process document using cascading multi-strategy approach

        Strategy order:
        1. Specialized parsers (best for specific formats)
        2. Multi-format parser (Docling - handles complex layouts)
        3. Vision model (GPT-4 Vision - handles scanned/unusual formats)
        4. Basic text extraction + advanced AI (fallback)

        Args:
            document_id: Database ID of document
            file_path: Path to uploaded file
            db_pool: Database connection pool

        Returns:
            Dict with extracted requirements and metadata
        """
        logger.info(f"Enhanced processing for document {document_id}: {file_path}")
        start_time = time.time()

        try:
            # Update status to processing
            await self._update_status(
                db_pool, document_id, 'processing', None, None, None, None
            )

            # Execute cascading strategy
            best_result = await self._execute_cascading_strategies(file_path)

            # Calculate processing time
            processing_time_ms = int((time.time() - start_time) * 1000)

            # Prepare results
            results = {
                'extracted_text_preview': best_result['text'][:1000],
                'full_text_length': len(best_result['text']),
                'requirements': best_result['requirements'],
                'processed_at': datetime.utcnow().isoformat(),
                'processor_version': '2.0.0',  # Enhanced multi-strategy version
                'extraction_method': best_result['method'],
                'extraction_confidence': best_result['confidence'],
                'processing_time_ms': processing_time_ms,
                'strategies_attempted': best_result.get('strategies_attempted', [])
            }

            # Update database
            await self._update_status(
                db_pool,
                document_id,
                'completed',
                results,
                best_result['method'],
                best_result['confidence'],
                processing_time_ms
            )

            logger.info(
                f"Successfully processed document {document_id}, "
                f"found {len(best_result['requirements'].get('items', []))} items, "
                f"confidence: {best_result['confidence']:.2f}, "
                f"method: {best_result['method']}, "
                f"time: {processing_time_ms}ms"
            )

            return results

        except Exception as e:
            logger.error(f"Error processing document {document_id}: {str(e)}", exc_info=True)

            processing_time_ms = int((time.time() - start_time) * 1000)
            error_data = {
                'error': str(e),
                'error_type': type(e).__name__,
                'failed_at': datetime.utcnow().isoformat(),
                'processing_time_ms': processing_time_ms
            }

            await self._update_status(
                db_pool,
                document_id,
                'failed',
                error_data,
                'error',
                0.0,
                processing_time_ms
            )

            raise

    async def _execute_cascading_strategies(self, file_path: str) -> Dict[str, Any]:
        """
        Execute extraction strategies in order, return best result

        Returns:
            Dict with: text, requirements, method, confidence, strategies_attempted
        """
        strategies = [
            ('specialized', self._try_specialized_parser),
            ('docling', self._try_docling_parser),
            ('vision', self._try_vision_extraction),
            ('fallback', self._try_basic_extraction)
        ]

        best_result = None
        best_score = 0.0
        strategies_attempted = []

        for strategy_name, strategy_func in strategies:
            # Skip disabled strategies
            if strategy_name == 'docling' and not self.enable_docling:
                logger.info(f"Skipping {strategy_name} (disabled in config)")
                continue
            if strategy_name == 'vision' and not self.enable_vision_fallback:
                logger.info(f"Skipping {strategy_name} (disabled in config)")
                continue

            try:
                logger.info(f"Trying strategy: {strategy_name}")
                strategies_attempted.append(strategy_name)

                result = await strategy_func(file_path)
                score = self._calculate_confidence(result)

                logger.info(
                    f"{strategy_name}: "
                    f"extracted {len(result['requirements'].get('items', []))} items, "
                    f"confidence: {score:.2f}"
                )

                if score > best_score:
                    best_result = result
                    best_result['confidence'] = score
                    best_result['method'] = strategy_name
                    best_score = score

                # Early exit if we got high confidence
                if score >= self.confidence_threshold:
                    logger.info(
                        f"Strategy {strategy_name} achieved target confidence "
                        f"({score:.2f} >= {self.confidence_threshold})"
                    )
                    break

            except Exception as e:
                logger.warning(f"Strategy {strategy_name} failed: {e}")
                continue

        if not best_result or best_score < 0.3:
            raise ValueError(
                "Unable to extract meaningful data from document. "
                f"Attempted strategies: {strategies_attempted}, "
                f"best score: {best_score:.2f}"
            )

        best_result['strategies_attempted'] = strategies_attempted
        return best_result

    async def _try_specialized_parser(self, file_path: str) -> Dict[str, Any]:
        """
        Strategy 1: Use specialized parsers for specific formats
        - PDF: Camelot + PDFPlumber
        - Excel: openpyxl
        - Word: python-docx with tables
        """
        file_ext = Path(file_path).suffix.lower()

        if file_ext == '.pdf':
            text, method = await self._extract_pdf_specialized(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            text = await self._extract_excel(file_path)
            method = 'specialized_excel'
        elif file_ext in ['.docx', '.doc']:
            text = await self._extract_docx(file_path)
            method = 'specialized_docx'
        elif file_ext == '.txt':
            text = await self._extract_txt(file_path)
            method = 'specialized_txt'
        else:
            raise ValueError(f"Unsupported format for specialized parser: {file_ext}")

        # Analyze with OpenAI
        requirements = await self._analyze_requirements(text)

        return {
            'text': text,
            'requirements': requirements,
            'method': method
        }

    async def _try_docling_parser(self, file_path: str) -> Dict[str, Any]:
        """
        Strategy 2: Use Docling multi-format parser

        Docling is IBM's open-source document parser optimized for:
        - Complex table layouts
        - Multi-column documents
        - Mixed content (tables + text)
        - All major formats (PDF, Word, Excel, Images)
        """
        try:
            from docling.document_converter import DocumentConverter

            logger.info("Using Docling for comprehensive document parsing")
            converter = DocumentConverter()

            # Convert document to structured format
            result = converter.convert(file_path)

            # Extract tables with spatial awareness
            content_parts = []

            # Get text blocks
            content_parts.append("=== DOCUMENT TEXT ===")
            for text_block in result.text_blocks:
                if text_block.text.strip():
                    content_parts.append(text_block.text)

            # Extract all tables
            if result.tables:
                logger.info(f"Docling found {len(result.tables)} tables")
                for i, table in enumerate(result.tables, 1):
                    content_parts.append(f"\n=== TABLE {i} START ===")
                    content_parts.append(f"Table Location: Page {table.page if hasattr(table, 'page') else 'unknown'}")

                    # Convert table to structured format
                    table_data = table.to_dict() if hasattr(table, 'to_dict') else str(table)
                    content_parts.append(json.dumps(table_data, indent=2))
                    content_parts.append(f"=== TABLE {i} END ===\n")

            text = '\n'.join(content_parts)

            # Analyze with OpenAI
            requirements = await self._analyze_requirements(text)

            return {
                'text': text,
                'requirements': requirements,
                'method': 'docling'
            }

        except ImportError:
            raise ImportError(
                "Docling not installed. Install with: pip install docling\n"
                "See: https://github.com/DS4SD/docling"
            )
        except Exception as e:
            logger.error(f"Docling extraction failed: {e}")
            raise

    async def _try_vision_extraction(self, file_path: str) -> Dict[str, Any]:
        """
        Strategy 3: Use GPT-4 Vision for visual document understanding

        Advantages:
        - Handles scanned documents
        - Understands visual layout (no OCR needed)
        - Works with any visual format
        - Processes entire page as context

        Best for: Complex layouts, scanned RFPs, unusual formats
        """
        logger.info("Using GPT-4 Vision for document extraction")

        # Convert document to images
        images = await self._convert_document_to_images(file_path)

        if not images:
            raise ValueError("Could not convert document to images")

        logger.info(f"Converted document to {len(images)} images")

        # Process each page with GPT-4 Vision
        all_items = []
        extracted_text_parts = []

        for page_num, image_base64 in enumerate(images, 1):
            logger.info(f"Processing page {page_num}/{len(images)} with GPT-4 Vision")

            response = await self.openai_client.chat.completions.create(
                model=self.vision_model,
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": f"""Extract ALL product line items from this RFP page (Page {page_num} of {len(images)}).

Look for:
- Tables with product descriptions, quantities, units
- Bulleted or numbered lists of items
- Text paragraphs describing materials needed
- Headers, footers, and annotations

Return JSON with:
{{
    "page_text": "all text visible on this page",
    "items": [
        {{
            "description": "full product name",
            "quantity": number,
            "unit": "pieces/boxes/meters/etc",
            "specifications": ["brand", "model", "specs"],
            "category": "inferred category"
        }}
    ]
}}

Extract EVERY product mentioned - no summaries."""
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{image_base64}",
                                "detail": "high"  # High detail for tables
                            }
                        }
                    ]
                }],
                temperature=0.1,
                max_tokens=2000,
                response_format={"type": "json_object"}
            )

            page_data = json.loads(response.choices[0].message.content)

            # Collect text and items
            extracted_text_parts.append(f"=== PAGE {page_num} ===")
            extracted_text_parts.append(page_data.get('page_text', ''))

            page_items = page_data.get('items', [])
            all_items.extend(page_items)

            logger.info(f"Page {page_num}: extracted {len(page_items)} items")

        # Combine all results
        full_text = '\n'.join(extracted_text_parts)

        requirements = {
            'customer_name': None,
            'project_name': None,
            'deadline': None,
            'items': all_items,
            'additional_requirements': [],
            'delivery_address': None,
            'contact_email': None
        }

        # Try to extract metadata from first page
        if all_items and extracted_text_parts:
            # Use AI to extract metadata from combined text
            metadata_prompt = f"""From this RFP text, extract metadata:

{full_text[:2000]}

Return JSON with:
{{
    "customer_name": "company name or null",
    "project_name": "project name or null",
    "deadline": "date in YYYY-MM-DD or null",
    "delivery_address": "address or null",
    "contact_email": "email or null"
}}"""

            try:
                meta_response = await self.openai_client.chat.completions.create(
                    model=self.openai_model,
                    messages=[{"role": "user", "content": metadata_prompt}],
                    temperature=0.1,
                    response_format={"type": "json_object"}
                )

                metadata = json.loads(meta_response.choices[0].message.content)
                requirements.update(metadata)
            except:
                pass  # Metadata extraction is optional

        return {
            'text': full_text,
            'requirements': requirements,
            'method': 'vision',
            'pages_processed': len(images)
        }

    async def _try_basic_extraction(self, file_path: str) -> Dict[str, Any]:
        """
        Strategy 4: Basic text extraction + advanced AI prompt
        Fallback when all other strategies fail
        """
        logger.info("Using basic extraction fallback")

        # Simple text extraction
        file_ext = Path(file_path).suffix.lower()

        if file_ext == '.txt':
            text = await self._extract_txt(file_path)
        elif file_ext == '.pdf':
            text, _ = await self._extract_pdf_basic(file_path)
        elif file_ext in ['.docx', '.doc']:
            text = await self._extract_docx_basic(file_path)
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

        # Use advanced AI prompt to extract requirements
        requirements = await self._analyze_requirements(text)

        return {
            'text': text,
            'requirements': requirements,
            'method': 'basic_fallback'
        }

    # ==================== Helper Methods ====================

    async def _convert_document_to_images(self, file_path: str) -> List[str]:
        """
        Convert document to high-quality images for vision processing

        Returns:
            List of base64-encoded PNG images (one per page)
        """
        file_ext = Path(file_path).suffix.lower()

        try:
            if file_ext == '.pdf':
                from pdf2image import convert_from_path

                # Convert PDF to images at 200 DPI for quality
                images = convert_from_path(file_path, dpi=200, fmt='PNG')

            elif file_ext in ['.docx', '.doc']:
                # Convert Word to PDF first, then to images
                # This requires libreoffice or similar
                try:
                    import subprocess
                    import tempfile

                    with tempfile.TemporaryDirectory() as tmpdir:
                        # Convert to PDF using libreoffice
                        subprocess.run([
                            'libreoffice',
                            '--headless',
                            '--convert-to', 'pdf',
                            '--outdir', tmpdir,
                            file_path
                        ], check=True, capture_output=True)

                        pdf_path = Path(tmpdir) / f"{Path(file_path).stem}.pdf"

                        from pdf2image import convert_from_path
                        images = convert_from_path(str(pdf_path), dpi=200, fmt='PNG')
                except Exception as e:
                    logger.warning(f"Could not convert Word to images: {e}")
                    return []

            elif file_ext in ['.xlsx', '.xls']:
                # For Excel, we'd need to screenshot each sheet
                # This is complex - for now, skip vision for Excel
                logger.warning("Vision extraction not supported for Excel files")
                return []

            elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                # Already an image
                from PIL import Image
                images = [Image.open(file_path)]

            else:
                logger.warning(f"Vision extraction not supported for {file_ext}")
                return []

            # Convert images to base64
            base64_images = []
            for img in images:
                buffer = BytesIO()
                img.save(buffer, format='PNG')
                buffer.seek(0)
                img_base64 = base64.b64encode(buffer.read()).decode('utf-8')
                base64_images.append(img_base64)

            return base64_images

        except ImportError as e:
            logger.error(
                f"Missing dependency for image conversion: {e}\n"
                "Install with: pip install pdf2image pillow"
            )
            return []
        except Exception as e:
            logger.error(f"Error converting document to images: {e}")
            return []

    def _calculate_confidence(self, result: Dict[str, Any]) -> float:
        """
        Calculate confidence score for extraction result

        Scoring criteria:
        - Number of items found (more is better)
        - Data completeness (all required fields present)
        - Text length (sufficient content extracted)
        - Field quality (valid quantities, meaningful descriptions)
        """
        requirements = result.get('requirements', {})
        items = requirements.get('items', [])
        text = result.get('text', '')

        if len(items) == 0:
            return 0.0

        score = 0.5  # Base score for finding any items

        # Bonus for multiple items (indicates real extraction)
        if len(items) >= 3:
            score += 0.15
        if len(items) >= 5:
            score += 0.10
        if len(items) >= 10:
            score += 0.05

        # Check data completeness for each item
        complete_items = 0
        for item in items:
            has_description = bool(item.get('description') and len(item.get('description', '')) > 10)
            has_quantity = bool(item.get('quantity') and item.get('quantity') > 0)
            has_valid_fields = has_description and has_quantity

            if has_valid_fields:
                complete_items += 1

        completeness_ratio = complete_items / len(items) if items else 0
        score += completeness_ratio * 0.15

        # Bonus for sufficient text extraction
        if len(text) > 500:
            score += 0.05

        return min(score, 1.0)

    # ==================== Extraction Methods (from original) ====================

    async def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    async def _extract_pdf_specialized(self, file_path: str) -> Tuple[str, str]:
        """Extract PDF with Camelot + PDFPlumber (Phase 1 implementation)"""
        # Try Camelot first
        try:
            import camelot
            tables = camelot.read_pdf(file_path, pages='all', flavor='lattice')

            if len(tables) > 0 and tables[0].accuracy > 50:
                content_parts = []

                # Get text from pdfplumber
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        for page_num, page in enumerate(pdf.pages, 1):
                            text = page.extract_text()
                            if text:
                                content_parts.append(f"=== PAGE {page_num} TEXT ===")
                                content_parts.append(text)
                except:
                    pass

                # Add tables
                for i, table in enumerate(tables):
                    content_parts.append(f"\n=== TABLE {i+1} (Page {table.page}) ===")
                    content_parts.append(self._format_camelot_table(table))
                    content_parts.append(f"=== TABLE {i+1} END ===\n")

                return ('\n'.join(content_parts), 'pdf_camelot')
        except Exception as e:
            logger.info(f"Camelot failed: {e}")

        # Try PDFPlumber
        try:
            import pdfplumber
            content_parts = []

            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"=== PAGE {page_num} ===")
                        content_parts.append(text)

                    tables = page.extract_tables()
                    if tables:
                        for table_num, table in enumerate(tables, 1):
                            content_parts.append(f"\n=== TABLE (Page {page_num}, Table {table_num}) ===")
                            content_parts.append(self._format_list_table(table))
                            content_parts.append("=== TABLE END ===\n")

            return ('\n'.join(content_parts), 'pdf_pdfplumber')
        except Exception as e:
            logger.warning(f"PDFPlumber failed: {e}")

        # Fallback to PyPDF2
        import PyPDF2
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text:
                    text_parts.append(f"=== PAGE {page_num} ===")
                    text_parts.append(text)

        return ('\n'.join(text_parts), 'pdf_pypdf2')

    async def _extract_pdf_basic(self, file_path: str) -> Tuple[str, str]:
        """Basic PDF text extraction for fallback"""
        import PyPDF2
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
        return ('\n'.join(text_parts), 'pdf_basic')

    async def _extract_docx(self, file_path: str) -> str:
        """Extract text AND tables from DOCX"""
        from docx import Document
        doc = Document(file_path)
        content_parts = []

        # Paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                content_parts.append(p.text)

        # Tables
        for table in doc.tables:
            content_parts.append("\n=== TABLE START ===")
            for i, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    prefix = "HEADERS: " if i == 0 else "ROW: "
                    content_parts.append(prefix + " | ".join(row_data))
            content_parts.append("=== TABLE END ===\n")

        return '\n'.join(content_parts)

    async def _extract_docx_basic(self, file_path: str) -> str:
        """Basic DOCX extraction (paragraphs only)"""
        from docx import Document
        doc = Document(file_path)
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())

    async def _extract_excel(self, file_path: str) -> str:
        """Extract from Excel with multiple sheets"""
        from openpyxl import load_workbook
        wb = load_workbook(file_path, data_only=True)
        content_parts = []

        for sheet in wb.worksheets:
            content_parts.append(f"\n=== SHEET: {sheet.title} ===")
            content_parts.append("=== TABLE START ===")

            rows_with_data = [
                row for row in sheet.iter_rows(values_only=True)
                if any(cell is not None and str(cell).strip() for cell in row)
            ]

            if rows_with_data:
                headers = [str(cell) if cell else "" for cell in rows_with_data[0]]
                content_parts.append("HEADERS: " + " | ".join(headers))

                for row in rows_with_data[1:]:
                    row_data = [str(cell) if cell else "" for cell in row]
                    content_parts.append("ROW: " + " | ".join(row_data))

            content_parts.append("=== TABLE END ===\n")

        return '\n'.join(content_parts)

    def _format_camelot_table(self, table) -> str:
        """Format Camelot table to text"""
        try:
            df = table.df
            headers = df.iloc[0].tolist()
            content_parts = ["HEADERS: " + " | ".join(str(h) for h in headers)]

            for idx in range(1, len(df)):
                row = df.iloc[idx].tolist()
                content_parts.append("ROW: " + " | ".join(str(cell) for cell in row))

            return '\n'.join(content_parts)
        except Exception as e:
            logger.warning(f"Error formatting Camelot table: {e}")
            return "(Table formatting error)"

    def _format_list_table(self, table: List[List]) -> str:
        """Format list-of-lists table to text"""
        if not table:
            return "(Empty table)"

        content_parts = []
        if table[0]:
            headers = [str(cell) if cell else "" for cell in table[0]]
            content_parts.append("HEADERS: " + " | ".join(headers))

        for row in table[1:]:
            if row:
                row_data = [str(cell) if cell else "" for cell in row]
                content_parts.append("ROW: " + " | ".join(row_data))

        return '\n'.join(content_parts)

    async def _analyze_requirements(self, text: str) -> Dict[str, Any]:
        """
        Analyze document text using OpenAI to extract structured requirements
        Enhanced prompt for better extraction across all document types
        """
        prompt = f"""Analyze this Request for Quotation (RFQ) document and extract ALL product line items.

Document text:
{text[:8000]}

IMPORTANT INSTRUCTIONS - Product Extraction:
1. TABLES: Look for tables marked with "=== TABLE START ===" and "=== TABLE END ==="
   - Each "ROW:" line represents ONE product item - extract ALL of them
   - Table columns typically include: Product Description, Brand, Quantity, Unit

2. BULLETED LISTS: Look for product lists with bullet points (-, *, •)
   - Extract each bulleted item as a separate product
   - Parse quantity and description from each line

3. NUMBERED LISTS: Look for numbered lists (1., 2., etc.)
   - Extract each numbered item as a separate product
   - Parse quantity and description from each line

4. TEXT DESCRIPTIONS: Look for product specifications in paragraphs
   - Extract any mentioned product names, quantities, and specifications
   - Look for patterns like "10 pieces of...", "50 units of...", etc.

5. MULTIPLE SHEETS/PAGES: Document may have multiple sheets or pages
   - Extract products from ALL sheets/pages, not just the first
   - Look for "=== SHEET N:" or "=== PAGE N ===" markers

EXTRACTION RULES:
- Extract EVERY single product mentioned - do NOT summarize or group
- If there are 15 line items, return 15 separate items
- Parse quantities from text (e.g., "10 pcs" → quantity: 10, unit: "pcs")
- If quantity is not specified, use 1 as default
- Include brand names in specifications if mentioned

Extract and return a JSON object with this exact structure:
{{
    "customer_name": "company name from the RFQ header",
    "project_name": "project name from the document",
    "deadline": "quotation due date in YYYY-MM-DD format or null",
    "items": [
        {{
            "description": "exact product name/description",
            "quantity": actual_number_or_1_if_not_specified,
            "unit": "unit (pieces, boxes, units, pcs, etc.)",
            "specifications": ["brand name if available", "model", "other specs"],
            "category": "infer from product name (tools, cleaning, safety, electrical, hardware, etc.)"
        }}
        // ... one object for EACH product item found
    ],
    "additional_requirements": ["payment terms", "delivery terms", "warranty", "certifications", etc],
    "delivery_address": "delivery location if mentioned",
    "contact_email": "contact email from document"
}}

CRITICAL: Extract EVERY product from ALL formats (tables, lists, text). No summaries - individual items only.
Return valid JSON only."""

        try:
            response = await self.openai_client.chat.completions.create(
                model=self.openai_model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert at analyzing RFPs and extracting structured product requirements. Always return valid JSON."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )

            requirements_json = response.choices[0].message.content
            requirements = json.loads(requirements_json)

            logger.info(f"OpenAI extracted {len(requirements.get('items', []))} requirement items")

            if not requirements.get('items'):
                logger.warning("No items extracted from document")

            return requirements

        except Exception as e:
            logger.error(f"OpenAI API error: {str(e)}")
            raise RuntimeError(f"AI requirement extraction failed: {str(e)}")

    async def _update_status(
        self,
        db_pool: asyncpg.Pool,
        document_id: int,
        status: str,
        extracted_data: Optional[Dict[str, Any]],
        extraction_method: Optional[str],
        extraction_confidence: Optional[float],
        processing_time_ms: Optional[int]
    ) -> None:
        """Update document processing status in database with analytics"""
        async with db_pool.acquire() as conn:
            await conn.execute("""
                UPDATE documents
                SET ai_status = $1,
                    ai_extracted_data = $2,
                    extraction_method = $3,
                    extraction_confidence = $4,
                    processing_time_ms = $5,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = $6
            """,
            status,
            json.dumps(extracted_data) if extracted_data else None,
            extraction_method,
            extraction_confidence,
            processing_time_ms,
            document_id)

            logger.info(f"Updated document {document_id} status to: {status}")
