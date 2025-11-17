"""
Document Extraction API - Extract job details from uploaded PDF/DOCX files

Provides endpoints for uploading job description documents and extracting
structured job information using OpenAI.
"""

import logging
import io
from typing import Optional, List
from fastapi import APIRouter, UploadFile, File, Depends, HTTPException, status
from pydantic import BaseModel, Field
import PyPDF2
from docx import Document as DocxDocument

from openai import OpenAI

from job_pricing.core.config import get_settings
from job_pricing.models.auth import User
from job_pricing.api.dependencies.auth import get_current_active_user

logger = logging.getLogger(__name__)
settings = get_settings()
router = APIRouter()


# Request/Response Models
class ExtractedJobDetails(BaseModel):
    """Extracted job details from document"""
    job_title: Optional[str] = None
    job_summary: Optional[str] = None
    key_responsibilities: List[str] = []
    skills_required: List[str] = []
    qualifications: List[str] = []
    experience_required: Optional[str] = None
    department: Optional[str] = None
    employment_type: Optional[str] = Field(None, description="Full-time, Part-time, Contract, etc.")
    job_family: Optional[str] = Field(None, description="Job function category")
    portfolio: Optional[str] = Field(None, description="TPC portfolio/business entity")
    raw_text: str = Field(..., description="Full extracted text from document")


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file.

    Args:
        file_content: PDF file content as bytes

    Returns:
        Extracted text from all pages

    Raises:
        ValueError: If PDF cannot be read
    """
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text_parts = []

        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return "\n\n".join(text_parts)

    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file.

    Args:
        file_content: DOCX file content as bytes

    Returns:
        Extracted text from document

    Raises:
        ValueError: If DOCX cannot be read
    """
    try:
        doc = DocxDocument(io.BytesIO(file_content))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n\n".join(text_parts)

    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def get_openai_client():
    """Get OpenAI client"""
    if not settings.OPENAI_API_KEY:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="OpenAI API key not configured"
        )
    return OpenAI(api_key=settings.OPENAI_API_KEY)


async def extract_job_details_with_ai(document_text: str) -> ExtractedJobDetails:
    """
    Use OpenAI to extract structured job details from document text.

    Args:
        document_text: Raw text extracted from document

    Returns:
        Structured job details
    """
    client = get_openai_client()

    system_prompt = """You are an expert HR analyst specializing in job description analysis.
Extract structured job information from the provided document text.

Focus on identifying:
1. Job Title - The main role title
2. Job Summary - A concise 2-3 sentence overview of the role
3. Key Responsibilities - List of main accountabilities (bullet points)
4. Skills Required - Technical and soft skills needed
5. Qualifications - Educational and professional requirements
6. Experience Required - Years and type of experience
7. Department - Which department/business unit
8. Employment Type - Full-time, Part-time, Contract, Temporary, etc.
9. Job Family - Job function category (e.g., "Human Resources Management", "Total Rewards", "Talent Acquisition", "Finance", "Operations", etc.)
10. Portfolio - Business entity/portfolio (e.g., "TPC HQ", "TPC Capital", "TPC Facilities", "TPC Health", etc.)

Return ONLY valid JSON in this exact format:
{
  "job_title": "extracted title or null",
  "job_summary": "concise summary paragraph",
  "key_responsibilities": ["resp1", "resp2", ...],
  "skills_required": ["skill1", "skill2", ...],
  "qualifications": ["qual1", "qual2", ...],
  "experience_required": "X years in Y field",
  "department": "department name or null",
  "employment_type": "Full-time or Part-time or Contract or null",
  "job_family": "job function category or null",
  "portfolio": "TPC entity/portfolio or null"
}

If any field cannot be determined, use null or empty array.
For employment_type, default to "Full-time" if not explicitly stated.
For job_family, infer from the job title and responsibilities."""

    user_prompt = f"""Extract job details from this document:

{document_text}

Remember to return ONLY the JSON object, no other text."""

    try:
        response = client.chat.completions.create(
            model=settings.OPENAI_MODEL_DEFAULT,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.1,
            max_tokens=2000
        )

        ai_response = response.choices[0].message.content

        # Parse JSON response
        import json
        extracted_data = json.loads(ai_response)

        return ExtractedJobDetails(
            job_title=extracted_data.get("job_title"),
            job_summary=extracted_data.get("job_summary"),
            key_responsibilities=extracted_data.get("key_responsibilities", []),
            skills_required=extracted_data.get("skills_required", []),
            qualifications=extracted_data.get("qualifications", []),
            experience_required=extracted_data.get("experience_required"),
            department=extracted_data.get("department"),
            employment_type=extracted_data.get("employment_type"),
            job_family=extracted_data.get("job_family"),
            portfolio=extracted_data.get("portfolio"),
            raw_text=document_text
        )

    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse AI response as JSON: {e}")
        logger.error(f"AI Response: {response}")
        raise ValueError("AI returned invalid JSON response")
    except Exception as e:
        logger.error(f"Error in AI extraction: {e}")
        raise


@router.post(
    "/extract",
    response_model=ExtractedJobDetails,
    summary="Extract job details from uploaded document",
    description="Upload a PDF or DOCX file and extract structured job information using AI",
    responses={
        200: {"description": "Job details extracted successfully"},
        400: {"description": "Invalid file format or content"},
        401: {"description": "Not authenticated"},
        413: {"description": "File too large"},
        500: {"description": "Extraction failed"}
    }
)
async def extract_document(
    file: UploadFile = File(..., description="PDF or DOCX file containing job description"),
    current_user: User = Depends(get_current_active_user)
):
    """
    Extract structured job details from uploaded PDF or DOCX document.

    The system will:
    1. Extract text from the document
    2. Use AI to analyze and structure the content
    3. Return job title, summary, responsibilities, skills, etc.

    Supported formats: PDF, DOCX
    Max file size: 10MB

    Returns:
        ExtractedJobDetails with structured job information
    """
    logger.info(f"Document extraction requested by user {current_user.email}")

    # Validate file type
    allowed_types = ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
    if file.content_type not in allowed_types and not (file.filename.endswith('.pdf') or file.filename.endswith('.docx')):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid file type. Only PDF and DOCX files are supported."
        )

    # Read file content
    try:
        content = await file.read()
    except Exception as e:
        logger.error(f"Error reading file: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to read uploaded file"
        )

    # Check file size (10MB limit)
    max_size = settings.MAX_UPLOAD_SIZE  # 10MB from config
    if len(content) > max_size:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large. Maximum size is {max_size / 1024 / 1024}MB"
        )

    # Extract text based on file type
    try:
        if file.filename.endswith('.pdf'):
            logger.info(f"Extracting text from PDF: {file.filename}")
            document_text = extract_text_from_pdf(content)
        elif file.filename.endswith('.docx'):
            logger.info(f"Extracting text from DOCX: {file.filename}")
            document_text = extract_text_from_docx(content)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Unsupported file format"
            )

        if not document_text.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No text could be extracted from the document"
            )

        logger.info(f"Extracted {len(document_text)} characters from document")

    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        logger.error(f"Unexpected error during text extraction: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to extract text from document"
        )

    # Use AI to extract structured job details
    try:
        logger.info("Analyzing document with AI...")
        extracted_details = await extract_job_details_with_ai(document_text)

        logger.info(
            f"Successfully extracted job details: "
            f"title={extracted_details.job_title}, "
            f"skills={len(extracted_details.skills_required)}, "
            f"responsibilities={len(extracted_details.key_responsibilities)}"
        )

        return extracted_details

    except Exception as e:
        logger.error(f"Error extracting job details with AI: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to extract job details: {str(e)}"
        )
